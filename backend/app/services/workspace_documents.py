"""受控 workspace 文档的导入、只读解析与精确定位服务。

这里刻意把“文件格式解析”放在 Tool/Agent 之外：所有调用方都只能得到 workspace 内的
相对文件名、受限文本和来源定位，不能借 PDF/DOCX/图片支持绕过原有的文件边界。
"""

from __future__ import annotations

import base64
from collections import OrderedDict
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
import re
from threading import RLock
from typing import Callable, Iterable, Literal

from app.core.config import settings
from app.services.ocr_adapter import OcrAdapter, OcrAdapterError, OcrDocumentResult
from app.schemas.workspace import (
    WorkspaceDocumentInfo,
    WorkspaceDocumentPreviewResponse,
    WorkspaceDocumentSearchMatch,
    WorkspaceDocumentSearchResponse,
)

try:  # 依赖缺失时仍由服务层给出可理解的安装提示，而不是让导入接口 500。
    import fitz
except ImportError:  # pragma: no cover - 正常 requirements 环境会安装此依赖。
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - 正常 requirements 环境会安装此依赖。
    DocxDocument = None


TEXT_DOCUMENT_SUFFIXES = {".txt", ".md", ".markdown"}
# 图片只作为本地 OCR 的受控输入，不向任何模型或接口暴露客户本机路径。它们仍沿用既有
# Base64 导入与二进制体积上限，而不是额外新增一条绕开 workspace 的文件通道。
IMAGE_DOCUMENT_SUFFIXES = {".png", ".jpg", ".jpeg"}
BINARY_DOCUMENT_SUFFIXES = {".pdf", ".docx"} | IMAGE_DOCUMENT_SUFFIXES
ALLOWED_DOCUMENT_SUFFIXES = TEXT_DOCUMENT_SUFFIXES | BINARY_DOCUMENT_SUFFIXES
MAX_WORKSPACE_DOCUMENT_BYTES = 1_000_000
MAX_WORKSPACE_BINARY_DOCUMENT_BYTES = 10_000_000
# 压缩包格式可以用很小的源文件展开出异常大的文本。解析结果也必须限幅，避免它变成模型
# 上下文、缓存和任务审计的内存放大入口。
MAX_EXTRACTED_DOCUMENT_CHARS = 1_000_000
WORKSPACE_DIR_NAME = "workspaces"
_PARSE_CACHE_MAX_ENTRIES = 16

SourceKind = Literal["line", "page", "paragraph", "table", "region", "mixed"]


class WorkspaceDocumentError(ValueError):
    pass


@dataclass(frozen=True)
class _DocumentSegment:
    """一段可回溯的已提取文本在完整文本中的字符范围。"""

    start_char: int
    end_char: int
    source_kind: SourceKind
    source_locator: str
    ordinal: int


@dataclass(frozen=True)
class _ParsedWorkspaceDocument:
    """解析器返回的内部表示，绝对路径永远不进入该对象。"""

    text: str
    document_type: Literal["text", "pdf", "docx", "image"]
    segments: tuple[_DocumentSegment, ...] = ()
    # OCR 页级统计不保存文字、坐标、图片或本机路径。它让知识库索引能说明“已有多少页可用”而
    # 不必为了展示状态再次读取材料；非 OCR 文档固定保留零值。
    ocr_page_count: int = 0
    ocr_completed_page_count: int = 0
    ocr_failed_page_count: int = 0
    ocr_retried_page_count: int = 0


# 知识库 K1/K7 复用同一套 TXT/Markdown/PDF/DOCX/图片解析和来源定位规则。别名明确这是后端
# 内部稳定契约，而不是新的 API 数据模型；其中从不携带原始文件路径。
ControlledDocumentSegment = _DocumentSegment
ParsedControlledDocument = _ParsedWorkspaceDocument


# 以“文件名 + mtime_ns + size”为键的 LRU 缓存。文件内容变更会自然 miss；不会主动扫描目录
# 或保存跨进程状态，因此不可能向用户展示旧版本解析结果超过一次当前请求。
_parse_cache: OrderedDict[tuple[str, int, int], _ParsedWorkspaceDocument] = OrderedDict()
_parse_cache_lock = RLock()


def workspace_documents_dir() -> Path:
    return settings.data_dir / WORKSPACE_DIR_NAME


def list_workspace_documents() -> list[WorkspaceDocumentInfo]:
    """列出用户显式导入的顶层 workspace 文档，不解析正文。"""

    root = workspace_documents_dir()
    if not root.exists():
        return []

    documents: list[WorkspaceDocumentInfo] = []
    for path in sorted(root.iterdir(), key=lambda item: item.name.lower()):
        if path.is_file() and path.suffix.lower() in ALLOWED_DOCUMENT_SUFFIXES:
            # 这个接口是页面初始化路径，必须只读取目录项和 stat 元数据。过去这里会为每份
            # PDF/DOCX 调用完整解析，扫描件还可能进入 OCR，导致客户只是打开下拉列表就等待
            # 数十秒。正文预览继续由单文件 preview API 和实际 Agent Tool 按需读取。
            documents.append(_document_info(path))
    return documents


def import_workspace_document(*, filename: str, content: str) -> WorkspaceDocumentInfo:
    """把一份 UTF-8 文本文档写入受控 workspace，兼容既有导入协议。"""

    safe_name = _safe_document_filename(filename)
    if Path(safe_name).suffix.lower() not in TEXT_DOCUMENT_SUFFIXES:
        raise WorkspaceDocumentError("PDF、DOCX 和图片请使用二进制导入协议。")
    encoded = content.encode("utf-8")
    _validate_source_size(len(encoded), is_binary=False)
    return _write_workspace_document(safe_name=safe_name, data=encoded)


def import_workspace_document_base64(*, filename: str, content_base64: str) -> WorkspaceDocumentInfo:
    """导入 PDF/DOCX/图片的受限 Base64 数据，不接受客户端文件系统路径。"""

    safe_name = _safe_document_filename(filename)
    if Path(safe_name).suffix.lower() not in BINARY_DOCUMENT_SUFFIXES:
        raise WorkspaceDocumentError("Base64 导入目前只支持 PDF、DOCX、PNG、JPG 和 JPEG 文件。")
    try:
        data = base64.b64decode(content_base64.encode("ascii"), validate=True)
    except (UnicodeEncodeError, ValueError) as exc:
        raise WorkspaceDocumentError("二进制文档内容不是有效 Base64 数据。") from exc
    _validate_source_size(len(data), is_binary=True)
    return _write_workspace_document(safe_name=safe_name, data=data)


def get_workspace_document_preview(
    *,
    relative_path: str,
    preview_chars: int = 2_400,
) -> WorkspaceDocumentPreviewResponse:
    """返回解析后的受控预览；PDF/DOCX/图片不会把原始二进制回传给客户端。"""

    path = _workspace_document_path(relative_path)
    parsed = _read_workspace_document(path)
    stat = path.stat()
    bounded_preview_chars = max(0, min(preview_chars, 8_000))
    return WorkspaceDocumentPreviewResponse(
        name=path.name,
        relative_path=path.name,
        size_bytes=stat.st_size,
        modified_at=datetime.fromtimestamp(stat.st_mtime, tz=UTC).isoformat(),
        document_type=parsed.document_type,
        preview_chars=bounded_preview_chars,
        truncated=len(parsed.text) > bounded_preview_chars,
        preview=parsed.text[:bounded_preview_chars],
    )


def resolve_workspace_document_path(
    relative_path: str,
    *,
    allowed_suffixes: set[str] | None = None,
) -> Path:
    """返回一份已导入文档的受控绝对路径，供本地确定性 Tool 使用。

    这个函数只给 Runtime 内部使用，不经过 API 返回给客户端。所有调用仍从文件名开始走
    既有的 workspace 清洗规则，再检查文件存在和工具允许的后缀，避免 PDF Tool 因为需要
    原始二进制而绕开受控目录边界。
    """

    path = _workspace_document_path(relative_path)
    if not path.is_file():
        raise WorkspaceDocumentError("未找到指定 workspace 文档。")
    if allowed_suffixes is not None and path.suffix.lower() not in allowed_suffixes:
        supported = "、".join(sorted(item.upper().lstrip(".") for item in allowed_suffixes))
        raise WorkspaceDocumentError(f"当前操作只支持 {supported} 文件。")
    return path


def parse_controlled_document(
    path: Path,
    *,
    on_ocr_started: Callable[[], None] | None = None,
) -> ParsedControlledDocument:
    """解析一份已由调用方验证边界的本地受控文件。

    workspace 预览、文档助手和知识库必须共享这个入口，才能保证 PDF 页码、DOCX 段落/表格
    与 UTF-8 文本行号的来源语义完全一致。调用方仍负责确认 ``path`` 位于自己的受控根目录；
    此函数不返回路径，也不接受来自 API 的任意路径字符串。
    """

    return _read_workspace_document(path, on_ocr_started=on_ocr_started)


def source_location_for_range(
    parsed: ParsedControlledDocument,
    start_char: int,
    end_char: int,
) -> tuple[SourceKind, str, int, int]:
    """返回解析文本范围的稳定来源定位，供知识库分块与引用共同使用。"""

    return _source_location_for_range(parsed, start_char, end_char)


def search_workspace_documents(
    *,
    query: str,
    limit: int = 20,
    case_sensitive: bool = False,
    context_chars: int = 80,
    allowed_relative_paths: Iterable[str] | None = None,
) -> WorkspaceDocumentSearchResponse:
    """在已导入材料中做确定性文本定位，并返回行/页/段落级来源。"""

    search_query = query.strip()
    if not search_query:
        raise WorkspaceDocumentError("搜索词不能为空。")

    bounded_limit = max(1, min(limit, 50))
    bounded_context = max(0, min(context_chars, 240))
    allowed_names = _allowed_workspace_names(allowed_relative_paths)
    root = workspace_documents_dir()
    matches: list[WorkspaceDocumentSearchMatch] = []
    searched_documents = 0

    if not root.exists():
        return WorkspaceDocumentSearchResponse(
            query=search_query,
            total=0,
            searched_documents=0,
            limit=bounded_limit,
            matched_documents=[],
            matches=[],
        )

    needle = search_query if case_sensitive else search_query.casefold()
    for path in sorted(root.iterdir(), key=lambda item: item.name.lower()):
        if not path.is_file() or path.suffix.lower() not in ALLOWED_DOCUMENT_SUFFIXES:
            continue
        if allowed_names is not None and path.name not in allowed_names:
            continue
        if path.stat().st_size > _max_document_bytes(path):
            continue

        try:
            parsed = _read_workspace_document(path)
        except WorkspaceDocumentError:
            # 其余可用文件仍应可搜索；用户选中损坏文件真正执行时会收到精确解析原因。
            continue
        searched_documents += 1
        char_offset = 0
        for line_number, line_with_end in enumerate(parsed.text.splitlines(keepends=True), start=1):
            line = line_with_end.rstrip("\r\n")
            haystack = line if case_sensitive else line.casefold()
            match_index = haystack.find(needle)
            if match_index >= 0:
                source_kind, source_locator, source_start, _source_end = _source_location_for_range(
                    parsed,
                    char_offset + match_index,
                    char_offset + match_index + len(search_query),
                )
                matches.append(
                    WorkspaceDocumentSearchMatch(
                        document_name=path.name,
                        relative_path=path.name,
                        # 对新格式这个字段兼容地表示页码/段落序号；source_locator 是最终显示依据。
                        line_number=source_start if parsed.document_type != "text" else line_number,
                        line_text=line,
                        preview=_line_preview(
                            line=line,
                            match_index=match_index,
                            query_length=len(search_query),
                            context_chars=bounded_context,
                        ),
                        source_kind=source_kind,
                        source_locator=source_locator,
                    )
                )
                if len(matches) >= bounded_limit:
                    return WorkspaceDocumentSearchResponse(
                        query=search_query,
                        total=len(matches),
                        searched_documents=searched_documents,
                        limit=bounded_limit,
                        limit_reached=True,
                        matched_documents=_matched_document_names(matches),
                        suggested_read_path=None,
                        matches=matches,
                    )
            char_offset += len(line_with_end)

    matched_documents = _matched_document_names(matches)
    return WorkspaceDocumentSearchResponse(
        query=search_query,
        total=len(matches),
        searched_documents=searched_documents,
        limit=bounded_limit,
        matched_documents=matched_documents,
        suggested_read_path=matched_documents[0] if len(matched_documents) == 1 else None,
        matches=matches,
    )


def read_workspace_document_excerpt(
    *,
    relative_path: str,
    start_char: int = 0,
    max_chars: int = 48_000,
) -> dict[str, object]:
    """读取一个受限文本窗口，并携带可展示的来源类型和定位。"""

    path = _workspace_document_path(relative_path)
    return _workspace_text_excerpt(
        path=path,
        parsed=_read_workspace_document(path),
        start_char=start_char,
        max_chars=max_chars,
    )


def read_workspace_document_chunks(
    *,
    relative_path: str,
    chunk_chars: int = 32_000,
) -> list[dict[str, object]]:
    """按字符和自然换行分块一次已解析的文档，避免长文档反复读盘或重复解析。"""

    path = _workspace_document_path(relative_path)
    parsed = _read_workspace_document(path)
    text = parsed.text
    bounded_chunk_chars = max(4_000, min(chunk_chars, 48_000))
    if not text:
        return [_workspace_text_excerpt(path=path, parsed=parsed, start_char=0, max_chars=bounded_chunk_chars)]

    chunks: list[dict[str, object]] = []
    start_char = 0
    while start_char < len(text):
        preferred_end = min(len(text), start_char + bounded_chunk_chars)
        end_char = preferred_end
        if preferred_end < len(text):
            newline_index = text.rfind("\n", start_char, preferred_end)
            if newline_index >= start_char + bounded_chunk_chars // 3:
                end_char = newline_index + 1
        chunk = _workspace_text_excerpt(
            path=path,
            parsed=parsed,
            start_char=start_char,
            max_chars=max(1_000, end_char - start_char),
        )
        chunks.append(chunk)
        next_start_char = int(chunk["end_char"])
        if next_start_char <= start_char:
            raise WorkspaceDocumentError("文档分块没有产生可继续的读取位置。")
        start_char = next_start_char
    return chunks


def read_workspace_document_preview(
    *,
    relative_path: str,
    preview_chars: int = 2_400,
) -> dict[str, object]:
    """供 Runtime 的受控预读使用，仍只暴露相对文件名和限长解析文本。"""

    path = _workspace_document_path(relative_path)
    parsed = _read_workspace_document(path)
    bounded_preview_chars = max(0, min(preview_chars, 8_000))
    return {
        "path": str(path.resolve()),
        "relative_path": path.name,
        "bytes": path.stat().st_size,
        "document_type": parsed.document_type,
        "preview": parsed.text[:bounded_preview_chars],
    }


def _write_workspace_document(*, safe_name: str, data: bytes) -> WorkspaceDocumentInfo:
    root = workspace_documents_dir()
    root.mkdir(parents=True, exist_ok=True)
    target = _available_path(root / safe_name)
    target.write_bytes(data)
    _invalidate_parse_cache(target)
    # 导入成功后的聊天提示只保留 UTF-8 文本的极短预览。二进制材料在导入阶段不解析，
    # 避免 PDF/DOCX/图片写入完成却因为首次解析或 OCR 使导入请求看似卡住。
    return _document_info(target, include_text_preview=True)


def _safe_document_filename(filename: str) -> str:
    raw_name = Path(filename.strip()).name
    if not raw_name:
        raise WorkspaceDocumentError("文件名为空。")
    suffix = Path(raw_name).suffix.lower()
    if suffix not in ALLOWED_DOCUMENT_SUFFIXES:
        raise WorkspaceDocumentError("当前支持 TXT、Markdown、PDF、DOCX、PNG、JPG 和 JPEG 文件。")
    stem = Path(raw_name).stem.strip()
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1F]+', "_", stem)
    stem = re.sub(r"\s+", "_", stem).strip("._ ")
    return f"{stem or 'document'}{suffix}"


def _validate_source_size(size_bytes: int, *, is_binary: bool) -> None:
    limit = MAX_WORKSPACE_BINARY_DOCUMENT_BYTES if is_binary else MAX_WORKSPACE_DOCUMENT_BYTES
    if size_bytes > limit:
        type_text = "PDF/DOCX/图片" if is_binary else "UTF-8 文本"
        raise WorkspaceDocumentError(f"文档超过 {limit // 1_000_000}MB，当前只读 {type_text} 导入不接受该文件。")


def _available_path(path: Path) -> Path:
    if not path.exists():
        return path
    for index in range(1, 1000):
        candidate = path.with_name(f"{path.stem}_{index}{path.suffix}")
        if not candidate.exists():
            return candidate
    raise WorkspaceDocumentError("同名文件过多，请清理 workspace 后再导入。")


def _document_info(path: Path, *, include_text_preview: bool = False) -> WorkspaceDocumentInfo:
    """返回轻量目录元数据；只有刚导入的文本才读取有限预览。"""

    stat = path.stat()
    document_type = _document_type(path)
    preview = ""
    if include_text_preview and path.suffix.lower() in TEXT_DOCUMENT_SUFFIXES:
        try:
            # 这不是文档解析器：仅为导入后的即时反馈读取前 400 个 UTF-8 字符。PDF/DOCX
            # 的真实页码、表格和 OCR 来源仍只由 _read_workspace_document 统一产出。
            with path.open("r", encoding="utf-8") as source:
                preview = source.read(400)
        except (OSError, UnicodeDecodeError):
            # 导入成功不能因为即时提示读失败而被回滚；后续真正预览会提供精确错误原因。
            preview = ""
    return WorkspaceDocumentInfo(
        name=path.name,
        relative_path=path.name,
        size_bytes=stat.st_size,
        modified_at=datetime.fromtimestamp(stat.st_mtime, tz=UTC).isoformat(),
        document_type=document_type,
        preview=preview,
    )


def _read_workspace_document(
    path: Path,
    *,
    on_ocr_started: Callable[[], None] | None = None,
) -> _ParsedWorkspaceDocument:
    """读取并解析已校验文件；缓存只存当前文件版本的不可变解析结果。"""

    if not path.exists() or not path.is_file():
        raise WorkspaceDocumentError("未找到指定 workspace 文档。")
    stat = path.stat()
    _validate_source_size(stat.st_size, is_binary=path.suffix.lower() in BINARY_DOCUMENT_SUFFIXES)
    cache_key = (str(path.resolve()), stat.st_mtime_ns, stat.st_size)
    with _parse_cache_lock:
        cached = _parse_cache.get(cache_key)
        if cached is not None:
            _parse_cache.move_to_end(cache_key)
            return cached

    parsed = _parse_workspace_document(path, on_ocr_started=on_ocr_started)
    with _parse_cache_lock:
        _parse_cache[cache_key] = parsed
        _parse_cache.move_to_end(cache_key)
        while len(_parse_cache) > _PARSE_CACHE_MAX_ENTRIES:
            _parse_cache.popitem(last=False)
    return parsed


def _parse_workspace_document(
    path: Path,
    *,
    on_ocr_started: Callable[[], None] | None,
) -> _ParsedWorkspaceDocument:
    suffix = path.suffix.lower()
    if suffix in TEXT_DOCUMENT_SUFFIXES:
        try:
            return _ParsedWorkspaceDocument(text=path.read_text(encoding="utf-8"), document_type="text")
        except UnicodeDecodeError as exc:
            raise WorkspaceDocumentError("文档不是有效 UTF-8 文本。") from exc
    if suffix == ".pdf":
        return _parse_pdf_document(path, on_ocr_started=on_ocr_started)
    if suffix == ".docx":
        return _parse_docx_document(path)
    if suffix in IMAGE_DOCUMENT_SUFFIXES:
        return _parse_image_document(path, on_ocr_started=on_ocr_started)
    raise WorkspaceDocumentError("当前不支持该文档格式。")


def _parse_pdf_document(
    path: Path,
    *,
    on_ocr_started: Callable[[], None] | None,
) -> _ParsedWorkspaceDocument:
    if fitz is None:
        raise WorkspaceDocumentError("PDF 解析依赖未安装，请重新安装 backend requirements。")
    try:
        pdf = fitz.open(path)
    except Exception as exc:
        raise WorkspaceDocumentError("PDF 无法打开，可能已损坏、加密或不是有效 PDF。") from exc
    try:
        parts: list[str] = []
        segments: list[_DocumentSegment] = []
        cursor = 0
        for page_index, page in enumerate(pdf, start=1):
            page_text = page.get_text("text").strip()
            if not page_text:
                continue
            if parts:
                parts.append("\n\n")
                cursor += 2
            start = cursor
            parts.append(page_text)
            cursor += len(page_text)
            segments.append(
                _DocumentSegment(start, cursor, "page", f"第 {page_index} 页", page_index)
            )
            _ensure_extracted_text_limit(cursor)
        text = "".join(parts)
    except WorkspaceDocumentError:
        raise
    except Exception as exc:
        raise WorkspaceDocumentError("PDF 文本提取失败。") from exc
    finally:
        pdf.close()
    # 可复制文本 PDF 保持原有快速解析路径；只有整份 PDF 完全没有文本层时才进入 OCR，避免
    # 对混合 PDF 进行无谓识别、引入两套互相矛盾的来源文字或触发可选组件的初始化。
    if not text.strip():
        return _parse_ocr_document(
            path,
            document_type="pdf",
            on_ocr_started=on_ocr_started,
        )
    return _ParsedWorkspaceDocument(text=text, document_type="pdf", segments=tuple(segments))


def _parse_docx_document(path: Path) -> _ParsedWorkspaceDocument:
    if DocxDocument is None:
        raise WorkspaceDocumentError("DOCX 解析依赖未安装，请重新安装 backend requirements。")
    try:
        document = DocxDocument(BytesIO(path.read_bytes()))
    except Exception as exc:
        raise WorkspaceDocumentError("DOCX 无法打开，可能已损坏或不是有效 Word 文档。") from exc

    parts: list[str] = []
    segments: list[_DocumentSegment] = []
    cursor = 0

    def append_segment(text: str, kind: SourceKind, locator: str, ordinal: int) -> None:
        nonlocal cursor
        normalized = text.strip()
        if not normalized:
            return
        if parts:
            parts.append("\n\n")
            cursor += 2
        start = cursor
        parts.append(normalized)
        cursor += len(normalized)
        _ensure_extracted_text_limit(cursor)
        segments.append(_DocumentSegment(start, cursor, kind, locator, ordinal))

    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        append_segment(paragraph.text, "paragraph", f"第 {paragraph_index} 段", paragraph_index)

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        append_segment("\n".join(rows), "table", f"表格 {table_index}", table_index)

    text = "".join(parts)
    if not text.strip():
        raise WorkspaceDocumentError("DOCX 没有可提取的段落或表格文本。")
    return _ParsedWorkspaceDocument(text=text, document_type="docx", segments=tuple(segments))


def _parse_image_document(
    path: Path,
    *,
    on_ocr_started: Callable[[], None] | None,
) -> _ParsedWorkspaceDocument:
    """把受控图片交给已准备的本地 OCR；模型准备绝不由这里触发。"""

    return _parse_ocr_document(
        path,
        document_type="image",
        on_ocr_started=on_ocr_started,
    )


def _create_ocr_adapter() -> OcrAdapter:
    """保留窄小工厂接缝，供离线解析回归替换假引擎而不安装或加载 Paddle。"""

    return OcrAdapter()


def _parse_ocr_document(
    path: Path,
    *,
    document_type: Literal["pdf", "image"],
    on_ocr_started: Callable[[], None] | None,
) -> _ParsedWorkspaceDocument:
    """把 OCR 的页/区域协议转为既有解析文本与来源分块契约。

    这里不保存原图、坐标或失败页细节，也不吞掉未准备状态。上层只得到可行动的短错误，成功
    区域则成为精确来源锚点；PDF 单页失败不会撤回其它已识别页面。
    """

    if on_ocr_started is not None:
        # 仅在已经确认材料实际需要 OCR 时通知上层；可复制文本 PDF 和缓存命中都不会走这里，
        # 因而 Qt 不会显示虚假的“正在识别”。
        on_ocr_started()

    adapter = _create_ocr_adapter()
    try:
        result = adapter.recognize_path(path)
    except OcrAdapterError as exc:
        raise WorkspaceDocumentError(_ocr_error_message(exc)) from exc
    if result.document_type != document_type:
        # 这是内部 Adapter 与解析层的协议保护，不向客户暴露引擎实现或实际文件路径。
        raise WorkspaceDocumentError("OCR 返回的材料类型与受控解析请求不一致。")
    retried_page_count = 0
    retryable_page_numbers = tuple(
        page.page_number
        for page in result.pages
        if page.status == "failed" and page.failure_code == "ocr_page_failed"
    )
    if document_type == "pdf" and retryable_page_numbers:
        # 只对引擎异常页再尝试一次；空白页重试没有信息增益，不能把 OCR 变成本地死循环。
        retried_page_count = len(retryable_page_numbers)
        try:
            retry_result = adapter.recognize_path(path, page_numbers=retryable_page_numbers)
        except OcrAdapterError:
            # 首轮已有其它可用页面时，重试自身失败不能撤回已识别文本。最终统计仍会如实保留
            # 未识别页数；错误细节不写入正文、事件或 SQLite。
            retry_result = None
        if retry_result is not None:
            if retry_result.document_type != document_type:
                raise WorkspaceDocumentError("OCR 页级重试返回的材料类型不一致。")
            result = _merge_ocr_retry_result(
                original=result,
                retry_result=retry_result,
                retried_page_numbers=retryable_page_numbers,
            )
    return _parsed_document_from_ocr_result(
        result,
        document_type=document_type,
        retried_page_count=retried_page_count,
    )


def _parsed_document_from_ocr_result(
    result: OcrDocumentResult,
    *,
    document_type: Literal["pdf", "image"],
    retried_page_count: int = 0,
) -> _ParsedWorkspaceDocument:
    """按识别区域拼接文本，并让每段文字保留稳定的页码/区域定位。"""

    parts: list[str] = []
    segments: list[_DocumentSegment] = []
    cursor = 0
    segment_ordinal = 0

    def append_region(text: str, locator: str) -> None:
        nonlocal cursor, segment_ordinal
        normalized = text.strip()
        if not normalized:
            return
        if parts:
            parts.append("\n")
            cursor += 1
        start = cursor
        parts.append(normalized)
        cursor += len(normalized)
        _ensure_extracted_text_limit(cursor)
        segment_ordinal += 1
        segments.append(_DocumentSegment(start, cursor, "region", locator, segment_ordinal))

    for page in result.pages:
        if page.status != "completed":
            continue
        if page.regions:
            for region in page.regions:
                append_region(region.text, f"第 {page.page_number} 页 · 区域 {region.ordinal}")
        else:
            # Adapter 正常会带区域；保留这个确定性兜底，避免未来兼容引擎只能返回整页文字时
            # 成功页面被误丢弃，同时来源仍至少可回到正确页码。
            append_region(page.text, f"第 {page.page_number} 页 · OCR 文本")

    text = "".join(parts)
    if not text.strip() or not segments:
        raise WorkspaceDocumentError("OCR 未识别到可用于分析的文字。")
    return _ParsedWorkspaceDocument(
        text=text,
        document_type=document_type,
        segments=tuple(segments),
        ocr_page_count=len(result.pages),
        ocr_completed_page_count=result.successful_page_count,
        ocr_failed_page_count=result.failed_page_count,
        ocr_retried_page_count=max(0, retried_page_count),
    )


def _merge_ocr_retry_result(
    *,
    original: OcrDocumentResult,
    retry_result: OcrDocumentResult,
    retried_page_numbers: tuple[int, ...],
) -> OcrDocumentResult:
    """仅用指定页的重试结果替换原失败页，防止恢复动作影响已经成功的页。"""

    retry_pages = {page.page_number: page for page in retry_result.pages}
    merged_pages = tuple(
        retry_pages.get(page.page_number, page)
        if page.page_number in retried_page_numbers
        else page
        for page in original.pages
    )
    return OcrDocumentResult(document_type=original.document_type, pages=merged_pages)


def _ocr_error_message(error: OcrAdapterError) -> str:
    """把 Adapter 的稳定错误码转换为客户可理解、无环境细节的解析说明。"""

    if error.code == "ocr_not_installed":
        return "该材料需要 OCR；本地 OCR 可选组件尚未安装。"
    if error.code == "ocr_not_ready":
        return "该材料需要 OCR；本地模型尚未准备，请先在界面确认下载。"
    if error.code == "ocr_unsupported_document":
        return "OCR 当前只支持 PNG、JPG、JPEG 和无文本层 PDF。"
    if error.code == "ocr_invalid_document":
        return "OCR 材料无法读取，可能已损坏、加密或格式无效。"
    return "OCR 未能从该材料识别可用于分析的文字。"


def _ensure_extracted_text_limit(char_count: int) -> None:
    if char_count > MAX_EXTRACTED_DOCUMENT_CHARS:
        raise WorkspaceDocumentError("解析后的文本超过 100 万字符，本次只读分析为保护性能未导入该文件。")


def _invalidate_parse_cache(path: Path) -> None:
    resolved = str(path.resolve())
    with _parse_cache_lock:
        stale_keys = [key for key in _parse_cache if key[0] == resolved]
        for key in stale_keys:
            _parse_cache.pop(key, None)


def _workspace_text_excerpt(
    *,
    path: Path,
    parsed: _ParsedWorkspaceDocument,
    start_char: int,
    max_chars: int,
) -> dict[str, object]:
    bounded_start_char = max(0, min(start_char, len(parsed.text)))
    bounded_max_chars = max(1_000, min(max_chars, 48_000))
    end_char = min(len(parsed.text), bounded_start_char + bounded_max_chars)
    excerpt = parsed.text[bounded_start_char:end_char]
    source_kind, source_locator, source_start, source_end = _source_location_for_range(
        parsed,
        bounded_start_char,
        max(bounded_start_char + 1, end_char),
    )
    total_lines = max(1, len(parsed.text.splitlines()))
    if parsed.document_type == "text":
        start_line = parsed.text.count("\n", 0, bounded_start_char) + 1
        if not excerpt:
            end_line = start_line
        else:
            end_line = start_line + excerpt.count("\n")
            if excerpt.endswith("\n"):
                end_line -= 1
            end_line = min(total_lines, max(start_line, end_line))
        source_start, source_end = start_line, end_line
    return {
        "relative_path": path.name,
        "bytes": path.stat().st_size,
        "document_type": parsed.document_type,
        "total_lines": total_lines,
        "start_line": source_start,
        "end_line": source_end,
        "source_kind": source_kind,
        "source_locator": source_locator,
        "start_char": bounded_start_char,
        "end_char": end_char,
        "next_start_char": end_char if end_char < len(parsed.text) else None,
        "text": excerpt,
        "truncated": end_char < len(parsed.text),
    }


def _source_location_for_range(
    parsed: _ParsedWorkspaceDocument,
    start_char: int,
    end_char: int,
) -> tuple[SourceKind, str, int, int]:
    if parsed.document_type == "text":
        start_line = parsed.text.count("\n", 0, max(0, start_char)) + 1
        end_line = parsed.text.count("\n", 0, max(0, end_char - 1)) + 1
        return "line", _line_locator(start_line, end_line), start_line, max(start_line, end_line)

    overlapping = [
        segment
        for segment in parsed.segments
        if segment.end_char > start_char and segment.start_char < end_char
    ]
    if not overlapping and parsed.segments:
        # 片段恰好落在分隔换行时，选择相邻的最近证据，而不是给模型一条无定位的来源。
        overlapping = [min(parsed.segments, key=lambda item: abs(item.start_char - start_char))]
    if not overlapping:
        return "mixed", "未定位到可提取内容", 1, 1

    kinds = {segment.source_kind for segment in overlapping}
    kind: SourceKind = overlapping[0].source_kind if len(kinds) == 1 else "mixed"
    start_ordinal = overlapping[0].ordinal
    end_ordinal = overlapping[-1].ordinal
    if kind == "page":
        locator = _page_locator(start_ordinal, end_ordinal)
    elif kind == "paragraph":
        locator = _paragraph_locator(start_ordinal, end_ordinal)
    elif kind == "table":
        locator = "、".join(segment.source_locator for segment in overlapping[:4])
    else:
        locator = "、".join(segment.source_locator for segment in overlapping[:4])
        if len(overlapping) > 4:
            locator += " 等"
    return kind, locator, start_ordinal, end_ordinal


def _line_locator(start: int, end: int) -> str:
    return f"第 {start} 行" if start == end else f"第 {start}-{end} 行"


def _page_locator(start: int, end: int) -> str:
    return f"第 {start} 页" if start == end else f"第 {start}-{end} 页"


def _paragraph_locator(start: int, end: int) -> str:
    return f"第 {start} 段" if start == end else f"第 {start}-{end} 段"


def _document_type(path: Path) -> Literal["text", "pdf", "docx", "image"]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix in IMAGE_DOCUMENT_SUFFIXES:
        return "image"
    return "text"


def _max_document_bytes(path: Path) -> int:
    return MAX_WORKSPACE_BINARY_DOCUMENT_BYTES if path.suffix.lower() in BINARY_DOCUMENT_SUFFIXES else MAX_WORKSPACE_DOCUMENT_BYTES


def _matched_document_names(matches: list[WorkspaceDocumentSearchMatch]) -> list[str]:
    names: list[str] = []
    for match in matches:
        if match.relative_path not in names:
            names.append(match.relative_path)
    return names


def _workspace_document_path(relative_path: str) -> Path:
    raw_name = Path(relative_path.strip()).name
    if not raw_name:
        raise WorkspaceDocumentError("文档路径为空。")
    if Path(raw_name).suffix.lower() not in ALLOWED_DOCUMENT_SUFFIXES:
        raise WorkspaceDocumentError("当前支持 TXT、Markdown、PDF、DOCX、PNG、JPG 和 JPEG 文件。")
    return workspace_documents_dir() / raw_name


def _allowed_workspace_names(relative_paths: Iterable[str] | None) -> set[str] | None:
    if relative_paths is None:
        return None
    names: set[str] = set()
    for relative_path in relative_paths:
        if not isinstance(relative_path, str):
            raise WorkspaceDocumentError("受限搜索范围包含无效文档名。")
        names.add(_workspace_document_path(relative_path).name)
    return names


def _line_preview(*, line: str, match_index: int, query_length: int, context_chars: int) -> str:
    start = max(0, match_index - context_chars)
    end = min(len(line), match_index + query_length + context_chars)
    return f"{'...' if start else ''}{line[start:end]}{'...' if end < len(line) else ''}"
