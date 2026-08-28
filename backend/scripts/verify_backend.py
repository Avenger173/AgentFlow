import atexit
import asyncio
import base64
from io import BytesIO
import json
import os
from pathlib import Path
import shutil
import sqlite3
import sys
import tempfile
import warnings

# 离线验证必须稳定、无成本；即使 backend/.env 配了真实模型，也强制使用 mock。
os.environ["AGENTFLOW_CHAT_MODE"] = "mock"

warnings.filterwarnings(
    "ignore",
    message="Using `httpx` with `starlette.testclient` is deprecated.*",
)
from fastapi.testclient import TestClient
import fitz
from docx import Document as DocxDocument

BACKEND_ROOT = Path(__file__).resolve().parents[1]
_VERIFY_DATA_DIR = Path(tempfile.mkdtemp(prefix="agentflow_verify_"))
os.environ["AGENTFLOW_DATA_DIR"] = str(_VERIFY_DATA_DIR)
atexit.register(lambda: shutil.rmtree(_VERIFY_DATA_DIR, ignore_errors=True))
sys.path.insert(0, str(BACKEND_ROOT))

from main import app
from app.agents.runner import AgentDefinition, AgentRunner, _parse_model_output
from app.core.config import settings
from app.schemas.chat import (
    WorkflowCommandPolicy,
    WorkflowPlan,
    WorkflowPlanPreferences,
    WorkflowStep,
)
from app.schemas.document_agent import DocumentModelOutput
from app.schemas.presentation_studio import (
    PresentationStudioExportRequest,
    PresentationStudioPlanRequest,
)
from app.services.llm_chat import build_personality_instruction
from app.services.agent_catalog import list_agents
from app.services.commander import create_commander_plan
from app.services import document_agent as document_agent_service
from app.services.model_gateway import (
    ModelConversationMessage,
    ModelGatewayConnectionError,
    ModelToolCall,
    ModelToolTurn,
    _anthropic_tool_messages,
    _extract_anthropic_tool_turn,
    _extract_openai_tool_turn,
    _openai_tool_messages,
)
from app.services.workspace_documents import read_workspace_document_excerpt
from app.workflow.dry_run import clear_dry_run_memory_cache, run_workflow_dry_run
from app.workflow.permission_policy import evaluate_permission_policy
from app.workflow.state_machine import allowed_next_statuses, can_transition, is_terminal_status
from app.workflow.validator import validate_workflow_plan


class _OutputRepairModel:
    """模拟 JSON mode 首次返回空内容，第二次在受控修复提示下正常收束。"""

    def __init__(self) -> None:
        self.calls: list[tuple[list[object], list[object]]] = []

    async def tool_turn(self, *, system_prompt: str, messages: list[object], tools: list[object]) -> ModelToolTurn:
        del system_prompt
        self.calls.append((list(messages), list(tools)))
        if len(self.calls) == 1:
            # DeepSeek 官方文档明确提示 JSON mode 可能偶发返回空 content。Runner 必须在
            # repair 前舍弃这条无 Tool 的无效 assistant 消息，不能把它回传给下一轮。
            return ModelToolTurn(content="")

        # 格式修复不允许重新暴露 Tool，也不能藉机扩展读取范围。
        assert tools == []
        assert [getattr(message, "role", "") for message in messages] == ["user", "user"]
        assert "仅重写最终结果" in str(messages[-1])
        return ModelToolTurn(
            content=(
                '{"answer":"已完成结构化修复","answer_source_ids":["src_001"],'
                '"summary":"修复后的摘要","requirements":[],"comparisons":[],'
                '"constraints":[],"todos":[],"entities":[],"open_questions":[],"confidence":"medium"}'
            )
        )


class _TimeoutModel:
    """模拟模型在网络层或供应商侧超过等待时限。"""

    async def tool_turn(self, *, system_prompt: str, messages: list[object], tools: list[object]) -> ModelToolTurn:
        del system_prompt, messages, tools
        raise TimeoutError("verification timeout")


class _ConnectionFailureModel:
    """模拟模型服务在请求发出前无法连接。"""

    async def tool_turn(self, *, system_prompt: str, messages: list[object], tools: list[object]) -> ModelToolTurn:
        del system_prompt, messages, tools
        raise ModelGatewayConnectionError("模型服务当前无法连接。")


class _InvalidRequirementsModel:
    """模拟读取成功后仍连续两次返回无效最终 JSON 的供应商。"""

    def __init__(self, *, request: object, selected_documents: list[str]) -> None:
        del request
        self.selected_documents = selected_documents

    async def tool_turn(self, *, system_prompt: str, messages: list[object], tools: list[object]) -> ModelToolTurn:
        del system_prompt, tools
        has_tool_result = any(getattr(message, "role", "") == "tool" for message in messages)
        if not has_tool_result:
            return ModelToolTurn(
                tool_calls=(
                    ModelToolCall(
                        call_id="invalid_requirements_read_001",
                        name="document_read_text",
                        arguments={"relative_path": self.selected_documents[0], "max_chars": 48_000},
                    ),
                )
            )
        # Runner 会先请求一次无工具 JSON 修复；此模型故意继续失败，以验证文档助手的
        # 需求专属保守降级不会丢掉已读取的来源。
        return ModelToolTurn(content="这不是合法 JSON。")


class _InvalidBriefModel:
    """模拟关键信息卡在读取成功后仍两次返回无效 JSON 的供应商。"""

    def __init__(self, *, request: object, selected_documents: list[str]) -> None:
        del request
        self.selected_documents = selected_documents

    async def tool_turn(self, *, system_prompt: str, messages: list[object], tools: list[object]) -> ModelToolTurn:
        del system_prompt, tools
        has_tool_result = any(getattr(message, "role", "") == "tool" for message in messages)
        if not has_tool_result:
            return ModelToolTurn(
                tool_calls=(
                    ModelToolCall(
                        call_id="invalid_brief_read_001",
                        name="document_read_text",
                        arguments={"relative_path": self.selected_documents[0], "max_chars": 48_000},
                    ),
                )
            )
        # 让 Runner 走完一次无工具修复，再由文档助手的只读保守降级接管。
        return ModelToolTurn(content="这不是合法的关键信息卡 JSON。")


class _InvalidOutlineModel:
    """模拟结构化大纲在读取成功后仍两次返回无效 JSON 的供应商。"""

    def __init__(self, *, request: object, selected_documents: list[str]) -> None:
        del request
        self.selected_documents = selected_documents

    async def tool_turn(self, *, system_prompt: str, messages: list[object], tools: list[object]) -> ModelToolTurn:
        del system_prompt, tools
        has_tool_result = any(getattr(message, "role", "") == "tool" for message in messages)
        if not has_tool_result:
            return ModelToolTurn(
                tool_calls=(
                    ModelToolCall(
                        call_id="invalid_outline_read_001",
                        name="document_read_text",
                        arguments={"relative_path": self.selected_documents[0], "max_chars": 48_000},
                    ),
                )
            )
        return ModelToolTurn(content="这不是合法的结构化大纲 JSON。")


def main() -> None:
    client = TestClient(app)
    # 文档解析已按规范转入线程池，异步 /start 任务会真正让出事件循环。TestClient 若不进入
    # context manager，会为每个请求短暂创建并销毁 portal，后台任务会被测试宿主取消；这里显式
    # 保持整轮离线验收共用一个 lifespan，才与真实 Uvicorn 的持续事件循环一致。
    client.__enter__()
    atexit.register(lambda: client.__exit__(None, None, None))

    # PPT V3 的公开资料是“计划内、确认后”的受控联网能力。主回归至少要固定其协议形状，
    # 防止未来调整通用模型配置或文档 API 时让客户端字段被静默忽略、误解为自动联网。
    studio_request = PresentationStudioPlanRequest(
        intent="为客户说明公开资料来源的边界。",
        visual_asset_provider="none",
        public_research_enabled=True,
        structured_data_enabled=True,
    )
    assert studio_request.public_research_enabled is True
    assert studio_request.structured_data_enabled is True
    studio_export = PresentationStudioExportRequest(
        plan_id="a" * 48,
        filename="公开资料协议验证.pptx",
        confirmed=True,
        fetch_public_research=True,
        fetch_structured_data=True,
        network_confirmed=True,
    )
    assert studio_export.fetch_public_research is True
    assert studio_export.fetch_structured_data is True
    assert studio_export.network_confirmed is True

    # AgentRunner 只处理统一的 ModelToolTurn，但真实供应商需要把同一轮“模型请求工具 ->
    # Runtime 返回结果”转换成不同的消息形状。这里固定验证 OpenAI-compatible 和 Anthropic
    # 两条协议，避免日后改 Kimi 思考回传、函数别名或 Tool Result 拼装时，只在真实额度测试里
    # 才发现文档助手无法继续下一轮。
    protocol_call = ModelToolCall(
        call_id="verify_read_001",
        name="document_read_text",
        arguments={"relative_path": "assignment.md", "max_chars": 48_000},
    )
    protocol_messages = [
        ModelConversationMessage(role="user", content="读取已选择的材料。"),
        ModelConversationMessage(
            role="assistant",
            tool_calls=(protocol_call,),
            reasoning_content="需要先读取受控材料。",
        ),
        ModelConversationMessage(
            role="tool",
            tool_call_id=protocol_call.call_id,
            tool_name=protocol_call.name,
            content='{"ok":true,"result":{"source":{"source_id":"src_001"}}}',
        ),
    ]
    openai_messages = _openai_tool_messages(
        system_prompt="验证 OpenAI-compatible Tool Calls。",
        messages=protocol_messages,
        preserve_reasoning=True,
    )
    assert openai_messages[1]["role"] == "user"
    assert openai_messages[2]["tool_calls"][0]["function"]["name"] == "document_read_text"
    assert openai_messages[2]["reasoning_content"] == "需要先读取受控材料。"
    assert openai_messages[3]["tool_call_id"] == "verify_read_001"
    # DeepSeek V4 思考模式会校验 Tool Call 回放的 content 非空字段；这条断言防止通用
    # OpenAI-compatible 适配再次把它退化成 null，导致第二轮模型调用 HTTP 400。
    deepseek_messages = _openai_tool_messages(
        system_prompt="验证 DeepSeek thinking Tool Calls。",
        messages=protocol_messages,
        preserve_reasoning=True,
        require_tool_call_content=True,
    )
    assert deepseek_messages[2]["content"] == ""
    assert deepseek_messages[2]["reasoning_content"] == "需要先读取受控材料。"
    openai_turn = _extract_openai_tool_turn(
        {
            "choices": [
                {
                    "message": {
                        "content": "",
                        "tool_calls": [
                            {
                                "id": "verify_read_002",
                                "type": "function",
                                "function": {
                                    "name": "document_read_text",
                                    "arguments": '{"relative_path":"assignment.md","max_chars":48000}',
                                },
                            }
                        ],
                    }
                }
            ]
        }
    )
    assert openai_turn.tool_calls == (
        ModelToolCall(
            call_id="verify_read_002",
            name="document_read_text",
            arguments={"relative_path": "assignment.md", "max_chars": 48_000},
        ),
    )

    anthropic_messages = _anthropic_tool_messages(protocol_messages)
    assert anthropic_messages[0]["role"] == "user"
    assert anthropic_messages[1]["content"][0]["type"] == "tool_use"
    assert anthropic_messages[1]["content"][0]["name"] == "document_read_text"
    assert anthropic_messages[2]["content"][0]["type"] == "tool_result"
    assert anthropic_messages[2]["content"][0]["tool_use_id"] == "verify_read_001"
    anthropic_turn = _extract_anthropic_tool_turn(
        {
            "content": [
                {"type": "text", "text": "先读取材料。"},
                {
                    "type": "tool_use",
                    "id": "verify_read_003",
                    "name": "document_read_text",
                    "input": {"relative_path": "assignment.md", "max_chars": 48_000},
                },
            ]
        }
    )
    assert anthropic_turn.content == "先读取材料。"
    assert anthropic_turn.tool_calls == (
        ModelToolCall(
            call_id="verify_read_003",
            name="document_read_text",
            arguments={"relative_path": "assignment.md", "max_chars": 48_000},
        ),
    )

    # 真实模型偶尔会在 JSON 前补一句自然语言，或把待确认项简写成字符串。Output Guardrail
    # 必须只接受其中的 JSON object，并把简写项绑定到模型已提交的来源，不能丢掉可追溯性。
    compact_model_output = _parse_model_output(
        "说明文字不会作为结果。\n```json\n"
        '{"answer":"已读取材料","answer_source_ids":["src_001"],'
        '"summary":"摘要","requirements":[],"constraints":[],"todos":[],"entities":[],'
        '"open_questions":["还需确认范围"],"confidence":"low"}\n```',
        DocumentModelOutput,
    )
    assert isinstance(compact_model_output, DocumentModelOutput)
    assert compact_model_output.open_questions[0].text == "还需确认范围"
    assert compact_model_output.open_questions[0].source_ids == ["src_001"]

    # 不能只尝试模型输出中的第一个 JSON object：模型有时先复述 schema 片段，真实结果在后面。
    trailing_model_output = _parse_model_output(
        '示例：{"answer":"示例，不是完整结果"}\n最终结果：'
        '{"answer":"已读取材料","answer_source_ids":["src_001"],"summary":"摘要",'
        '"requirements":[],"comparisons":[],"constraints":[],"todos":[],"entities":[],'
        '"open_questions":[],"confidence":"medium"}',
        DocumentModelOutput,
    )
    assert isinstance(trailing_model_output, DocumentModelOutput)
    assert trailing_model_output.answer == "已读取材料"

    # 草稿任务的模型有时会把“范围”等具体主题写进 category，或额外给关键信息卡增加
    # acceptance 标签。前者应保守降为 unknown，后者应被忽略；两种情况都不能因为标签
    # 不属于稳定 UI 契约而丢弃已有来源的草稿正文。
    normalized_contract_output = _parse_model_output(
        '{"answer":"草稿已生成","answer_source_ids":["src_001"],'
        '"requirements":[{"id":"req_01","text":"本期只覆盖审批流程",'
        '"category":"scope","priority":"must","source_ids":["src_001"]}],'
        '"brief_fields":[{"key":"acceptance","value":"三分钟完成审批",'
        '"source_ids":["src_001"]}]}',
        DocumentModelOutput,
    )
    assert isinstance(normalized_contract_output, DocumentModelOutput)
    assert normalized_contract_output.requirements[0].category == "unknown"
    assert normalized_contract_output.brief_fields == []

    repair_model = _OutputRepairModel()
    repair_result = asyncio.run(
        AgentRunner().run(
            definition=AgentDefinition(
                agent_id="verification_agent",
                system_prompt="只返回 DocumentModelOutput JSON。",
                tools=(),
                output_model=DocumentModelOutput,
                max_turns=2,
                max_output_repair_attempts=1,
            ),
            model=repair_model,
            user_message="验证结构化输出修复。",
        )
    )
    assert repair_result.status == "completed"
    assert repair_result.output is not None
    assert len(repair_model.calls) == 2
    assert repair_result.turn_traces[0].output_repair_requested is True

    # 模型超时必须有稳定停止原因，便于 Document Agent 映射成可重试的用户提示，而不是
    # 把网络慢响应、供应商错误和输出协议错误都混成同一种失败。
    timeout_result = asyncio.run(
        AgentRunner().run(
            definition=AgentDefinition(
                agent_id="verification_agent",
                system_prompt="验证模型超时。",
                tools=(),
                output_model=DocumentModelOutput,
            ),
            model=_TimeoutModel(),
            user_message="验证超时边界。",
        )
    )
    assert timeout_result.status == "failed"
    assert timeout_result.stop_reason == "model_timeout"
    assert "没有返回" in timeout_result.message

    # 网络不可达需要和超时、HTTP 响应错误分开记录，前端才能提示用户检查网络或 Base URL，
    # 而不是让客户面对 httpx 的 ConnectError 类名。
    connection_result = asyncio.run(
        AgentRunner().run(
            definition=AgentDefinition(
                agent_id="verification_agent",
                system_prompt="验证模型连接边界。",
                tools=(),
                output_model=DocumentModelOutput,
            ),
            model=_ConnectionFailureModel(),
            user_message="验证连接失败边界。",
        )
    )
    assert connection_result.status == "failed"
    assert connection_result.stop_reason == "model_connection_failed"
    assert connection_result.message == "模型服务当前无法连接。"

    health = client.get("/health")
    assert health.status_code == 200, health.text
    health_payload = health.json()
    assert health_payload["status"] == "ok"
    assert health_payload["capabilities"]["data_workspace"]["ready"] is True

    agents = client.get("/api/agents")
    assert agents.status_code == 200, agents.text
    agents_payload = agents.json()
    assert agents_payload["total"] >= 4
    commander = next(
        (agent for agent in agents_payload["agents"] if agent["id"] == "commander_agent"),
        None,
    )
    assert commander is not None
    assert commander["source"] == "builtin"
    assert commander["builtin"] is True
    assert commander["capabilities"]
    assert commander["runtime_ready"] is True
    document_agent = next(
        (agent for agent in agents_payload["agents"] if agent["id"] == "document_agent"),
        None,
    )
    assert document_agent is not None
    assert document_agent["runtime_ready"] is True
    assert document_agent["maturity"] == "mvp"
    assert document_agent["health"] == "ready"
    knowledge_agent = next(
        (agent for agent in agents_payload["agents"] if agent["id"] == "knowledge_agent"),
        None,
    )
    assert knowledge_agent is not None
    assert knowledge_agent["runtime_ready"] is True
    assert knowledge_agent["maturity"] == "mvp"
    assert knowledge_agent["permissions"] == {
        "file_read": False,
        "file_write": False,
        "network": False,
        "shell": False,
        "database": False,
    }
    for placeholder_id in ("code_agent", "report_agent"):
        placeholder = next(agent for agent in agents_payload["agents"] if agent["id"] == placeholder_id)
        assert placeholder["runtime_ready"] is False
        assert placeholder["maturity"] == "placeholder"

    registry = client.get("/api/agents/registry/status")
    assert registry.status_code == 200, registry.text
    registry_payload = registry.json()
    assert registry_payload["loaded_total"] >= 4
    assert registry_payload["errors"] == []

    models = client.get("/api/models/providers")
    assert models.status_code == 200, models.text
    models_payload = models.json()
    provider_ids = {provider["provider"] for provider in models_payload["providers"]}
    assert {"deepseek", "openai", "anthropic", "qwen", "openai_compatible"}.issubset(provider_ids)
    providers_by_id = {provider["provider"]: provider for provider in models_payload["providers"]}
    assert providers_by_id["deepseek"]["context_cache_mode"] == "automatic_observable"
    assert providers_by_id["anthropic"]["context_cache_mode"] == "explicit_request"
    assert providers_by_id["openai_compatible"]["context_cache_mode"] == "unknown"
    assert models_payload["current"]["provider"]

    node_contracts = client.get("/api/workflow/node-contracts")
    assert node_contracts.status_code == 200, node_contracts.text
    node_contracts_payload = node_contracts.json()
    assert node_contracts_payload["total"] >= 4
    contracts_by_tool = {
        contract["tool_name"]: contract
        for contract in node_contracts_payload["contracts"]
    }
    assert "document.read_text" in contracts_by_tool
    assert "document.search_text" in contracts_by_tool
    assert "document.extract_requirements" in contracts_by_tool
    assert "code.generate_code" in contracts_by_tool
    assert "report.compose_markdown" in contracts_by_tool
    assert "agent.knowledge_agent.answer" in contracts_by_tool
    assert "file_read" in contracts_by_tool["document.read_text"]["required_permissions"]
    assert "file_read" in contracts_by_tool["document.search_text"]["required_permissions"]
    assert "file_write" in contracts_by_tool["code.generate_code"]["required_permissions"]
    assert "relative_path" in contracts_by_tool["document.read_text"]["output_schema"]
    assert "tool_timeout" in contracts_by_tool["document.read_text"]["failure_codes"]
    assert "empty_query" in contracts_by_tool["document.search_text"]["failure_codes"]
    assert "context" in contracts_by_tool["document.extract_requirements"]["output_schema"]
    assert "document.context" in contracts_by_tool["document.extract_requirements"]["state_writes"]
    assert "missing_document_context" in contracts_by_tool["document.extract_requirements"]["failure_codes"]
    assert "search_match_total" in contracts_by_tool["document.extract_requirements"]["evaluation_signals"]
    assert "document_context" in contracts_by_tool["code.generate_code"]["input_schema"]
    assert "artifact_verification_failed" in contracts_by_tool["code.generate_code"]["failure_codes"]
    assert "document_context" in contracts_by_tool["report.compose_markdown"]["input_schema"]
    assert "verification" in contracts_by_tool["report.compose_markdown"]["output_schema"]
    assert "artifact_created" in contracts_by_tool["report.compose_markdown"]["evaluation_signals"]
    assert "artifact_verified" in contracts_by_tool["report.compose_markdown"]["evaluation_signals"]
    assert contracts_by_tool["agent.knowledge_agent.answer"]["required_permissions"] == []
    assert "knowledge_base_id" in contracts_by_tool["agent.knowledge_agent.answer"]["input_schema"]

    read_text_contract = client.get(
        "/api/workflow/node-contracts",
        params={"agent_id": "document_agent", "action": "read_text"},
    )
    assert read_text_contract.status_code == 200, read_text_contract.text
    read_text_contract_payload = read_text_contract.json()
    assert read_text_contract_payload["total"] == 1
    assert read_text_contract_payload["contracts"][0]["tool_name"] == "document.read_text"

    extract_contract = client.get(
        "/api/workflow/node-contracts",
        params={"agent_id": "document_agent", "action": "extract_requirements"},
    )
    assert extract_contract.status_code == 200, extract_contract.text
    extract_contract_payload = extract_contract.json()
    assert extract_contract_payload["total"] == 1
    assert extract_contract_payload["contracts"][0]["tool_name"] == "document.extract_requirements"
    assert extract_contract_payload["contracts"][0]["node_type"] == "data"

    document_agent_contract = client.get(
        "/api/workflow/node-contracts",
        params={"agent_id": "document_agent", "action": "analyze_document"},
    )
    assert document_agent_contract.status_code == 200, document_agent_contract.text
    document_agent_contract_payload = document_agent_contract.json()
    assert document_agent_contract_payload["total"] == 1
    assert document_agent_contract_payload["contracts"][0]["tool_name"] == "agent.document_agent.analyze"
    assert document_agent_contract_payload["contracts"][0]["required_permissions"] == ["file_read"]

    command_policy_rules = client.get("/api/workflow/command-policy/rules")
    assert command_policy_rules.status_code == 200, command_policy_rules.text
    command_policy_rules_payload = command_policy_rules.json()
    assert command_policy_rules_payload["total"] >= 10
    command_rule_ids = {
        rule["rule_id"]
        for rule in command_policy_rules_payload["rules"]
    }
    assert "shell.rm_recursive_force" in command_rule_ids
    assert "git.reset_hard" in command_rule_ids
    assert "database.drop_truncate" in command_rule_ids
    assert all(rule["risk_level"] == "high_risk" for rule in command_policy_rules_payload["rules"])
    assert all(rule["default_action"] == "block" for rule in command_policy_rules_payload["rules"])
    assert any(rule["safer_alternatives"] for rule in command_policy_rules_payload["rules"])
    assert "pattern" not in json.dumps(command_policy_rules_payload, ensure_ascii=False)

    read_only_command_policy = client.post(
        "/api/workflow/command-policy/check",
        json={"command": "rg -n \"WorkflowRun\" backend/app"},
    )
    assert read_only_command_policy.status_code == 200, read_only_command_policy.text
    read_only_command_policy_payload = read_only_command_policy.json()
    assert read_only_command_policy_payload["risk_level"] == "read_only"
    assert read_only_command_policy_payload["allowed"] is True
    assert read_only_command_policy_payload["requires_confirmation"] is False
    assert read_only_command_policy_payload["concurrency_safe"] is True
    assert read_only_command_policy_payload["effective_permission_policy"] == "smart_confirm"
    assert read_only_command_policy_payload["effective_action"] == "allow"
    assert read_only_command_policy_payload["execution_scope"] == "read_only"
    assert read_only_command_policy_payload["execution_route"] == "prefer_agentic_search_or_read_tool"
    assert read_only_command_policy_payload["runtime_ready"] is True
    assert read_only_command_policy_payload["permission_required"] is False
    assert read_only_command_policy_payload["runtime_request_status"] == "ready"
    assert read_only_command_policy_payload["audit_record_preview"]["policy_action"] == "allow"
    assert "cwd" in read_only_command_policy_payload["audit_fields"]
    assert "rg" in read_only_command_policy_payload["detected_commands"]

    diagnostic_command_policy = client.post(
        "/api/workflow/command-policy/check",
        json={"command": "python -m compileall -q backend/app backend/scripts"},
    )
    assert diagnostic_command_policy.status_code == 200, diagnostic_command_policy.text
    diagnostic_command_policy_payload = diagnostic_command_policy.json()
    assert diagnostic_command_policy_payload["risk_level"] == "diagnostic"
    assert diagnostic_command_policy_payload["allowed"] is True
    assert diagnostic_command_policy_payload["requires_confirmation"] is False
    assert diagnostic_command_policy_payload["effective_action"] == "allow"
    assert diagnostic_command_policy_payload["execution_scope"] == "diagnostic"
    assert diagnostic_command_policy_payload["execution_route"] == "diagnostic_runner_after_policy_check"
    assert diagnostic_command_policy_payload["default_timeout_ms"] >= 120000

    network_command_policy = client.post(
        "/api/workflow/command-policy/check",
        json={"command": "python -m pip install some-package"},
    )
    assert network_command_policy.status_code == 200, network_command_policy.text
    network_command_policy_payload = network_command_policy.json()
    assert network_command_policy_payload["risk_level"] == "network"
    assert network_command_policy_payload["allowed"] is True
    assert network_command_policy_payload["requires_confirmation"] is True
    assert network_command_policy_payload["effective_action"] == "confirm"
    assert network_command_policy_payload["execution_scope"] == "network"
    assert network_command_policy_payload["runtime_ready"] is False
    assert network_command_policy_payload["permission_required"] is True
    assert network_command_policy_payload["runtime_request_status"] == "needs_approval"
    assert "批准" in network_command_policy_payload["approval_prompt"]
    assert "network_target" in network_command_policy_payload["audit_fields"]
    assert "network" in network_command_policy_payload["categories"]

    network_full_access_policy = client.post(
        "/api/workflow/command-policy/check",
        json={
            "command": "python -m pip install some-package",
            "permission_policy": "full_access",
        },
    )
    assert network_full_access_policy.status_code == 200, network_full_access_policy.text
    network_full_access_payload = network_full_access_policy.json()
    assert network_full_access_payload["risk_level"] == "network"
    assert network_full_access_payload["effective_permission_policy"] == "full_access"
    assert network_full_access_payload["effective_action"] == "allow"
    assert network_full_access_payload["runtime_ready"] is True

    high_risk_command_policy = client.post(
        "/api/workflow/command-policy/check",
        json={"command": "rm -rf data && git reset --hard", "permission_policy": "full_access"},
    )
    assert high_risk_command_policy.status_code == 200, high_risk_command_policy.text
    high_risk_command_policy_payload = high_risk_command_policy.json()
    assert high_risk_command_policy_payload["risk_level"] == "high_risk"
    assert high_risk_command_policy_payload["allowed"] is False
    assert high_risk_command_policy_payload["requires_confirmation"] is True
    assert high_risk_command_policy_payload["effective_permission_policy"] == "full_access"
    assert high_risk_command_policy_payload["effective_action"] == "block"
    assert high_risk_command_policy_payload["execution_scope"] == "blocked"
    assert high_risk_command_policy_payload["execution_route"] == "blocked_by_command_governance"
    assert high_risk_command_policy_payload["runtime_ready"] is False
    assert high_risk_command_policy_payload["permission_required"] is False
    assert high_risk_command_policy_payload["runtime_request_status"] == "blocked"
    assert high_risk_command_policy_payload["block_reason_code"] == "command_governance_high_risk"
    assert high_risk_command_policy_payload["audit_record_preview"]["rule_ids"]
    assert "shell.rm_recursive_force" in high_risk_command_policy_payload["rule_ids"]
    assert "git.reset_hard" in high_risk_command_policy_payload["rule_ids"]
    assert high_risk_command_policy_payload["destructive_warnings"]
    assert high_risk_command_policy_payload["safer_alternatives"]
    assert high_risk_command_policy_payload["warnings"]

    destructive_git_policy = client.post(
        "/api/workflow/command-policy/check",
        json={"command": "git checkout -- ."},
    )
    assert destructive_git_policy.status_code == 200, destructive_git_policy.text
    destructive_git_payload = destructive_git_policy.json()
    assert destructive_git_payload["risk_level"] == "high_risk"
    assert destructive_git_payload["effective_action"] == "block"
    assert "git.checkout_restore" in destructive_git_payload["rule_ids"]
    assert destructive_git_payload["destructive_warnings"]
    assert destructive_git_payload["safer_alternatives"]

    destructive_database_policy = client.post(
        "/api/workflow/command-policy/check",
        json={"command": "psql -c \"DROP TABLE users;\""},
    )
    assert destructive_database_policy.status_code == 200, destructive_database_policy.text
    destructive_database_payload = destructive_database_policy.json()
    assert destructive_database_payload["risk_level"] == "high_risk"
    assert "database.drop_truncate" in destructive_database_payload["rule_ids"]
    assert destructive_database_payload["allowed"] is False
    assert destructive_database_payload["safer_alternatives"]

    controlled_write_step = WorkflowStep(
        id="permission-policy-write",
        agent="code_agent",
        action="generate_code",
        title="生成受控代码草稿",
        required_permissions=["file_read", "file_write"],
        risk_level="medium",
        requires_confirmation=True,
    )
    network_step = WorkflowStep(
        id="permission-policy-network",
        agent="commander_agent",
        action="fetch_remote_context",
        title="读取联网资料",
        required_permissions=["network"],
        risk_level="medium",
        requires_confirmation=True,
    )
    blocked_shell_step = WorkflowStep(
        id="permission-policy-shell",
        agent="code_agent",
        action="run_command",
        title="执行命令",
        required_permissions=["shell"],
        risk_level="high",
        requires_confirmation=True,
        command_policy=WorkflowCommandPolicy(
            may_run_command=True,
            risk_level="high_risk",
            requires_confirmation=True,
            allowed=False,
        ),
    )
    assert evaluate_permission_policy(
        permission_policy="always_ask", step=controlled_write_step
    ).action == "confirm"
    assert evaluate_permission_policy(
        permission_policy="smart_confirm", step=controlled_write_step
    ).action == "confirm"
    assert evaluate_permission_policy(
        permission_policy="auto_approve", step=controlled_write_step
    ).action == "allow"
    assert evaluate_permission_policy(
        permission_policy="auto_approve", step=network_step
    ).action == "confirm"
    assert evaluate_permission_policy(
        permission_policy="full_access", step=network_step
    ).action == "allow"
    assert evaluate_permission_policy(
        permission_policy="full_access", step=blocked_shell_step
    ).action == "block"
    assert evaluate_permission_policy(
        permission_policy="unknown", step=controlled_write_step
    ).policy == "smart_confirm"

    runtime_preferences = client.get("/api/settings/runtime-preferences")
    assert runtime_preferences.status_code == 200, runtime_preferences.text
    runtime_preferences_payload = runtime_preferences.json()
    assert runtime_preferences_payload["permission_policy"] == "smart_confirm"
    assert runtime_preferences_payload["personality"] == "professional"
    assert runtime_preferences_payload["permission_policy_options"]
    assert runtime_preferences_payload["personality_options"]
    assert "专业稳重" in build_personality_instruction("professional")
    assert "简洁直接" in build_personality_instruction("concise")
    assert "温和耐心" in build_personality_instruction("warm")
    assert "启发性" in build_personality_instruction("creative")
    assert build_personality_instruction("unknown") == build_personality_instruction("professional")

    update_runtime_preferences = client.put(
        "/api/settings/runtime-preferences",
        json={"permission_policy": "auto_approve", "personality": "warm"},
    )
    assert update_runtime_preferences.status_code == 200, update_runtime_preferences.text
    updated_runtime_preferences_payload = update_runtime_preferences.json()
    assert updated_runtime_preferences_payload["permission_policy"] == "auto_approve"
    assert updated_runtime_preferences_payload["personality"] == "warm"
    runtime_preferences_file = settings.data_dir / "runtime_preferences.json"
    assert runtime_preferences_file.exists()
    runtime_preferences_text = runtime_preferences_file.read_text(encoding="utf-8")
    assert "auto_approve" in runtime_preferences_text
    assert "warm" in runtime_preferences_text

    model_config = client.get("/api/models/config")
    assert model_config.status_code == 200, model_config.text
    model_config_payload = model_config.json()
    assert "ciphertext" not in json.dumps(model_config_payload, ensure_ascii=False)

    model_test = client.post(
        "/api/models/test",
        json={
            "provider": "openai_compatible",
            "model": "verify-test-model",
            "thinking": "disabled",
        },
    )
    assert model_test.status_code == 200, model_test.text
    model_test_payload = model_test.json()
    assert model_test_payload["ok"] is False
    assert model_test_payload["provider"] == "openai_compatible"
    assert "Base URL" in model_test_payload["message"]
    assert "ciphertext" not in json.dumps(model_test_payload, ensure_ascii=False)

    fake_api_key = "test-api-key-placeholder"
    update_model_config = client.put(
        "/api/models/config",
        json={
            "provider": "deepseek",
            "base_url": "https://api.deepseek.com",
            "model": "deepseek-v4-flash",
            "thinking": "enabled",
            "api_key": fake_api_key,
        },
    )
    if os.name == "nt":
        assert update_model_config.status_code == 200, update_model_config.text
        updated_config_payload = update_model_config.json()
        updated_config_json = json.dumps(updated_config_payload, ensure_ascii=False)
        assert fake_api_key not in updated_config_json
        assert "ciphertext" not in updated_config_json
        assert updated_config_payload["provider"] == "deepseek"
        assert updated_config_payload["thinking"] == "enabled"
        assert updated_config_payload["api_key_configured"] is True
        assert updated_config_payload["api_key_source"] == "local_config"
        assert updated_config_payload["configuration_source"] == "local_config"

        model_config_file = settings.data_dir / "model_config.json"
        assert model_config_file.exists()
        model_config_text = model_config_file.read_text(encoding="utf-8")
        assert fake_api_key not in model_config_text
        assert "api_key_secrets" in model_config_text
        assert "api_key_secret\"" not in model_config_text

        fake_kimi_key = "test-api-key-placeholder"
        update_kimi_config = client.put(
            "/api/models/config",
            json={
                "provider": "kimi",
                "base_url": "https://api.moonshot.cn/v1",
                "model": "kimi-k2.6",
                "thinking": "enabled",
                "api_key": fake_kimi_key,
            },
        )
        assert update_kimi_config.status_code == 200, update_kimi_config.text
        assert update_kimi_config.json()["provider"] == "kimi"
        assert update_kimi_config.json()["thinking"] == "enabled"
        assert update_kimi_config.json()["api_key_configured"] is True

        # 切回 DeepSeek 不传 Key：两份 provider 密文都必须保留，且不会跨 provider 误用。
        restore_deepseek_config = client.put(
            "/api/models/config",
            json={
                "provider": "deepseek",
                "base_url": "https://api.deepseek.com",
                "model": "deepseek-v4-flash",
                "thinking": "enabled",
            },
        )
        assert restore_deepseek_config.status_code == 200, restore_deepseek_config.text
        assert restore_deepseek_config.json()["provider"] == "deepseek"
        assert restore_deepseek_config.json()["thinking"] == "enabled"
        assert restore_deepseek_config.json()["api_key_configured"] is True

        model_config_text = model_config_file.read_text(encoding="utf-8")
        assert fake_api_key not in model_config_text
        assert fake_kimi_key not in model_config_text
        assert '"deepseek"' in model_config_text
        assert '"kimi"' in model_config_text

        models_after_config = client.get("/api/models/providers")
        assert models_after_config.status_code == 200, models_after_config.text
        models_after_config_payload = models_after_config.json()
        assert models_after_config_payload["current"]["provider"] == "deepseek"
        assert models_after_config_payload["current"]["thinking"] == "enabled"
        assert models_after_config_payload["current"]["api_key_source"] == "local_config"
        assert models_after_config_payload["current"]["configuration_source"] == "local_config"
        provider_keys = {
            item["provider"]: item["api_key_configured"]
            for item in models_after_config_payload["providers"]
        }
        assert provider_keys["deepseek"] is True
        assert provider_keys["kimi"] is True
    else:
        assert update_model_config.status_code == 400, update_model_config.text

    workspace_import = client.post(
        "/api/workspace/documents",
        json={
            "filename": "assignment.md",
            "content": "# 作业要求\n\n请生成 Python 示例和 README 报告。",
        },
    )
    assert workspace_import.status_code == 200, workspace_import.text
    workspace_import_payload = workspace_import.json()
    assert workspace_import_payload["name"] == "assignment.md"
    assert workspace_import_payload["relative_path"] == "assignment.md"
    assert workspace_import_payload["size_bytes"] > 0
    assert "作业要求" in workspace_import_payload["preview"]
    assert (settings.data_dir / "workspaces" / "assignment.md").exists()

    workspace_preview = client.get(
        "/api/workspace/documents/assignment.md",
        params={"preview_chars": 12},
    )
    assert workspace_preview.status_code == 200, workspace_preview.text
    workspace_preview_payload = workspace_preview.json()
    assert workspace_preview_payload["name"] == "assignment.md"
    assert workspace_preview_payload["relative_path"] == "assignment.md"
    assert workspace_preview_payload["preview_chars"] == 12
    assert workspace_preview_payload["truncated"] is True
    assert "作业要求" in workspace_preview_payload["preview"]
    assert "path" not in workspace_preview_payload
    assert str(settings.data_dir) not in json.dumps(workspace_preview_payload, ensure_ascii=False)

    missing_workspace_preview = client.get("/api/workspace/documents/missing.md")
    assert missing_workspace_preview.status_code == 404

    unsafe_workspace_preview = client.get("/api/workspace/documents/unsafe.exe")
    assert unsafe_workspace_preview.status_code == 400

    workspace_duplicate = client.post(
        "/api/workspace/documents",
        json={
            "filename": "assignment.md",
            "content": "# 第二份作业\n\n验证同名文件自动避让。",
        },
    )
    assert workspace_duplicate.status_code == 200, workspace_duplicate.text
    assert workspace_duplicate.json()["name"] == "assignment_1.md"

    workspace_list = client.get("/api/workspace/documents")
    assert workspace_list.status_code == 200, workspace_list.text
    workspace_list_payload = workspace_list.json()
    assert workspace_list_payload["total"] >= 2
    assert {"assignment.md", "assignment_1.md"}.issubset(
        {item["name"] for item in workspace_list_payload["documents"]}
    )

    # 首个正式 Agent 的离线闭环：选定文档 -> 受控读取 -> 结构化结论 -> 可回放审计。
    # 验证脚本强制 mock，不会消耗用户配置的模型额度，但会走与真实模式相同的 Runner/Tool 协议。
    document_agent_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "提取作业的功能、输出和验收要求。",
            "document_refs": ["assignment.md"],
            "output_mode": "requirements",
        },
    )
    assert document_agent_run.status_code == 200, document_agent_run.text
    document_agent_payload = document_agent_run.json()
    assert document_agent_payload["mode"] == "mock"
    assert document_agent_payload["status"] == "completed"
    assert document_agent_payload["workflow_run"]["mode"] == "runtime"
    assert document_agent_payload["workflow_run"]["status"] == "completed"
    document_context = document_agent_payload["document_context"]
    assert document_context["schema_version"] == "agentflow.document_context.v1"
    assert document_context["documents"] == ["assignment.md"]
    assert document_context["requirements"]
    assert document_context["sources"]
    assert all(item["source_refs"] for item in document_context["requirements"])
    assert str(settings.data_dir) not in json.dumps(document_context, ensure_ascii=False)
    document_agent_task_id = document_agent_payload["task_id"]
    document_tool_calls = client.get(f"/api/tasks/{document_agent_task_id}/tool-calls")
    assert document_tool_calls.status_code == 200, document_tool_calls.text
    assert document_tool_calls.json()["total"] == 1
    assert document_tool_calls.json()["tool_calls"][0]["tool_name"] == "document.read_text"
    assert document_tool_calls.json()["tool_calls"][0]["status"] == "completed"
    document_agent_updates = client.get(f"/api/tasks/{document_agent_task_id}/updates")
    assert document_agent_updates.status_code == 200, document_agent_updates.text
    assert any(
        update["event"] == "task_state_snapshot"
        for update in document_agent_updates.json()["updates"]
    )

    # 用户选择“提取需求”时，模型在一次无工具格式修复后仍给不出 JSON，不能丢弃已经
    # 读取的原文。只允许从明确要求词生成带来源条目；这不是 QA/摘要/对比的通用替代。
    fallback_import = client.post(
        "/api/workspace/documents",
        json={
            "filename": "requirements_fallback.md",
            "content": (
                "# 交付约束\n\n"
                "系统必须保留来源位置。\n"
                "输出需要包含可复核的需求清单。\n"
                "不得修改用户导入的原始文档。\n"
                "The service must retain source references for every conclusion.\n"
            ),
        },
    )
    assert fallback_import.status_code == 200, fallback_import.text
    fallback_name = fallback_import.json()["relative_path"]
    original_mock_model = document_agent_service._DeterministicDocumentModel
    document_agent_service._DeterministicDocumentModel = _InvalidRequirementsModel
    try:
        requirements_fallback_run = client.post(
            "/api/agents/document_agent/run",
            json={
                "task_goal": "提取文档中的明确需求与约束。",
                "document_refs": [fallback_name],
                "output_mode": "requirements",
            },
        )
    finally:
        document_agent_service._DeterministicDocumentModel = original_mock_model
    assert requirements_fallback_run.status_code == 200, requirements_fallback_run.text
    requirements_fallback_payload = requirements_fallback_run.json()
    assert requirements_fallback_payload["status"] == "completed"
    assert requirements_fallback_payload["stop_reason"] == "completed_with_conservative_requirements_fallback"
    fallback_context = requirements_fallback_payload["document_context"]
    assert fallback_context["requirements"]
    assert all(item["source_refs"] for item in fallback_context["requirements"])
    assert any(
        "must retain source references" in item["text"].lower()
        and item["priority"] == "must"
        for item in fallback_context["requirements"]
    )
    assert any("保守需求条目" in warning for warning in fallback_context["warnings"])
    assert requirements_fallback_payload["workflow_run"]["status"] == "completed"

    # 精确搜索零命中不等于文档没有相关语义。单文档任务应明确记录零命中后回读材料，
    # 再生成带来源的结论；不能因 answer_source_ids 为空退化成通用 JSON 协议错误。
    no_hit_import = client.post(
        "/api/workspace/documents",
        json={
            "filename": "no_hit_fallback.md",
            "content": "# 可追溯要求\n\n系统需要保留来源位置，便于用户复核结论。\n",
        },
    )
    assert no_hit_import.status_code == 200, no_hit_import.text
    no_hit_name = no_hit_import.json()["relative_path"]
    no_hit_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "说明材料是否包含审计能力，并给出可追溯来源。",
            "document_refs": [no_hit_name],
            "query": "审计",
            "output_mode": "qa",
        },
    )
    assert no_hit_run.status_code == 200, no_hit_run.text
    no_hit_payload = no_hit_run.json()
    assert no_hit_payload["status"] == "completed"
    assert no_hit_payload["document_context"]["sources"]
    assert any(
        "没有精确文本命中" in warning
        for warning in no_hit_payload["document_context"]["warnings"]
    )
    no_hit_calls = client.get(f"/api/tasks/{no_hit_payload['task_id']}/tool-calls")
    assert no_hit_calls.status_code == 200, no_hit_calls.text
    no_hit_call_records = no_hit_calls.json()["tool_calls"]
    assert {item["tool_name"] for item in no_hit_call_records} == {
        "document.search_text",
        "document.read_text",
    }
    # 历史接口按持久化顺序返回，不把 SQL 的排序实现误当成 Agent 的行为契约。
    no_hit_search_call = next(
        item for item in no_hit_call_records if item["tool_name"] == "document.search_text"
    )
    no_hit_read_call = next(
        item for item in no_hit_call_records if item["tool_name"] == "document.read_text"
    )
    assert no_hit_search_call["result"]["total"] == 0
    assert no_hit_search_call["result"]["recommended_fallback_read_path"] == no_hit_name
    assert no_hit_read_call["result"]["source"]["relative_path"] == no_hit_name

    # 文档页的实时体验：/start 必须立即受理任务，WebSocket 再按实际阶段推送，最终结果
    # 仍通过已持久化的结构化协议读取，不能把未经来源校验的模型中间文本直接给 UI。
    document_agent_start = client.post(
        "/api/agents/document_agent/start",
        json={
            "task_goal": "概括这份作业的核心目标。",
            "document_refs": ["assignment.md"],
            "output_mode": "summary",
        },
    )
    assert document_agent_start.status_code == 202, document_agent_start.text
    started_task_id = document_agent_start.json()["task_id"]
    streamed_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{started_task_id}") as websocket:
        while True:
            try:
                streamed_events.append(websocket.receive_json())
            except Exception:
                break
    assert streamed_events
    assert streamed_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in streamed_events)
    assert streamed_events[-1]["event"] == "task_completed"
    document_agent_started_result = client.get(f"/api/agents/document_agent/{started_task_id}/result")
    assert document_agent_started_result.status_code == 200, document_agent_started_result.text
    assert document_agent_started_result.json()["status"] == "completed"
    assert document_agent_started_result.json()["result"]["document_context"]["sources"]

    # 回归：来源展示片段必须受 schema 上限约束。长文档仍可读取并返回可追溯结论， 
    # 不允许在 Tool 成功后因来源 excerpt 过长而触发连续重试。
    long_source_import = client.post(
        "/api/workspace/documents",
        json={
            "filename": "long_source.md",
            "content": "# 长来源验证\n\n" + "这是一段用于验证来源截断的材料。" * 80,
        },
    )
    assert long_source_import.status_code == 200, long_source_import.text
    long_source_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "请总结这份长材料。",
            "document_refs": ["long_source.md"],
            "output_mode": "summary",
        },
    )
    assert long_source_run.status_code == 200, long_source_run.text
    long_source_payload = long_source_run.json()
    assert long_source_payload["status"] == "completed"
    assert long_source_payload["document_context"]["sources"]
    assert all(
        len(source["excerpt"]) <= 360
        for source in long_source_payload["document_context"]["sources"]
    )

    # 文档工作台 v1 的解析扩展：PDF/DOCX 仍走同一份受控 workspace 协议，但必须在导入、
    # 预览、精确定位与最终 Agent 来源四个环节保持“页码/段落”事实，不能伪装成文本行号。
    pdf = fitz.open()
    pdf_page_one = pdf.new_page()
    pdf_page_one.insert_text((72, 72), "PDF requirement: preserve page source.")
    pdf_page_two = pdf.new_page()
    pdf_page_two.insert_text((72, 72), "PDF acceptance: show page two.")
    pdf_import = client.post(
        "/api/workspace/documents",
        json={
            "filename": "document_agent_pdf_verification.pdf",
            "content_base64": base64.b64encode(pdf.tobytes()).decode("ascii"),
        },
    )
    pdf.close()
    assert pdf_import.status_code == 200, pdf_import.text
    pdf_name = pdf_import.json()["relative_path"]
    assert pdf_import.json()["document_type"] == "pdf"
    pdf_preview = client.get(f"/api/workspace/documents/{pdf_name}")
    assert pdf_preview.status_code == 200, pdf_preview.text
    assert pdf_preview.json()["document_type"] == "pdf"
    assert "PDF acceptance" in pdf_preview.json()["preview"]
    pdf_search = client.post("/api/workspace/search", json={"query": "acceptance"})
    assert pdf_search.status_code == 200, pdf_search.text
    assert pdf_search.json()["matches"][0]["source_kind"] == "page"
    assert pdf_search.json()["matches"][0]["source_locator"] == "第 2 页"
    pdf_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "提取 PDF 的需求与验收信息，并给出来源。",
            "document_refs": [pdf_name],
            "output_mode": "requirements",
        },
    )
    assert pdf_run.status_code == 200, pdf_run.text
    assert pdf_run.json()["status"] == "completed"
    assert all(
        source["source_kind"] == "page"
        and source["source_locator"].startswith("第 ")
        and source["source_locator"].endswith("页")
        for source in pdf_run.json()["document_context"]["sources"]
    )

    docx = DocxDocument()
    docx.add_paragraph("DOCX 第一段：系统必须记录可追溯来源。")
    docx.add_paragraph("DOCX 第二段：验收需要显示对应段落位置。")
    docx_bytes = BytesIO()
    docx.save(docx_bytes)
    docx_import = client.post(
        "/api/workspace/documents",
        json={
            "filename": "document_agent_docx_verification.docx",
            "content_base64": base64.b64encode(docx_bytes.getvalue()).decode("ascii"),
        },
    )
    assert docx_import.status_code == 200, docx_import.text
    docx_name = docx_import.json()["relative_path"]
    assert docx_import.json()["document_type"] == "docx"
    docx_search = client.post("/api/workspace/search", json={"query": "验收"})
    assert docx_search.status_code == 200, docx_search.text
    docx_match = next(
        match for match in docx_search.json()["matches"] if match["relative_path"] == docx_name
    )
    assert docx_match["source_kind"] == "paragraph"
    assert docx_match["source_locator"] == "第 2 段"
    docx_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "提取 Word 材料的需求和验收项，并给出段落来源。",
            "document_refs": [docx_name],
            "output_mode": "requirements",
        },
    )
    assert docx_run.status_code == 200, docx_run.text
    assert docx_run.json()["status"] == "completed"
    assert all(
        source["source_kind"] == "paragraph"
        and source["source_locator"].startswith("第 ")
        and source["source_locator"].endswith("段")
        for source in docx_run.json()["document_context"]["sources"]
    )
    # 调度台导入入口会复用同一份文件选择器；Commander 至少要能把 PDF/DOCX 文件名交给
    # 已就绪的文档助手，而不是因为旧正则只识别 txt/markdown 而要求用户重新输入。
    docx_commander = client.post(
        "/api/chat",
        json={"message": f"请读取 {docx_name}，归纳其中的验收要求。"},
    )
    assert docx_commander.status_code == 200, docx_commander.text
    docx_commander_step = next(
        step
        for step in docx_commander.json()["workflow_plan"]["steps"]
        if step["agent"] == "document_agent"
    )
    assert docx_commander_step["input"]["document_refs"] == [docx_name], (
        docx_name,
        docx_commander_step["input"],
    )

    # 多文档对比必须逐份读取，并且每个比较结论至少同时引用两份不同材料。
    document_comparison_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "比较两份作业材料的共识和差异，并保留来源。",
            "document_refs": ["assignment.md", "assignment_1.md"],
            "output_mode": "comparison",
        },
    )
    assert document_comparison_run.status_code == 200, document_comparison_run.text
    document_comparison_payload = document_comparison_run.json()
    assert document_comparison_payload["status"] == "completed"
    comparison_context = document_comparison_payload["document_context"]
    assert comparison_context["documents"] == ["assignment.md", "assignment_1.md"]
    assert comparison_context["comparisons"]
    for comparison in comparison_context["comparisons"]:
        assert len({source["relative_path"] for source in comparison["source_refs"]}) >= 2
    comparison_tool_calls = client.get(
        f"/api/tasks/{document_comparison_payload['task_id']}/tool-calls"
    )
    assert comparison_tool_calls.status_code == 200, comparison_tool_calls.text
    assert [item["tool_name"] for item in comparison_tool_calls.json()["tool_calls"]] == [
        "document.read_text",
        "document.read_text",
    ]

    # 跨文档问答与“多文档对比”共用逐份读取和来源覆盖规则，但用户只需要问题答案，
    # 不应被强迫生成共识/差异报告。answer 的来源必须直接覆盖两份材料，而非仅让某个
    # 附带条目碰巧引用多份材料后就放行。
    document_cross_qa_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "两份作业材料都要求保留来源吗？请说明依据。",
            "document_refs": ["assignment.md", "assignment_1.md"],
            "output_mode": "cross_qa",
        },
    )
    assert document_cross_qa_run.status_code == 200, document_cross_qa_run.text
    document_cross_qa_payload = document_cross_qa_run.json()
    assert document_cross_qa_payload["status"] == "completed"
    cross_qa_context = document_cross_qa_payload["document_context"]
    assert not cross_qa_context["comparisons"]
    assert {
        source["relative_path"]
        for source in cross_qa_context["sources"]
    } >= {"assignment.md", "assignment_1.md"}
    cross_qa_tool_calls = client.get(
        f"/api/tasks/{document_cross_qa_payload['task_id']}/tool-calls"
    )
    assert cross_qa_tool_calls.status_code == 200, cross_qa_tool_calls.text
    assert [item["tool_name"] for item in cross_qa_tool_calls.json()["tool_calls"]] == [
        "document.read_text",
        "document.read_text",
    ]

    cross_qa_needs_more_material = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "比较这份材料与另一份材料的权限要求。",
            "document_refs": ["assignment.md"],
            "output_mode": "cross_qa",
        },
    )
    assert cross_qa_needs_more_material.status_code == 200, cross_qa_needs_more_material.text
    assert cross_qa_needs_more_material.json()["status"] == "needs_clarification"
    assert "至少需要选择两份材料" in cross_qa_needs_more_material.json()["reply"]

    # 跨文档整合是只读的“合并证据与结构化条目”，不是生成或覆盖文件。确定性模式只对
    # 完全相同的可见条目做字面归并，并应把两份材料的来源保留在同一 requirement 中。
    synthesis_refs: list[str] = []
    for filename, content in (
        (
            "synthesis_a.md",
            "# 材料 A\n\n系统必须保留可追溯来源。\n接口调用需要记录任务标识。\n",
        ),
        (
            "synthesis_b.md",
            "# 材料 B\n\n系统必须保留可追溯来源。\n高风险操作不得绕过权限确认。\n",
        ),
    ):
        imported = client.post("/api/workspace/documents", json={"filename": filename, "content": content})
        assert imported.status_code == 200, imported.text
        synthesis_refs.append(imported.json()["relative_path"])

    document_synthesis_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "整合两份材料的需求，合并重复条目并保留来源。",
            "document_refs": synthesis_refs,
            "output_mode": "synthesis",
        },
    )
    assert document_synthesis_run.status_code == 200, document_synthesis_run.text
    document_synthesis_payload = document_synthesis_run.json()
    assert document_synthesis_payload["status"] == "completed"
    synthesis_context = document_synthesis_payload["document_context"]
    assert synthesis_context["requirements"]
    merged_requirement = next(
        item
        for item in synthesis_context["requirements"]
        if "必须保留可追溯来源" in item["text"]
    )
    assert {
        source["relative_path"]
        for source in merged_requirement["source_refs"]
    } == set(synthesis_refs)
    assert not synthesis_context["comparisons"]
    synthesis_tool_calls = client.get(
        f"/api/tasks/{document_synthesis_payload['task_id']}/tool-calls"
    )
    assert synthesis_tool_calls.status_code == 200, synthesis_tool_calls.text
    assert [item["tool_name"] for item in synthesis_tool_calls.json()["tool_calls"]] == [
        "document.read_text",
        "document.read_text",
    ]

    # 关键信息卡是单文档只读模板提取，不应退化成散乱的 requirements。每个字段仍须携带
    # 实际来源，后续 Commander 才能安全消费这些机器可读字段。
    brief_import = client.post(
        "/api/workspace/documents",
        json={
            "filename": "project_brief.md",
            "content": (
                "# 星河平台升级\n\n"
                "项目目标：为运营团队统一处理客户工单。\n"
                "项目范围：覆盖工单创建、分派和状态查询，不包含外部支付。\n"
                "相关角色：运营负责人、客服人员和客户管理员。\n"
                "交付物：Web 管理界面、接口说明和验收记录。\n"
                "时间节点：第一阶段在 2026-08-15 前完成。\n"
                "风险：依赖旧工单系统提供稳定的数据接口。\n"
            ),
        },
    )
    assert brief_import.status_code == 200, brief_import.text
    brief_reference = brief_import.json()["relative_path"]
    document_brief_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "按关键信息卡梳理这份项目材料。",
            "document_refs": [brief_reference],
            "output_mode": "brief",
        },
    )
    assert document_brief_run.status_code == 200, document_brief_run.text
    document_brief_payload = document_brief_run.json()
    assert document_brief_payload["status"] == "completed"
    brief_context = document_brief_payload["document_context"]
    brief_fields = {item["key"]: item for item in brief_context["brief_fields"]}
    assert {"subject", "purpose", "scope", "stakeholders", "deliverables", "milestones", "risks"} <= set(brief_fields)
    assert "星河平台升级" in brief_fields["subject"]["value"]
    assert "运营团队" in brief_fields["purpose"]["value"]
    assert not brief_context["requirements"]
    assert all(
        item["source_refs"][0]["relative_path"] == brief_reference
        for item in brief_fields.values()
    )
    brief_tool_calls = client.get(f"/api/tasks/{document_brief_payload['task_id']}/tool-calls")
    assert brief_tool_calls.status_code == 200, brief_tool_calls.text
    assert [item["tool_name"] for item in brief_tool_calls.json()["tool_calls"]] == [
        "document.read_text"
    ]

    # 真实 Provider 偶发漏字段或夹带自然语言时，关键信息卡不能让用户面对一张空白结果。
    # 仅在读取已成功且字段可由原文显式标记得出时，才允许只读保守降级。
    original_mock_model = document_agent_service._DeterministicDocumentModel
    document_agent_service._DeterministicDocumentModel = _InvalidBriefModel
    try:
        brief_fallback_run = client.post(
            "/api/agents/document_agent/run",
            json={
                "task_goal": "按关键信息卡梳理这份项目材料。",
                "document_refs": [brief_reference],
                "output_mode": "brief",
            },
        )
    finally:
        document_agent_service._DeterministicDocumentModel = original_mock_model
    assert brief_fallback_run.status_code == 200, brief_fallback_run.text
    brief_fallback_payload = brief_fallback_run.json()
    assert brief_fallback_payload["status"] == "completed"
    assert brief_fallback_payload["stop_reason"].startswith(
        "completed_with_conservative_brief_fallback:"
    )
    brief_fallback_context = brief_fallback_payload["document_context"]
    assert brief_fallback_context["brief_fields"]
    assert all(item["source_refs"] for item in brief_fallback_context["brief_fields"])
    assert any("保守结果" in warning for warning in brief_fallback_context["warnings"])
    brief_fallback_tool_calls = client.get(
        f"/api/tasks/{brief_fallback_payload['task_id']}/tool-calls"
    )
    assert brief_fallback_tool_calls.status_code == 200, brief_fallback_tool_calls.text
    assert [item["tool_name"] for item in brief_fallback_tool_calls.json()["tool_calls"]] == [
        "document.read_text"
    ]

    # 结构化大纲只生成待审阅的章节蓝图，不创建文件。它同样必须带来源、只读一次所选材料，
    # 并在模型协议不稳定时退回原文显式线索，避免客户面对空白的“创作前计划”。
    outline_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "为这份项目材料制定可审阅的结构化文档大纲。",
            "document_refs": [brief_reference],
            "output_mode": "outline",
        },
    )
    assert outline_run.status_code == 200, outline_run.text
    outline_payload = outline_run.json()
    assert outline_payload["status"] == "completed"
    outline_context = outline_payload["document_context"]
    assert outline_context["outline_sections"]
    assert not outline_context["requirements"]
    assert all(item["source_refs"] for item in outline_context["outline_sections"])
    assert all(
        item["source_refs"][0]["relative_path"] == brief_reference
        for item in outline_context["outline_sections"]
    )
    outline_tool_calls = client.get(f"/api/tasks/{outline_payload['task_id']}/tool-calls")
    assert outline_tool_calls.status_code == 200, outline_tool_calls.text
    assert [item["tool_name"] for item in outline_tool_calls.json()["tool_calls"]] == [
        "document.read_text"
    ]

    original_mock_model = document_agent_service._DeterministicDocumentModel
    document_agent_service._DeterministicDocumentModel = _InvalidOutlineModel
    try:
        outline_fallback_run = client.post(
            "/api/agents/document_agent/run",
            json={
                "task_goal": "为这份项目材料制定可审阅的结构化文档大纲。",
                "document_refs": [brief_reference],
                "output_mode": "outline",
            },
        )
    finally:
        document_agent_service._DeterministicDocumentModel = original_mock_model
    assert outline_fallback_run.status_code == 200, outline_fallback_run.text
    outline_fallback_payload = outline_fallback_run.json()
    assert outline_fallback_payload["status"] == "completed"
    assert outline_fallback_payload["stop_reason"].startswith(
        "completed_with_conservative_outline_fallback:"
    )
    outline_fallback_context = outline_fallback_payload["document_context"]
    assert outline_fallback_context["outline_sections"]
    assert all(item["source_refs"] for item in outline_fallback_context["outline_sections"])
    assert any("保守只读蓝图" in warning for warning in outline_fallback_context["warnings"])

    # 正式创作的首步只提供可审阅草稿，不触发文件写入。草稿正文必须逐章节保留来源；与
    # 字段/大纲不同，模型 JSON 连续失效时不能由规则擅自补写正文，必须如实失败。
    draft_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "为这份项目材料生成可审阅的 Markdown 草稿预览。",
            "document_refs": [brief_reference],
            "output_mode": "draft",
        },
    )
    assert draft_run.status_code == 200, draft_run.text
    draft_payload = draft_run.json()
    assert draft_payload["status"] == "completed"
    draft_context = draft_payload["document_context"]
    assert draft_context["draft_title"] == "星河平台升级"
    assert draft_context["draft_sections"]
    base_version = draft_context["draft_version"]
    assert base_version["version_id"] == draft_payload["task_id"]
    assert base_version["root_task_id"] == draft_payload["task_id"]
    assert base_version["parent_task_id"] == ""
    assert base_version["kind"] == "base_draft"
    assert not draft_context["requirements"]
    assert all(item["body"] and item["source_refs"] for item in draft_context["draft_sections"])
    draft_tool_calls = client.get(f"/api/tasks/{draft_payload['task_id']}/tool-calls")
    assert draft_tool_calls.status_code == 200, draft_tool_calls.text
    assert [item["tool_name"] for item in draft_tool_calls.json()["tool_calls"]] == [
        "document.read_text"
    ]

    # 分章节创作从已完成草稿的稳定章节 ID 派生新任务：它会重新读取同一受控材料，只返回
    # 一个固定身份的章节预览，不修改原草稿、不写入文件。错误章节必须明确拒绝，不能被模型
    # 当成新的标题或自由指令继续执行。
    missing_section_start = client.post(
        f"/api/agents/document_agent/{draft_payload['task_id']}/draft-sections/start",
        json={"section_id": "missing_section", "instruction": "请补充本章说明。"},
    )
    assert missing_section_start.status_code == 404, missing_section_start.text

    # 批量修订需要两个不重叠的原文片段；选择含多个段落的章节，避免测试依赖草稿章节顺序。
    source_section = next(
        (item for item in draft_context["draft_sections"] if "\n\n" in item["body"]),
        draft_context["draft_sections"][0],
    )
    section_draft_start = client.post(
        f"/api/agents/document_agent/{draft_payload['task_id']}/draft-sections/start",
        json={
            "section_id": source_section["id"],
            "instruction": "请面向项目负责人扩展本章说明，保持材料事实且不要加入未给出的承诺。",
        },
    )
    assert section_draft_start.status_code == 202, section_draft_start.text
    section_draft_task_id = section_draft_start.json()["task_id"]
    section_draft_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{section_draft_task_id}") as websocket:
        while True:
            try:
                section_draft_events.append(websocket.receive_json())
            except Exception:
                break
    assert section_draft_events
    assert section_draft_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in section_draft_events)
    assert section_draft_events[-1]["event"] == "task_completed"

    section_draft_result = client.get(
        f"/api/agents/document_agent/{section_draft_task_id}/result"
    )
    assert section_draft_result.status_code == 200, section_draft_result.text
    section_draft_payload = section_draft_result.json()["result"]
    assert section_draft_payload["status"] == "completed"
    section_draft_context = section_draft_payload["document_context"]
    assert len(section_draft_context["draft_sections"]) == 1
    assert section_draft_context["draft_sections"][0]["id"] == source_section["id"]
    assert section_draft_context["draft_sections"][0]["heading"] == source_section["heading"]
    assert section_draft_context["draft_sections"][0]["source_refs"]
    assert "分章节创作预览" in section_draft_context["draft_title"]
    section_draft_calls = client.get(f"/api/tasks/{section_draft_task_id}/tool-calls")
    assert section_draft_calls.status_code == 200, section_draft_calls.text
    assert [item["tool_name"] for item in section_draft_calls.json()["tool_calls"]] == [
        "document.read_text"
    ]

    # 草稿事实核验同样从已完成草稿派生，但它不续写任何章节。异步终态必须保留原草稿，
    # 并把材料可支持的表述和待确认问题写进独立任务，供 Qt 详情页和历史审计同时读取。
    draft_review_start = client.post(
        f"/api/agents/document_agent/{draft_payload['task_id']}/draft-review/start",
        json={"focus": "只核验草稿表述能否由材料支持，不改写正文。"},
    )
    assert draft_review_start.status_code == 202, draft_review_start.text
    draft_review_task_id = draft_review_start.json()["task_id"]
    draft_review_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{draft_review_task_id}") as websocket:
        while True:
            try:
                draft_review_events.append(websocket.receive_json())
            except Exception:
                break
    assert draft_review_events
    assert draft_review_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in draft_review_events)
    assert draft_review_events[-1]["event"] == "task_completed"

    draft_review_result = client.get(
        f"/api/agents/document_agent/{draft_review_task_id}/result"
    )
    assert draft_review_result.status_code == 200, draft_review_result.text
    draft_review_payload = draft_review_result.json()["result"]
    assert draft_review_payload["status"] == "completed"
    draft_review_context = draft_review_payload["document_context"]
    assert draft_review_context["review_target_title"]
    assert draft_review_context["draft_title"] == draft_context["draft_title"]
    assert draft_review_context["draft_sections"] == draft_context["draft_sections"]
    assert draft_review_context["constraints"]
    assert all(item["source_refs"] for item in draft_review_context["constraints"])
    assert not draft_review_context["requirements"]
    assert draft_review_payload["workflow_run"]["steps"][-1]["action"] == "review_draft_facts"
    draft_review_calls = client.get(f"/api/tasks/{draft_review_task_id}/tool-calls")
    assert draft_review_calls.status_code == 200, draft_review_calls.text
    assert [item["tool_name"] for item in draft_review_calls.json()["tool_calls"]] == [
        "document.read_text"
    ]
    draft_review_artifacts = client.get(f"/api/tasks/{draft_review_task_id}/artifacts")
    assert draft_review_artifacts.status_code == 200, draft_review_artifacts.text
    assert not draft_review_artifacts.json()["artifacts"]

    # 本章审校只返回候选建议，不是“把模型回复直接覆盖进 Markdown”。它从同一原草稿恢复
    # 章节身份和完整快照，重新读取材料后给出带来源的建议；错误章节仍必须在启动前拒绝。
    missing_section_review = client.post(
        f"/api/agents/document_agent/{draft_payload['task_id']}/draft-sections/review/start",
        json={"section_id": "missing_section", "focus": "检查清晰度。"},
    )
    assert missing_section_review.status_code == 404, missing_section_review.text

    section_review_start = client.post(
        f"/api/agents/document_agent/{draft_payload['task_id']}/draft-sections/review/start",
        json={
            "section_id": source_section["id"],
            "focus": "检查本章表述是否清晰、是否引入材料外事实；只给建议，不改写正文。",
        },
    )
    assert section_review_start.status_code == 202, section_review_start.text
    section_review_task_id = section_review_start.json()["task_id"]
    section_review_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{section_review_task_id}") as websocket:
        while True:
            try:
                section_review_events.append(websocket.receive_json())
            except Exception:
                break
    assert section_review_events
    assert section_review_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in section_review_events)
    assert section_review_events[-1]["event"] == "task_completed"

    section_review_result = client.get(
        f"/api/agents/document_agent/{section_review_task_id}/result"
    )
    assert section_review_result.status_code == 200, section_review_result.text
    section_review_payload = section_review_result.json()["result"]
    assert section_review_payload["status"] == "completed"
    section_review_context = section_review_payload["document_context"]
    assert section_review_context["revision_target_section_id"] == source_section["id"]
    assert "本章审校" in section_review_context["revision_target_title"]
    assert section_review_context["draft_title"] == draft_context["draft_title"]
    assert section_review_context["draft_sections"] == draft_context["draft_sections"]
    section_review_version = section_review_context["draft_version"]
    assert section_review_version["root_task_id"] == draft_payload["task_id"]
    assert section_review_version["parent_task_id"] == draft_payload["task_id"]
    assert section_review_version["kind"] == "section_review"
    assert section_review_context["revision_suggestions"]
    assert all(item["source_refs"] for item in section_review_context["revision_suggestions"])
    assert all(
        "".join(item["original_excerpt"].split())
        in "".join(source_section["body"].split())
        for item in section_review_context["revision_suggestions"]
    )
    assert section_review_payload["workflow_run"]["steps"][-1]["action"] == "review_draft_section"
    section_review_calls = client.get(f"/api/tasks/{section_review_task_id}/tool-calls")
    assert section_review_calls.status_code == 200, section_review_calls.text
    assert [item["tool_name"] for item in section_review_calls.json()["tool_calls"]] == [
        "document.read_text"
    ]
    section_review_artifacts = client.get(f"/api/tasks/{section_review_task_id}/artifacts")
    assert section_review_artifacts.status_code == 200, section_review_artifacts.text
    assert not section_review_artifacts.json()["artifacts"]

    # “应用建议”的第一版不是覆盖旧稿：用户只能从本章审校结果选择一个稳定 suggestion_id，
    # Runtime 精确替换唯一原文片段并生成新的任务快照。它不调用模型/Tool、不写文件；原草稿
    # 和审校任务都必须保持原样，后续仍需通过 save-draft 的明确确认另存新 Markdown。
    missing_revision_preview = client.post(
        f"/api/agents/document_agent/{section_review_task_id}/draft-sections/revision-preview/start",
        json={"suggestion_id": "missing_suggestion"},
    )
    assert missing_revision_preview.status_code == 404, missing_revision_preview.text

    selected_suggestion = section_review_context["revision_suggestions"][0]
    revision_preview_start = client.post(
        f"/api/agents/document_agent/{section_review_task_id}/draft-sections/revision-preview/start",
        json={"suggestion_id": selected_suggestion["id"]},
    )
    assert revision_preview_start.status_code == 202, revision_preview_start.text
    revision_preview_task_id = revision_preview_start.json()["task_id"]
    revision_preview_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{revision_preview_task_id}") as websocket:
        while True:
            try:
                revision_preview_events.append(websocket.receive_json())
            except Exception:
                break
    assert revision_preview_events
    assert revision_preview_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in revision_preview_events)
    assert revision_preview_events[-1]["event"] == "task_completed"

    revision_preview_result = client.get(
        f"/api/agents/document_agent/{revision_preview_task_id}/result"
    )
    assert revision_preview_result.status_code == 200, revision_preview_result.text
    revision_preview_payload = revision_preview_result.json()["result"]
    assert revision_preview_payload["status"] == "completed"
    revision_preview_context = revision_preview_payload["document_context"]
    revision_preview = revision_preview_context["revision_preview"]
    assert revision_preview["source_review_task_id"] == section_review_task_id
    assert revision_preview["suggestion_id"] == selected_suggestion["id"]
    assert revision_preview["section_id"] == source_section["id"]
    assert revision_preview["original_body"] == source_section["body"]
    assert revision_preview["revised_body"] != source_section["body"]
    assert revision_preview_context["draft_title"] == draft_context["draft_title"]
    assert len(revision_preview_context["draft_sections"]) == len(draft_context["draft_sections"])
    revision_version = revision_preview_context["draft_version"]
    assert revision_version["root_task_id"] == draft_payload["task_id"]
    assert revision_version["parent_task_id"] == section_review_task_id
    assert revision_version["kind"] == "revision_preview"
    revised_section = next(
        item for item in revision_preview_context["draft_sections"] if item["id"] == source_section["id"]
    )
    assert revised_section["body"] == revision_preview["revised_body"]
    assert selected_suggestion["original_excerpt"] not in revised_section["body"]
    assert revision_preview_payload["workflow_run"]["steps"][-1]["action"] == "create_section_revision_preview"
    revision_preview_calls = client.get(f"/api/tasks/{revision_preview_task_id}/tool-calls")
    assert revision_preview_calls.status_code == 200, revision_preview_calls.text
    assert not revision_preview_calls.json()["tool_calls"]
    revision_preview_artifacts = client.get(f"/api/tasks/{revision_preview_task_id}/artifacts")
    assert revision_preview_artifacts.status_code == 200, revision_preview_artifacts.text
    assert not revision_preview_artifacts.json()["artifacts"]
    # 再取原审校任务，确保派生版本没有回写已有审校快照或原草稿。
    original_review_after_preview = client.get(
        f"/api/agents/document_agent/{section_review_task_id}/result"
    )
    assert original_review_after_preview.status_code == 200, original_review_after_preview.text
    assert original_review_after_preview.json()["result"]["document_context"]["draft_sections"] == draft_context["draft_sections"]

    # 多建议预览沿用同一份审校快照，只接受同章节、唯一且互不重叠的候选片段。它按原文位置
    # 倒序替换，避免前段变长或变短影响后段位置；任意冲突应在后台启动前安全拒绝。
    assert len(section_review_context["revision_suggestions"]) >= 2
    batch_suggestions = section_review_context["revision_suggestions"][:2]
    duplicate_batch_revision = client.post(
        f"/api/agents/document_agent/{section_review_task_id}/draft-sections/revision-batch-preview/start",
        json={"suggestion_ids": [batch_suggestions[0]["id"], batch_suggestions[0]["id"]]},
    )
    assert duplicate_batch_revision.status_code == 422, duplicate_batch_revision.text
    batch_revision_start = client.post(
        f"/api/agents/document_agent/{section_review_task_id}/draft-sections/revision-batch-preview/start",
        json={"suggestion_ids": [item["id"] for item in batch_suggestions]},
    )
    assert batch_revision_start.status_code == 202, batch_revision_start.text
    batch_revision_task_id = batch_revision_start.json()["task_id"]
    batch_revision_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{batch_revision_task_id}") as websocket:
        while True:
            try:
                batch_revision_events.append(websocket.receive_json())
            except Exception:
                break
    assert batch_revision_events
    assert batch_revision_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in batch_revision_events)
    assert batch_revision_events[-1]["event"] == "task_completed"
    batch_revision_result = client.get(
        f"/api/agents/document_agent/{batch_revision_task_id}/result"
    )
    assert batch_revision_result.status_code == 200, batch_revision_result.text
    batch_revision_payload = batch_revision_result.json()["result"]
    assert batch_revision_payload["status"] == "completed"
    batch_revision_context = batch_revision_payload["document_context"]
    batch_preview = batch_revision_context["revision_preview"]
    assert batch_preview["source_review_task_id"] == section_review_task_id
    assert batch_preview["suggestion_ids"] == [item["id"] for item in batch_suggestions]
    assert batch_preview["section_id"] == source_section["id"]
    assert batch_preview["original_body"] == source_section["body"]
    assert batch_preview["revised_body"] != source_section["body"]
    batch_version = batch_revision_context["draft_version"]
    assert batch_version["root_task_id"] == draft_payload["task_id"]
    assert batch_version["parent_task_id"] == section_review_task_id
    assert batch_version["kind"] == "revision_batch_preview"
    for suggestion in batch_suggestions:
        assert suggestion["original_excerpt"] not in batch_preview["revised_body"]
    assert batch_revision_payload["workflow_run"]["steps"][-1]["action"] == "create_section_revision_batch_preview"
    batch_revision_calls = client.get(f"/api/tasks/{batch_revision_task_id}/tool-calls")
    assert batch_revision_calls.status_code == 200, batch_revision_calls.text
    assert not batch_revision_calls.json()["tool_calls"]
    batch_revision_artifacts = client.get(f"/api/tasks/{batch_revision_task_id}/artifacts")
    assert batch_revision_artifacts.status_code == 200, batch_revision_artifacts.text
    assert not batch_revision_artifacts.json()["artifacts"]

    # 同根版本合并采用三方比较，而不是把候选版本直接覆盖到当前版本。两个修订预览都从同一
    # 审校快照派生，且它们对同一章节有不同修改，因此必须出现冲突；未提交逐项选择时后端
    # 应安全拒绝。这里选择保留当前版本，验证合并只会创建新的可审计预览，不会回写任一旧稿。
    merge_candidates = client.get(
        f"/api/agents/document_agent/{revision_preview_task_id}/merge-candidates"
    )
    assert merge_candidates.status_code == 200, merge_candidates.text
    merge_candidates_payload = merge_candidates.json()
    assert merge_candidates_payload["task_id"] == revision_preview_task_id
    assert merge_candidates_payload["root_task_id"] == draft_payload["task_id"]
    assert batch_revision_task_id in {
        item["task_id"] for item in merge_candidates_payload["candidates"]
    }

    merge_plan_response = client.get(
        f"/api/agents/document_agent/{revision_preview_task_id}/merge-plan/{batch_revision_task_id}"
    )
    assert merge_plan_response.status_code == 200, merge_plan_response.text
    merge_plan = merge_plan_response.json()
    assert merge_plan["primary_task_id"] == revision_preview_task_id
    assert merge_plan["secondary_task_id"] == batch_revision_task_id
    assert merge_plan["common_ancestor_task_id"] == section_review_task_id
    assert merge_plan["conflicts"], "两条不同修订分支必须要求用户选择冲突方案。"

    merge_without_resolution = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/merge-preview/start",
        json={"other_task_id": batch_revision_task_id},
    )
    assert merge_without_resolution.status_code == 400, merge_without_resolution.text
    merge_with_unknown_resolution = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/merge-preview/start",
        json={
            "other_task_id": batch_revision_task_id,
            "resolutions": [{"conflict_id": "unknown_conflict", "choice": "primary"}],
        },
    )
    assert merge_with_unknown_resolution.status_code == 400, merge_with_unknown_resolution.text
    merge_with_extra_field = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/merge-preview/start",
        json={"other_task_id": batch_revision_task_id, "draft_body": "客户端不得提交正文"},
    )
    assert merge_with_extra_field.status_code == 422, merge_with_extra_field.text

    merge_resolutions = [
        {"conflict_id": item["conflict_id"], "choice": "primary"}
        for item in merge_plan["conflicts"]
    ]
    merge_preview_start = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/merge-preview/start",
        json={"other_task_id": batch_revision_task_id, "resolutions": merge_resolutions},
    )
    assert merge_preview_start.status_code == 202, merge_preview_start.text
    merge_preview_task_id = merge_preview_start.json()["task_id"]
    merge_preview_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{merge_preview_task_id}") as websocket:
        while True:
            try:
                merge_preview_events.append(websocket.receive_json())
            except Exception:
                break
    assert merge_preview_events
    assert merge_preview_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in merge_preview_events)
    assert any(event["event"] == "draft_merging" for event in merge_preview_events)
    assert merge_preview_events[-1]["event"] == "task_completed"

    merge_preview_result = client.get(
        f"/api/agents/document_agent/{merge_preview_task_id}/result"
    )
    assert merge_preview_result.status_code == 200, merge_preview_result.text
    merge_preview_payload = merge_preview_result.json()["result"]
    assert merge_preview_payload["status"] == "completed"
    merge_preview_context = merge_preview_payload["document_context"]
    merge_preview = merge_preview_context["merge_preview"]
    assert merge_preview["primary_task_id"] == revision_preview_task_id
    assert merge_preview["secondary_task_id"] == batch_revision_task_id
    assert merge_preview["common_ancestor_task_id"] == section_review_task_id
    assert merge_preview["resolved_conflict_count"] == len(merge_plan["conflicts"])
    assert merge_preview_context["draft_sections"] == revision_preview_context["draft_sections"]
    merged_version = merge_preview_context["draft_version"]
    assert merged_version["root_task_id"] == draft_payload["task_id"]
    assert merged_version["parent_task_id"] == revision_preview_task_id
    assert merged_version["kind"] == "merge_preview"
    assert merge_preview_payload["workflow_run"]["steps"][-1]["action"] == "create_draft_merge_preview"
    merge_preview_calls = client.get(f"/api/tasks/{merge_preview_task_id}/tool-calls")
    assert merge_preview_calls.status_code == 200, merge_preview_calls.text
    assert not merge_preview_calls.json()["tool_calls"]
    merge_preview_artifacts = client.get(f"/api/tasks/{merge_preview_task_id}/artifacts")
    assert merge_preview_artifacts.status_code == 200, merge_preview_artifacts.text
    assert not merge_preview_artifacts.json()["artifacts"]
    revision_after_merge = client.get(
        f"/api/agents/document_agent/{revision_preview_task_id}/result"
    )
    assert revision_after_merge.status_code == 200, revision_after_merge.text
    assert revision_after_merge.json()["result"]["document_context"]["draft_sections"] == revision_preview_context["draft_sections"]
    batch_after_merge = client.get(
        f"/api/agents/document_agent/{batch_revision_task_id}/result"
    )
    assert batch_after_merge.status_code == 200, batch_after_merge.text
    assert batch_after_merge.json()["result"]["document_context"]["draft_sections"] == batch_revision_context["draft_sections"]

    # 版本恢复不能把“回退”实现成覆盖旧任务或文件。客户端只给出历史 task_id，Runtime 从
    # SQLite 重新取已验证快照，生成一个零模型、零 Tool、零写入的新任务；这里选择已经修订
    # 过的历史版本，验证恢复的是该版本正文而非根草稿或当前建议列表。
    missing_restore_preview = client.post(
        "/api/agents/document_agent/task_document_missing/restore-preview/start",
        json={},
    )
    assert missing_restore_preview.status_code == 404, missing_restore_preview.text
    restore_preview_with_untrusted_body = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/restore-preview/start",
        json={"draft_title": "这不是恢复接口的输入"},
    )
    assert restore_preview_with_untrusted_body.status_code == 422, restore_preview_with_untrusted_body.text
    restore_preview_start = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/restore-preview/start",
        json={},
    )
    assert restore_preview_start.status_code == 202, restore_preview_start.text
    restore_preview_task_id = restore_preview_start.json()["task_id"]
    restore_preview_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{restore_preview_task_id}") as websocket:
        while True:
            try:
                restore_preview_events.append(websocket.receive_json())
            except Exception:
                break
    assert restore_preview_events
    assert restore_preview_events[0]["event"] == "task_queued"
    assert any(event["event"] == "scope_ready" for event in restore_preview_events)
    assert restore_preview_events[-1]["event"] == "task_completed"

    restore_preview_result = client.get(
        f"/api/agents/document_agent/{restore_preview_task_id}/result"
    )
    assert restore_preview_result.status_code == 200, restore_preview_result.text
    restore_preview_payload = restore_preview_result.json()["result"]
    assert restore_preview_payload["status"] == "completed"
    restore_preview_context = restore_preview_payload["document_context"]
    assert restore_preview_context["draft_title"] == revision_preview_context["draft_title"]
    assert restore_preview_context["draft_sections"] == revision_preview_context["draft_sections"]
    assert not restore_preview_context["revision_preview"]
    assert not restore_preview_context["revision_suggestions"]
    restored_version = restore_preview_context["draft_version"]
    assert restored_version["root_task_id"] == draft_payload["task_id"]
    assert restored_version["parent_task_id"] == revision_preview_task_id
    assert restored_version["kind"] == "restored_preview"
    assert restore_preview_payload["workflow_run"]["steps"][-1]["action"] == "restore_draft_preview"
    restore_preview_calls = client.get(f"/api/tasks/{restore_preview_task_id}/tool-calls")
    assert restore_preview_calls.status_code == 200, restore_preview_calls.text
    assert not restore_preview_calls.json()["tool_calls"]
    restore_preview_artifacts = client.get(f"/api/tasks/{restore_preview_task_id}/artifacts")
    assert restore_preview_artifacts.status_code == 200, restore_preview_artifacts.text
    assert not restore_preview_artifacts.json()["artifacts"]
    revision_after_restore = client.get(
        f"/api/agents/document_agent/{revision_preview_task_id}/result"
    )
    assert revision_after_restore.status_code == 200, revision_after_restore.text
    assert revision_after_restore.json()["result"]["document_context"]["draft_sections"] == revision_preview_context["draft_sections"]

    # 版本差异是只读查询：初版没有父版本会明确拒绝；修订版本能看见正文变化；恢复预览
    # 与它的直接父快照应完全一致。查询本身不创建任务、不调用模型/Tool，也不会写文件。
    base_version_diff = client.get(
        f"/api/agents/document_agent/{draft_payload['task_id']}/version-diff"
    )
    assert base_version_diff.status_code == 400, base_version_diff.text
    revision_version_diff = client.get(
        f"/api/agents/document_agent/{revision_preview_task_id}/version-diff"
    )
    assert revision_version_diff.status_code == 200, revision_version_diff.text
    revision_version_diff_payload = revision_version_diff.json()
    assert revision_version_diff_payload["parent_task_id"] == section_review_task_id
    assert any(
        section["change_kind"] == "modified"
        for section in revision_version_diff_payload["sections"]
    )
    restored_version_diff = client.get(
        f"/api/agents/document_agent/{restore_preview_task_id}/version-diff"
    )
    assert restored_version_diff.status_code == 200, restored_version_diff.text
    restored_version_diff_payload = restored_version_diff.json()
    assert restored_version_diff_payload["parent_task_id"] == revision_preview_task_id
    assert all(
        section["change_kind"] == "unchanged"
        for section in restored_version_diff_payload["sections"]
    )

    # 模板化交付只接受固定模板 ID 和已核验的草稿快照：服务端只重组现有章节与来源，
    # 不调用模型、Tool、workspace 或写文件。未命中的模板章节必须作为待补充信息公开，
    # 不能用通用文本伪装成交付完成。
    template_preview_with_extra_field = client.post(
        f"/api/agents/document_agent/{draft_payload['task_id']}/template-preview/start",
        json={"template_id": "project_proposal", "untrusted_body": "不应被接受"},
    )
    assert template_preview_with_extra_field.status_code == 422, template_preview_with_extra_field.text
    template_preview_start = client.post(
        f"/api/agents/document_agent/{draft_payload['task_id']}/template-preview/start",
        json={"template_id": "project_proposal"},
    )
    assert template_preview_start.status_code == 202, template_preview_start.text
    template_preview_task_id = template_preview_start.json()["task_id"]
    template_preview_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{template_preview_task_id}") as websocket:
        while True:
            try:
                template_preview_events.append(websocket.receive_json())
            except Exception:
                break
    assert template_preview_events
    assert template_preview_events[0]["event"] == "task_queued"
    assert any(event["event"] == "template_composing" for event in template_preview_events)
    assert template_preview_events[-1]["event"] == "task_completed"
    template_preview_result = client.get(
        f"/api/agents/document_agent/{template_preview_task_id}/result"
    )
    assert template_preview_result.status_code == 200, template_preview_result.text
    template_preview_payload = template_preview_result.json()["result"]
    template_preview_context = template_preview_payload["document_context"]
    assert template_preview_payload["status"] == "completed"
    assert template_preview_payload["mode"] == "mock"
    assert template_preview_context["draft_title"].startswith("项目方案：")
    assert len(template_preview_context["draft_sections"]) == len(draft_context["draft_sections"])
    assert template_preview_context["template_preview"]["template_id"] == "project_proposal"
    assert template_preview_context["template_preview"]["template_name"] == "项目方案"
    assert isinstance(template_preview_context["template_preview"]["missing_sections"], list)
    template_version = template_preview_context["draft_version"]
    assert template_version["root_task_id"] == draft_payload["task_id"]
    assert template_version["parent_task_id"] == draft_payload["task_id"]
    assert template_version["kind"] == "template_preview"
    assert template_preview_payload["workflow_run"]["steps"][-1]["action"] == "create_template_delivery_preview"
    template_preview_calls = client.get(f"/api/tasks/{template_preview_task_id}/tool-calls")
    assert template_preview_calls.status_code == 200, template_preview_calls.text
    assert not template_preview_calls.json()["tool_calls"]
    template_preview_artifacts = client.get(f"/api/tasks/{template_preview_task_id}/artifacts")
    assert template_preview_artifacts.status_code == 200, template_preview_artifacts.text
    assert not template_preview_artifacts.json()["artifacts"]
    source_after_template = client.get(
        f"/api/agents/document_agent/{draft_payload['task_id']}/result"
    )
    assert source_after_template.status_code == 200, source_after_template.text
    assert source_after_template.json()["result"]["document_context"]["draft_sections"] == draft_context["draft_sections"]

    # 用户手动修订不能覆盖当前草稿、更不能借用旧来源直接落盘。接口只接受章节 ID 与新正文，
    # Runtime 重新绑定父版本，创建零模型/零 Tool 的待核验预览；保存端必须在事实核验前拒绝。
    manual_source_section = revision_preview_context["draft_sections"][0]
    manual_revision_with_extra_field = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/draft-sections/manual-revision-preview/start",
        json={
            "section_id": manual_source_section["id"],
            "revised_body": manual_source_section["body"] + "\n\n人工补充：交付前应完成来源核验。",
            "untrusted_path": "C:/outside.md",
        },
    )
    assert manual_revision_with_extra_field.status_code == 422, manual_revision_with_extra_field.text
    unchanged_manual_revision = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/draft-sections/manual-revision-preview/start",
        json={"section_id": manual_source_section["id"], "revised_body": manual_source_section["body"]},
    )
    assert unchanged_manual_revision.status_code == 400, unchanged_manual_revision.text
    manual_revision_start = client.post(
        f"/api/agents/document_agent/{revision_preview_task_id}/draft-sections/manual-revision-preview/start",
        json={
            "section_id": manual_source_section["id"],
            "revised_body": manual_source_section["body"] + "\n\n人工补充：交付前应完成来源核验。",
        },
    )
    assert manual_revision_start.status_code == 202, manual_revision_start.text
    manual_revision_task_id = manual_revision_start.json()["task_id"]
    manual_revision_events: list[dict] = []
    with client.websocket_connect(f"/ws/tasks/{manual_revision_task_id}") as websocket:
        while True:
            try:
                manual_revision_events.append(websocket.receive_json())
            except Exception:
                break
    assert manual_revision_events
    assert manual_revision_events[0]["event"] == "task_queued"
    assert any(event["event"] == "manual_revision_previewing" for event in manual_revision_events)
    assert manual_revision_events[-1]["event"] == "task_completed"
    manual_revision_result = client.get(
        f"/api/agents/document_agent/{manual_revision_task_id}/result"
    )
    assert manual_revision_result.status_code == 200, manual_revision_result.text
    manual_revision_payload = manual_revision_result.json()["result"]
    manual_revision_context = manual_revision_payload["document_context"]
    assert manual_revision_payload["status"] == "completed"
    assert manual_revision_context["draft_verification_state"] == "requires_review"
    assert manual_revision_context["manual_revision_preview"]["original_body"] == manual_source_section["body"]
    assert manual_revision_context["manual_revision_preview"]["revised_body"].endswith("来源核验。")
    manual_revision_version = manual_revision_context["draft_version"]
    assert manual_revision_version["root_task_id"] == draft_payload["task_id"]
    assert manual_revision_version["parent_task_id"] == revision_preview_task_id
    assert manual_revision_version["kind"] == "manual_revision_pending_review"
    assert manual_revision_payload["workflow_run"]["steps"][-1]["action"] == "create_section_manual_revision_preview"
    manual_revision_calls = client.get(f"/api/tasks/{manual_revision_task_id}/tool-calls")
    assert manual_revision_calls.status_code == 200, manual_revision_calls.text
    assert not manual_revision_calls.json()["tool_calls"]
    manual_revision_artifacts = client.get(f"/api/tasks/{manual_revision_task_id}/artifacts")
    assert manual_revision_artifacts.status_code == 200, manual_revision_artifacts.text
    assert not manual_revision_artifacts.json()["artifacts"]
    blocked_manual_save = client.post(
        f"/api/agents/document_agent/{manual_revision_task_id}/save-draft",
        json={"filename": "不应写入的手动修订.md", "confirmed": True},
    )
    assert blocked_manual_save.status_code == 400, blocked_manual_save.text
    manual_version_diff = client.get(
        f"/api/agents/document_agent/{manual_revision_task_id}/version-diff"
    )
    assert manual_version_diff.status_code == 200, manual_version_diff.text
    assert manual_version_diff.json()["parent_task_id"] == revision_preview_task_id
    assert any(section["change_kind"] == "modified" for section in manual_version_diff.json()["sections"])
    revision_after_manual = client.get(
        f"/api/agents/document_agent/{revision_preview_task_id}/result"
    )
    assert revision_after_manual.status_code == 200, revision_after_manual.text
    assert revision_after_manual.json()["result"]["document_context"]["draft_sections"] == revision_preview_context["draft_sections"]

    # 草稿预览不会自动写盘；只有客户端明确 confirmed 后，才允许把同一份已验证章节保存到
    # 受控输出目录。相同任务可以另存为多个不同文件名，但任何同名覆盖都必须失败。
    # 验证使用临时目录，避免离线回归污染项目根 output/。
    draft_output_dir = _VERIFY_DATA_DIR / "document_drafts"
    previous_draft_output_dir = os.environ.get("AGENTFLOW_DOCUMENT_DRAFT_OUTPUT_DIR")
    os.environ["AGENTFLOW_DOCUMENT_DRAFT_OUTPUT_DIR"] = str(draft_output_dir)
    try:
        unconfirmed_save = client.post(
            f"/api/agents/document_agent/{draft_payload['task_id']}/save-draft",
            json={"filename": "星河平台草稿.md", "confirmed": False},
        )
        assert unconfirmed_save.status_code == 409, unconfirmed_save.text
        assert not draft_output_dir.exists()

        saved_draft = client.post(
            f"/api/agents/document_agent/{draft_payload['task_id']}/save-draft",
            json={"filename": "星河平台草稿.md", "confirmed": True},
        )
        assert saved_draft.status_code == 200, saved_draft.text
        saved_draft_payload = saved_draft.json()
        saved_draft_path = draft_output_dir / "星河平台草稿.md"
        assert saved_draft_payload["relative_path"] == "output/document_drafts/星河平台草稿.md"
        assert saved_draft_path.exists()
        saved_draft_text = saved_draft_path.read_text(encoding="utf-8")
        assert "# 星河平台升级" in saved_draft_text
        assert "> 来源：" in saved_draft_text
        assert "版本 草稿初版" in saved_draft_text

        duplicate_save = client.post(
            f"/api/agents/document_agent/{draft_payload['task_id']}/save-draft",
            json={"filename": "星河平台草稿.md", "confirmed": True},
        )
        assert duplicate_save.status_code == 409, duplicate_save.text

        saved_copy = client.post(
            f"/api/agents/document_agent/{draft_payload['task_id']}/save-draft",
            json={"filename": "星河平台草稿-副本.md", "confirmed": True},
        )
        assert saved_copy.status_code == 200, saved_copy.text
        saved_copy_payload = saved_copy.json()
        saved_copy_path = draft_output_dir / "星河平台草稿-副本.md"
        assert saved_copy_payload["artifact_id"] != saved_draft_payload["artifact_id"]
        assert saved_copy_payload["relative_path"] == "output/document_drafts/星河平台草稿-副本.md"
        assert saved_copy_path.read_text(encoding="utf-8") == saved_draft_text

        artifact_response = client.get(f"/api/tasks/{draft_payload['task_id']}/artifacts")
        assert artifact_response.status_code == 200, artifact_response.text
        saved_artifact = next(
            item
            for item in artifact_response.json()["artifacts"]
            if item["artifact_id"] == saved_draft_payload["artifact_id"]
        )
        assert saved_artifact["metadata"]["output_scope"] == "document_drafts"
        assert saved_artifact["metadata"]["document_version_id"] == draft_payload["task_id"]
        assert saved_artifact["metadata"]["document_root_task_id"] == draft_payload["task_id"]
        assert saved_artifact["metadata"]["document_parent_task_id"] == ""
        assert saved_artifact["metadata"]["document_version_kind"] == "base_draft"
        saved_artifacts = [
            item
            for item in artifact_response.json()["artifacts"]
            if item["metadata"].get("output_scope") == "document_drafts"
        ]
        assert {item["artifact_id"] for item in saved_artifacts} == {
            saved_draft_payload["artifact_id"],
            saved_copy_payload["artifact_id"],
        }

        artifact_preview = client.get(
            f"/api/tasks/{draft_payload['task_id']}/artifacts/"
            f"{saved_draft_payload['artifact_id']}/preview"
        )
        assert artifact_preview.status_code == 200, artifact_preview.text
        assert artifact_preview.json()["available"] is True
        assert "星河平台升级" in artifact_preview.json()["text"]

        restored_save = client.post(
            f"/api/agents/document_agent/{restore_preview_task_id}/save-draft",
            json={"filename": "星河平台恢复预览.md", "confirmed": True},
        )
        assert restored_save.status_code == 200, restored_save.text
        restored_save_payload = restored_save.json()
        restored_path = draft_output_dir / "星河平台恢复预览.md"
        assert "版本 恢复预览" in restored_path.read_text(encoding="utf-8")
        restored_artifacts = client.get(f"/api/tasks/{restore_preview_task_id}/artifacts")
        assert restored_artifacts.status_code == 200, restored_artifacts.text
        restored_artifact = next(
            item
            for item in restored_artifacts.json()["artifacts"]
            if item["artifact_id"] == restored_save_payload["artifact_id"]
        )
        assert restored_artifact["metadata"]["document_version_kind"] == "restored_preview"
        assert restored_artifact["metadata"]["document_parent_task_id"] == revision_preview_task_id
    finally:
        if previous_draft_output_dir is None:
            os.environ.pop("AGENTFLOW_DOCUMENT_DRAFT_OUTPUT_DIR", None)
        else:
            os.environ["AGENTFLOW_DOCUMENT_DRAFT_OUTPUT_DIR"] = previous_draft_output_dir

    original_mock_model = document_agent_service._DeterministicDocumentModel
    document_agent_service._DeterministicDocumentModel = _InvalidOutlineModel
    try:
        invalid_draft_run = client.post(
            "/api/agents/document_agent/run",
            json={
                "task_goal": "为这份项目材料生成可审阅的 Markdown 草稿预览。",
                "document_refs": [brief_reference],
                "output_mode": "draft",
            },
        )
    finally:
        document_agent_service._DeterministicDocumentModel = original_mock_model
    assert invalid_draft_run.status_code == 200, invalid_draft_run.text
    invalid_draft_payload = invalid_draft_run.json()
    assert invalid_draft_payload["status"] == "failed"
    assert invalid_draft_payload["stop_reason"] == "model_output_invalid"
    assert not invalid_draft_payload["document_context"]["draft_sections"]

    # 回归用户真实场景：四份接近规划文档大小的材料必须在一次比较内完成。旧版固定 4 轮
    # 且每份仅给 12k 字，必然会在最后一份材料或最终 JSON 前停住；现在每份应完整覆盖到
    # 最后一行，并保留跨文档来源。
    four_document_refs: list[str] = []
    for index in range(1, 5):
        large_material = (
            f"# 第 {index} 份大材料\n\n"
            f"这份材料用于验证多文档完整读取，差异编号为 {index}。\n\n"
            + (f"第 {index} 份材料的普通说明行，用于覆盖真实的长文本阅读范围。\n" * 900)
            + f"\n尾部验收标记：第 {index} 份材料已完整读取。\n"
        )
        response = client.post(
            "/api/workspace/documents",
            json={"filename": f"comparison_large_{index}.md", "content": large_material},
        )
        assert response.status_code == 200, response.text
        four_document_refs.append(response.json()["relative_path"])

    four_document_comparison = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "比较四份材料的共同点和差异，并确认每份材料均被完整读取。",
            "document_refs": four_document_refs,
            "output_mode": "comparison",
        },
    )
    assert four_document_comparison.status_code == 200, four_document_comparison.text
    four_document_payload = four_document_comparison.json()
    assert four_document_payload["status"] == "completed"
    assert len(four_document_payload["document_context"]["sources"]) >= 4
    assert not any("尚有后续内容待继续读取" in warning for warning in four_document_payload["document_context"]["warnings"])
    four_document_tool_calls = client.get(
        f"/api/tasks/{four_document_payload['task_id']}/tool-calls"
    )
    assert four_document_tool_calls.status_code == 200, four_document_tool_calls.text
    four_document_results = [item["result"] for item in four_document_tool_calls.json()["tool_calls"]]
    assert len(four_document_results) == 4
    assert all(result["truncated"] is False for result in four_document_results)
    assert all(
        result["source"]["end_line"] == result["total_lines"]
        for result in four_document_results
    )

    # Qt 文档页走 /start 异步入口；四份合法材料必须被立即受理，不能因为旧版接口契约或
    # 任务轮次估算问题返回 422。这里只验证受理协议，完整分析已经在上面的 /run 回归覆盖。
    four_document_start = client.post(
        "/api/agents/document_agent/start",
        json={
            "task_goal": "异步比较四份材料。",
            "document_refs": four_document_refs,
            "output_mode": "comparison",
        },
    )
    assert four_document_start.status_code == 202, four_document_start.text

    # 分页协议不依赖模型猜测：单页容量不够时，Runtime 必须返回可继续的字符偏移和真实行号。
    paged_import = client.post(
        "/api/workspace/documents",
        json={"filename": "paged_material.md", "content": "标题\n" + ("分页读取内容\n" * 8_000)},
    )
    assert paged_import.status_code == 200, paged_import.text
    paged_name = paged_import.json()["relative_path"]
    first_page = read_workspace_document_excerpt(
        relative_path=paged_name,
        start_char=0,
        max_chars=48_000,
    )
    assert first_page["truncated"] is True
    assert isinstance(first_page["next_start_char"], int)
    second_page = read_workspace_document_excerpt(
        relative_path=paged_name,
        start_char=int(first_page["next_start_char"]),
        max_chars=48_000,
    )
    assert second_page["truncated"] is False
    assert second_page["end_line"] == second_page["total_lines"]

    # 超过直接两页读取范围时不再把原文持续塞入同一模型会话：每个连续分块先压缩，最终
    # 归并只消费带 source_id 的短摘要。离线 mock 也必须完整走这条 Harness 路径。
    compacted_content = (
        "# 超长材料验收\n\n"
        + ("本段包含功能需求、验收条件与约束说明，用于验证连续分块压缩的来源完整性。\n" * 4_000)
        + "\n尾部验收标记：超长材料的最后一段也必须进入可追溯上下文。\n"
    )
    compacted_import = client.post(
        "/api/workspace/documents",
        json={"filename": "compacted_long.md", "content": compacted_content},
    )
    assert compacted_import.status_code == 200, compacted_import.text
    compacted_name = compacted_import.json()["relative_path"]
    compacted_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "提取超长材料中的需求和验收信息，并保留来源。",
            "document_refs": [compacted_name],
            "output_mode": "requirements",
        },
    )
    assert compacted_run.status_code == 200, compacted_run.text
    compacted_payload = compacted_run.json()
    assert compacted_payload["status"] == "completed"
    assert any("连续分段完成上下文压缩" in warning for warning in compacted_payload["document_context"]["warnings"])
    assert compacted_payload["document_context"]["requirements"]
    compacted_workflow = compacted_payload["workflow_run"]
    assert compacted_workflow["limits"]["max_steps"] >= compacted_workflow["metrics"]["step_total"]
    assert compacted_workflow["limits"]["max_tool_calls"] >= compacted_workflow["metrics"]["tool_call_total"]
    assert compacted_workflow["metrics"]["estimated_input_tokens"] > 1_000
    compacted_calls = client.get(f"/api/tasks/{compacted_payload['task_id']}/tool-calls")
    assert compacted_calls.status_code == 200, compacted_calls.text
    compacted_results = [item["result"] for item in compacted_calls.json()["tool_calls"]]
    assert len(compacted_results) >= 3
    assert all(item["context_strategy"] == "chunk_summary" for item in compacted_results)
    assert any(
        item["source"]["end_line"] == item["total_lines"]
        for item in compacted_results
    )

    # 两份材料都超过直接上下文范围时，不能退化为“各自摘要”或只引用最后一个分块。比较结论
    # 必须同时覆盖两个文件的连续分块，并在审计里保留每份材料读到尾部的事实。
    common_long_line = "共同要求：系统必须保留可追溯来源，并展示验收状态。\n"
    long_comparison_contents = {
        "compacted_compare_a.md": (
            "# 材料 A\n\n"
            + common_long_line * 1_700
            + "材料 A 额外约束：不得修改原始文档。\n"
        ),
        "compacted_compare_b.md": (
            "# 材料 B\n\n"
            + common_long_line * 1_700
            + "材料 B 额外要求：需要记录任务历史。\n"
        ),
    }
    compacted_comparison_refs: list[str] = []
    for filename, content in long_comparison_contents.items():
        assert len(content.encode("utf-8")) > 96_000
        imported = client.post("/api/workspace/documents", json={"filename": filename, "content": content})
        assert imported.status_code == 200, imported.text
        compacted_comparison_refs.append(imported.json()["relative_path"])

    compacted_comparison_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "比较两份超长材料的共同要求和差异，并给出跨文件来源。",
            "document_refs": compacted_comparison_refs,
            "output_mode": "comparison",
        },
    )
    assert compacted_comparison_run.status_code == 200, compacted_comparison_run.text
    compacted_comparison_payload = compacted_comparison_run.json()
    assert compacted_comparison_payload["status"] == "completed"
    comparisons = compacted_comparison_payload["document_context"]["comparisons"]
    assert comparisons
    assert {
        source["relative_path"]
        for source in comparisons[0]["source_refs"]
    } == set(compacted_comparison_refs)
    compacted_comparison_calls = client.get(
        f"/api/tasks/{compacted_comparison_payload['task_id']}/tool-calls"
    )
    assert compacted_comparison_calls.status_code == 200, compacted_comparison_calls.text
    comparison_chunk_results = [item["result"] for item in compacted_comparison_calls.json()["tool_calls"]]
    assert len(comparison_chunk_results) >= 4
    assert all(item["context_strategy"] == "chunk_summary" for item in comparison_chunk_results)
    for document_ref in compacted_comparison_refs:
        assert any(
            item["document"] == document_ref
            and item["source"]["end_line"] == item["total_lines"]
            for item in comparison_chunk_results
        )

    # 单个任务最多压缩 12 块。超过预算必须返回可解释的终态，不能后台无限请求模型。
    over_budget_content = "# over-compaction budget\n\n" + ("budget context validation line.\n" * 20_000)
    over_budget_import = client.post(
        "/api/workspace/documents",
        json={"filename": "over_compaction_budget.md", "content": over_budget_content},
    )
    assert over_budget_import.status_code == 200, over_budget_import.text
    over_budget_run = client.post(
        "/api/agents/document_agent/run",
        json={
            "task_goal": "分析这份超预算材料。",
            "document_refs": [over_budget_import.json()["relative_path"]],
            "output_mode": "summary",
        },
    )
    assert over_budget_run.status_code == 200, over_budget_run.text
    assert over_budget_run.json()["status"] == "budget_exhausted"
    assert "12 块" in over_budget_run.json()["reply"]

    # 多文档但未选择时，Agent 必须请求澄清，不能擅自读取其中任意一份。
    ambiguous_document_run = client.post(
        "/api/agents/document_agent/run",
        json={"task_goal": "请总结文档", "output_mode": "summary"},
    )
    assert ambiguous_document_run.status_code == 200, ambiguous_document_run.text
    assert ambiguous_document_run.json()["status"] == "needs_clarification"
    assert ambiguous_document_run.json()["workflow_run"]["status"] == "blocked"

    missing_document_run = client.post(
        "/api/agents/document_agent/run",
        json={"task_goal": "请分析", "document_refs": ["missing.md"]},
    )
    assert missing_document_run.status_code == 200, missing_document_run.text
    assert missing_document_run.json()["status"] == "insufficient_context"

    workspace_search = client.post(
        "/api/workspace/search",
        json={"query": "作业要求", "limit": 5, "context_chars": 20},
    )
    assert workspace_search.status_code == 200, workspace_search.text
    workspace_search_payload = workspace_search.json()
    assert workspace_search_payload["query"] == "作业要求"
    assert workspace_search_payload["total"] >= 1
    assert workspace_search_payload["searched_documents"] >= 2
    assert workspace_search_payload["matched_documents"] == ["assignment.md"]
    assert workspace_search_payload["suggested_read_path"] == "assignment.md"
    first_workspace_match = workspace_search_payload["matches"][0]
    assert first_workspace_match["document_name"] == "assignment.md"
    assert first_workspace_match["line_number"] == 1
    assert "作业要求" in first_workspace_match["preview"]

    empty_workspace_search = client.post(
        "/api/workspace/search",
        json={"query": "   "},
    )
    assert empty_workspace_search.status_code == 400

    unsupported_workspace_import = client.post(
        "/api/workspace/documents",
        json={"filename": "unsafe.exe", "content": "nope"},
    )
    assert unsupported_workspace_import.status_code == 400

    chat = client.post(
        "/api/chat",
        json={"message": "帮我根据这个作业要求生成 Python 代码和 README 报告。"},
    )
    assert chat.status_code == 200, chat.text
    chat_payload = chat.json()
    assert chat_payload["mode"] == "mock"
    task_id = chat_payload["task_id"]
    workflow_plan = chat_payload["workflow_plan"]
    assert workflow_plan["version"] == "1.0"
    assert workflow_plan["schema_version"] == "agentflow.workflow_plan.v1"
    assert workflow_plan["plan_id"].startswith("plan_")
    assert workflow_plan["plan_version"] == 1
    # C0 后总指挥不会因“这个作业”就扫描或猜测本机材料。代码/报告占位能力同样不能
    # 进入客户计划；这里应得到直接答复加材料澄清，而不是一条伪文档委派。
    assert workflow_plan["intent"] == "direct_answer"
    assert workflow_plan["user_goal"] == "帮我根据这个作业要求生成 Python 代码和 README 报告。"
    assert workflow_plan["workflow_name"] == "commander_manager_plan"
    assert workflow_plan["summary"]
    assert workflow_plan["definition_of_done"]
    assert workflow_plan["preference_applied"]["permission_policy"] == "auto_approve"
    assert workflow_plan["preference_applied"]["personality"] == "warm"
    assert workflow_plan["budget_estimate"]["step_count"] == len(workflow_plan["steps"])
    assert workflow_plan["budget_estimate"]["time_level"] == "low"
    assert workflow_plan["workspace_scope"]["write_paths"] == []
    assert workflow_plan["next_action"] == "ask_clarifying_questions"
    assert workflow_plan["clarifying_questions"]
    assert workflow_plan["max_risk_level"] == "low"
    assert workflow_plan["requires_confirmation"] is False
    assert workflow_plan["validation_errors"] == []
    planned_agents = {step["agent"] for step in workflow_plan["steps"]}
    # 文本同时命中代码、报告和文档关键词，但没有显式材料时不能调度任何专业 Agent。
    assert "commander_agent" in planned_agents
    assert "document_agent" not in planned_agents
    assert "code_agent" not in planned_agents
    assert "report_agent" not in planned_agents
    steps_by_agent = {step["agent"]: step for step in workflow_plan["steps"]}
    for step in workflow_plan["steps"]:
        assert step["reason"]
        assert step["expected_output"]
        assert step["tool_name"]
        assert step["success_criteria"]
        assert step["retry_policy"]["max_attempts"] >= 1
    assert steps_by_agent["commander_agent"]["required_permissions"] == []
    assert steps_by_agent["commander_agent"]["execution_mode"] == "planning_only"
    assert steps_by_agent["commander_agent"]["admission_status"] == "ready"

    clarify_chat = client.post("/api/chat", json={"message": "帮我整理这个"})
    assert clarify_chat.status_code == 200, clarify_chat.text
    clarify_plan = clarify_chat.json()["workflow_plan"]
    assert clarify_plan["clarifying_questions"]
    assert clarify_plan["next_action"] == "ask_clarifying_questions"

    commander_workspace_dir = settings.data_dir / "workspaces"
    commander_workspace_dir.mkdir(parents=True, exist_ok=True)
    commander_document = commander_workspace_dir / "commander_input.md"
    commander_document.write_text(
        "# Commander 文档读取验证\n\n请生成一个 Python 示例和 README。",
        encoding="utf-8",
    )
    document_chat = client.post(
        "/api/chat",
        json={
            "message": f"请读取这个文档 {commander_document.name}，然后生成 Python 代码和 README 报告。"
        },
    )
    assert document_chat.status_code == 200, document_chat.text
    document_plan = document_chat.json()["workflow_plan"]
    document_step = next(step for step in document_plan["steps"] if step["agent"] == "document_agent")
    assert document_step["action"] == "analyze_document"
    assert document_step["input"]["document_refs"] == [commander_document.name]
    assert document_step["requires_confirmation"] is False
    assert document_step["tool_name"] == "agent.document_agent.analyze"
    assert [step["action"] for step in document_plan["steps"] if step["agent"] == "document_agent"] == ["analyze_document"]

    document_extract_chat = client.post(
        "/api/chat",
        json={"message": f"请读取这个文档 {commander_document.name}，并归纳要点。"},
    )
    assert document_extract_chat.status_code == 200, document_extract_chat.text
    document_extract_plan = document_extract_chat.json()["workflow_plan"]
    document_extract_steps = [
        step for step in document_extract_plan["steps"] if step["agent"] == "document_agent"
    ]
    assert [step["action"] for step in document_extract_steps] == ["analyze_document"]
    assert document_extract_steps[0]["input"]["document_refs"] == [commander_document.name]

    # Commander 的执行入口必须委派给正式 Document Agent，而不是回退到旧的
    # read_text/extract_requirements 演示节点。父任务保存关联 ID，子任务保存完整 Tool trace。
    commander_document_execute = client.post(
        f"/api/tasks/{document_extract_chat.json()['task_id']}/execute"
    )
    assert commander_document_execute.status_code == 200, commander_document_execute.text
    commander_document_payload = commander_document_execute.json()
    assert commander_document_payload["status"] == "completed"
    commander_document_step = next(
        step
        for step in commander_document_payload["workflow_run"]["steps"]
        if step["agent"] == "document_agent"
    )
    commander_document_result = commander_document_step["output"]["result"]
    assert commander_document_result["agent_status"] == "completed"
    assert commander_document_result["document_context"]["documents"] == [commander_document.name]
    delegated_task_id = commander_document_result["delegated_task_id"]
    delegated_task = client.get(f"/api/tasks/{delegated_task_id}")
    assert delegated_task.status_code == 200, delegated_task.text
    assert delegated_task.json()["status"] == "completed"

    selected_assignment_material = {
        "binding_id": "verify_selected_assignment",
        "kind": "document",
        "ref": "assignment.md",
        "display_name": "assignment.md",
        "origin": "client_selected",
        "usage": "验证总指挥只在明确选取材料内搜索。",
    }
    search_chat = client.post(
        "/api/chat",
        json={
            "message": "请在已选文档中搜索《作业要求》。",
            "materials": [selected_assignment_material],
        },
    )
    assert search_chat.status_code == 200, search_chat.text
    search_plan = search_chat.json()["workflow_plan"]
    search_step = next(step for step in search_plan["steps"] if step["agent"] == "document_agent")
    assert search_step["action"] == "search_text"
    assert search_step["input"]["query"] == "作业要求"
    assert search_step["requires_confirmation"] is False
    assert [step["action"] for step in search_plan["steps"] if step["agent"] == "document_agent"] == ["search_text"]
    assert search_chat.json()["workflow_run"]["status"] == "completed"

    search_extract_chat = client.post(
        "/api/chat",
        json={
            "message": "请搜索《作业要求》并归纳已选文档的要点。",
            "materials": [selected_assignment_material],
        },
    )
    assert search_extract_chat.status_code == 200, search_extract_chat.text
    search_extract_plan = search_extract_chat.json()["workflow_plan"]
    document_actions = [
        step["action"]
        for step in search_extract_plan["steps"]
        if step["agent"] == "document_agent"
    ]
    # 理解任务由文档 Agent 的受控读取/来源链一次完成，不先做无收益的全局搜索。
    assert document_actions == ["analyze_document"]
    document_steps = [
        step
        for step in search_extract_plan["steps"]
        if step["agent"] == "document_agent"
    ]
    assert document_steps[0]["input"]["document_refs"] == ["assignment.md"]
    assert document_steps[0]["execution_mode"] == "execute"
    assert document_steps[0]["admission_status"] == "ready"
    assert search_extract_plan["next_action"] == "execute_after_confirm"

    # Runtime 权限、暂停恢复和产物链路仍需独立回归，但不能由客户入口调度未完成的 Agent。
    # 因此这里显式构造一个仅供底层回归的历史计划，并不代表 Code/Report 已正式上线。
    legacy_runtime_agents = [
        agent.model_copy(update={"runtime_ready": True}) for agent in list_agents()
    ]
    legacy_runtime_plan = WorkflowPlan(
        workflow_name="legacy_runtime_regression_plan",
        description="只验证历史 Runtime 节点的权限、产物与恢复链路，不代表这些 Agent 已上架。",
        summary="验证受控文档上下文、代码草稿和报告草稿的底层 Runtime 链路。",
        max_risk_level="medium",
        requires_confirmation=True,
        preference_applied=WorkflowPlanPreferences(
            permission_policy="auto_approve",
            personality="warm",
        ),
        steps=[
            WorkflowStep(
                id="step_1",
                agent="commander_agent",
                action="analyze_task",
                title="分析底层回归任务",
                input={"message": "验证历史 Runtime 节点。"},
                reason="仅为回归验证建立受控上下文。",
                expected_output="有效的底层执行计划。",
                execution_mode="planning_only",
                admission_status="ready",
                admission_reason="回归计划使用已准入的 Commander 分析动作。",
                verification_scope="WorkflowPlan Validator。",
                recovery_hint="回归失败时查看对应步骤审计。",
            ),
            WorkflowStep(
                id="step_2",
                agent="document_agent",
                action="read_text",
                title="读取受控验证材料",
                depends_on=["step_1"],
                input={"path": commander_document.name},
                reason="为遗留执行器提供最小的受控文档上下文。",
                expected_output="受控文档短上下文。",
                required_permissions=["file_read"],
            ),
            WorkflowStep(
                id="step_3",
                agent="code_agent",
                action="generate_code",
                title="生成历史代码草稿",
                depends_on=["step_2"],
                reason="保留旧代码产物执行器的权限回归覆盖。",
                expected_output="受控 outputs 内的代码草稿。",
                required_permissions=["file_read", "file_write"],
                risk_level="medium",
                requires_confirmation=True,
            ),
            WorkflowStep(
                id="step_4",
                agent="report_agent",
                action="generate_report",
                title="生成历史报告草稿",
                depends_on=["step_2", "step_3"],
                reason="保留旧报告产物执行器的权限回归覆盖。",
                expected_output="受控 outputs 内的 Markdown 报告。",
                required_permissions=["file_read", "file_write"],
                risk_level="medium",
                requires_confirmation=True,
            ),
        ],
    )
    task_id = "verify_legacy_runtime_permissions"
    workflow_run = run_workflow_dry_run(
        task_id=task_id,
        plan=legacy_runtime_plan,
        available_agents=legacy_runtime_agents,
    ).model_dump(mode="json")
    workflow_plan = legacy_runtime_plan.model_dump(mode="json")
    assert workflow_run["task_id"] == task_id
    assert workflow_run["mode"] == "dry_run"
    assert workflow_run["status"] == "completed"
    assert workflow_run["requires_confirmation"] is True
    assert workflow_run["max_risk_level"] == "medium"
    assert workflow_run["validation_errors"] == []
    assert len(workflow_run["steps"]) == len(workflow_plan["steps"])
    assert workflow_run["limits"]["max_retries_per_tool"] == 2
    assert workflow_run["limits"]["tool_timeout_ms"] == 30000
    assert workflow_run["limits"]["task_timeout_ms"] == 120000
    assert workflow_run["metrics"]["step_total"] == len(workflow_plan["steps"])
    assert workflow_run["metrics"]["step_completed"] == len(workflow_run["steps"])
    assert workflow_run["metrics"]["tool_call_total"] == len(workflow_run["steps"])
    assert workflow_run["metrics"]["tool_call_simulated"] == len(workflow_run["steps"])
    assert workflow_run["metrics"]["permission_request_total"] == 2
    assert workflow_run["metrics"]["validation_error_total"] == 0
    assert workflow_run["metrics"]["retry_total"] == 0
    assert workflow_run["metrics"]["budget_exceeded"] is False
    assert workflow_run["metrics"]["estimated_input_tokens"] > 0
    assert workflow_run["metrics"]["started_at"]
    assert workflow_run["metrics"]["finished_at"]
    code_run = next(step for step in workflow_run["steps"] if step["agent"] == "code_agent")
    assert code_run["status"] == "completed"
    assert code_run["output"]["dry_run"] is True
    assert "file_write" in code_run["output"]["required_permissions"]
    assert code_run["output"]["confirmation_required"] is True
    assert "文件写入" in code_run["output"]["permission_summary"]
    assert "需要用户确认" in workflow_run["summary"]
    confirmation_step_ids = {
        step["step_id"] for step in workflow_run["steps"] if step["requires_confirmation"]
    }
    confirmation_step_count = len(confirmation_step_ids)
    assert confirmation_step_count == 2

    task_status = client.get(f"/api/tasks/{task_id}")
    assert task_status.status_code == 200, task_status.text
    task_status_payload = task_status.json()
    assert task_status_payload == workflow_run

    task_plan = client.get(f"/api/tasks/{task_id}/plan")
    assert task_plan.status_code == 200, task_plan.text
    task_plan_payload = task_plan.json()
    assert task_plan_payload["task_id"] == task_id
    assert task_plan_payload["workflow_plan"]["plan_id"] == workflow_plan["plan_id"]
    assert task_plan_payload["workflow_plan"]["definition_of_done"] == workflow_plan["definition_of_done"]
    assert len(task_plan_payload["workflow_plan"]["steps"]) == len(workflow_plan["steps"])

    task_steps = client.get(f"/api/tasks/{task_id}/steps")
    assert task_steps.status_code == 200, task_steps.text
    task_steps_payload = task_steps.json()
    assert task_steps_payload["task_id"] == task_id
    assert task_steps_payload["total"] == len(workflow_run["steps"])
    assert task_steps_payload["steps"] == workflow_run["steps"]

    task_metrics = client.get(f"/api/tasks/{task_id}/metrics")
    assert task_metrics.status_code == 200, task_metrics.text
    task_metrics_payload = task_metrics.json()
    assert task_metrics_payload["task_id"] == task_id
    assert task_metrics_payload["limits"] == workflow_run["limits"]
    assert task_metrics_payload["metrics"] == workflow_run["metrics"]

    task_evaluation = client.get(f"/api/tasks/{task_id}/evaluation")
    assert task_evaluation.status_code == 200, task_evaluation.text
    task_evaluation_payload = task_evaluation.json()
    assert task_evaluation_payload["task_id"] == task_id
    assert task_evaluation_payload["mode"] == "dry_run"
    assert task_evaluation_payload["outcome"] == "dry_run_ready"
    assert task_evaluation_payload["step_success_rate"] == 1.0
    assert task_evaluation_payload["tool_success_rate"] == 1.0
    assert any("dry-run" in warning for warning in task_evaluation_payload["warnings"])

    task_runtime_state = client.get(f"/api/tasks/{task_id}/runtime-state")
    assert task_runtime_state.status_code == 200, task_runtime_state.text
    task_runtime_state_payload = task_runtime_state.json()
    assert task_runtime_state_payload["task_id"] == task_id
    assert task_runtime_state_payload["mode"] == "dry_run"
    assert task_runtime_state_payload["status"] == "completed"
    assert task_runtime_state_payload["terminal"] is True
    assert task_runtime_state_payload["allowed_actions"] == ["retry"]
    assert task_runtime_state_payload["allowed_next_statuses"] == []
    assert "retry" in task_runtime_state_payload["message"]
    assert can_transition("pending", "running") is True
    assert can_transition("running", "waiting_permission") is True
    assert can_transition("completed", "running") is False
    assert is_terminal_status("completed") is True
    assert is_terminal_status("blocked") is False
    assert "running" in allowed_next_statuses("blocked")

    task_artifacts = client.get(f"/api/tasks/{task_id}/artifacts")
    assert task_artifacts.status_code == 200, task_artifacts.text
    task_artifacts_payload = task_artifacts.json()
    assert task_artifacts_payload["task_id"] == task_id
    assert task_artifacts_payload["total"] == len(workflow_run["steps"])
    assert len(task_artifacts_payload["artifacts"]) == len(workflow_run["steps"])
    artifact_step_ids = {
        artifact["step_id"] for artifact in task_artifacts_payload["artifacts"]
    }
    assert artifact_step_ids == {step["step_id"] for step in workflow_run["steps"]}
    assert all(
        artifact["metadata"]["dry_run"] is True
        for artifact in task_artifacts_payload["artifacts"]
    )
    assert all(
        artifact["uri"].startswith("artifact://dry-run/")
        for artifact in task_artifacts_payload["artifacts"]
    )

    dry_run_preview = client.get(
        f"/api/tasks/{task_id}/artifacts/{task_artifacts_payload['artifacts'][0]['artifact_id']}/preview"
    )
    assert dry_run_preview.status_code == 200, dry_run_preview.text
    dry_run_preview_payload = dry_run_preview.json()
    assert dry_run_preview_payload["available"] is False
    assert dry_run_preview_payload["source"] == "dry_run"
    assert "dry-run" in dry_run_preview_payload["reason"]
    assert "output_path" not in dry_run_preview_payload["metadata"]

    task_tool_calls = client.get(f"/api/tasks/{task_id}/tool-calls")
    assert task_tool_calls.status_code == 200, task_tool_calls.text
    task_tool_calls_payload = task_tool_calls.json()
    assert task_tool_calls_payload["task_id"] == task_id
    assert task_tool_calls_payload["total"] == len(workflow_run["steps"])
    assert len(task_tool_calls_payload["tool_calls"]) == len(workflow_run["steps"])
    tool_call_step_ids = {
        tool_call["step_id"] for tool_call in task_tool_calls_payload["tool_calls"]
    }
    assert tool_call_step_ids == {step["step_id"] for step in workflow_run["steps"]}
    for tool_call in task_tool_calls_payload["tool_calls"]:
        assert tool_call["status"] == "simulated"
        assert tool_call["attempt"] == 1
        assert tool_call["max_attempts"] == workflow_run["limits"]["max_retries_per_tool"] + 1
        assert tool_call["timeout_ms"] == workflow_run["limits"]["tool_timeout_ms"]
        assert tool_call["duration_ms"] == 0
        assert tool_call["failure_count"] == 0
        assert tool_call["request"]["dry_run"] is True
        assert tool_call["result"]["simulated"] is True
        assert "." in tool_call["tool_name"]

    task_logs = client.get(f"/api/tasks/{task_id}/logs")
    assert task_logs.status_code == 200, task_logs.text
    task_logs_payload = task_logs.json()
    assert task_logs_payload["task_id"] == task_id
    assert task_logs_payload["total"] == 2 + len(workflow_run["steps"]) * 2 + 1 + confirmation_step_count
    assert task_logs_payload["events"][1]["agent_id"] == "workflow_engine"
    confirmation_events = [
        event for event in task_logs_payload["events"] if event["event"] == "confirmation_required"
    ]
    assert len(confirmation_events) == confirmation_step_count
    assert {event["step_id"] for event in confirmation_events} == confirmation_step_ids
    for event in confirmation_events:
        assert event["level"] == "warning"
        assert "真实执行前需要用户确认" in event["message"]

    task_updates = client.get(f"/api/tasks/{task_id}/updates")
    assert task_updates.status_code == 200, task_updates.text
    task_updates_payload = task_updates.json()
    assert task_updates_payload["task_id"] == task_id
    assert task_updates_payload["total"] > task_logs_payload["total"]
    assert task_updates_payload["updates"][0]["event"] == "connected"
    update_events = {update["event"] for update in task_updates_payload["updates"]}
    update_types = {update["update_type"] for update in task_updates_payload["updates"]}
    assert {"confirmation_required", "step_completed", "artifact_planned", "task_state_snapshot"}.issubset(update_events)
    assert {"permission", "step", "artifact", "state"}.issubset(update_types)
    assert sum(1 for update in task_updates_payload["updates"] if update["event"] == "artifact_planned") == len(workflow_run["steps"])
    permission_updates = [
        update
        for update in task_updates_payload["updates"]
        if update["event"] == "confirmation_required"
    ]
    assert permission_updates
    assert "permissions" in permission_updates[0]["payload"]
    dry_run_state_update = next(
        update
        for update in task_updates_payload["updates"]
        if update["event"] == "task_state_snapshot"
    )
    assert dry_run_state_update["payload"]["evaluation"]["outcome"] == "dry_run_ready"
    assert (
        dry_run_state_update["payload"]["task_retrospective"]["summary"]
        == dry_run_state_update["payload"]["evaluation"]["summary"]
    )

    assert settings.database_path.exists()
    with sqlite3.connect(settings.database_path) as connection:
        persisted_step_count = connection.execute(
            "SELECT COUNT(*) FROM workflow_steps WHERE task_id = ?",
            (task_id,),
        ).fetchone()[0]
        persisted_artifact_count = connection.execute(
            "SELECT COUNT(*) FROM workflow_artifacts WHERE task_id = ?",
            (task_id,),
        ).fetchone()[0]
        persisted_tool_call_count = connection.execute(
            "SELECT COUNT(*) FROM workflow_tool_calls WHERE task_id = ?",
            (task_id,),
        ).fetchone()[0]
    assert persisted_step_count == len(workflow_run["steps"])
    assert persisted_artifact_count == len(workflow_run["steps"])
    assert persisted_tool_call_count == len(workflow_run["steps"])

    permission_list = client.get(f"/api/tasks/{task_id}/permissions")
    assert permission_list.status_code == 200, permission_list.text
    permission_payload = permission_list.json()
    assert permission_payload["task_id"] == task_id
    assert permission_payload["total"] == confirmation_step_count
    assert len(permission_payload["permissions"]) == confirmation_step_count
    permission_items = permission_payload["permissions"]
    assert {
        item["request"]["step_id"] for item in permission_items
    } == confirmation_step_ids
    for item in permission_items:
        assert item["decision"]["decision"] == "pending"
        assert item["request"]["risk_level"] == "medium"
        assert "file_write" in item["request"]["permissions"]
        assert item["request"]["details"]["dry_run"] is True

    pending_permission_list = client.get(
        f"/api/tasks/{task_id}/permissions",
        params={"decision": "pending"},
    )
    assert pending_permission_list.status_code == 200, pending_permission_list.text
    assert pending_permission_list.json()["total"] == confirmation_step_count

    first_permission_request_id = permission_items[0]["request"]["request_id"]
    approve_permission = client.post(
        f"/api/tasks/{task_id}/permissions/{first_permission_request_id}/decision",
        json={
            "decision": "approved",
            "decided_by": "verify_backend",
            "note": "验证权限确认审计记录。",
        },
    )
    assert approve_permission.status_code == 200, approve_permission.text
    approved_payload = approve_permission.json()
    assert approved_payload["decision"]["decision"] == "approved"
    assert approved_payload["decision"]["decided_by"] == "verify_backend"
    assert approved_payload["decision"]["decided_at"]

    approved_permission_list = client.get(
        f"/api/tasks/{task_id}/permissions",
        params={"decision": "approved"},
    )
    assert approved_permission_list.status_code == 200, approved_permission_list.text
    assert approved_permission_list.json()["total"] == 1

    task_list = client.get("/api/tasks?limit=10&offset=0")
    assert task_list.status_code == 200, task_list.text
    task_list_payload = task_list.json()
    assert task_list_payload["total"] >= 1
    assert task_list_payload["limit"] == 10
    assert task_list_payload["offset"] == 0
    listed_task = next(
        (task for task in task_list_payload["tasks"] if task["task_id"] == task_id),
        None,
    )
    assert listed_task is not None
    assert listed_task["mode"] == "dry_run"
    assert listed_task["status"] == workflow_run["status"]
    assert listed_task["summary"] == workflow_run["summary"]
    assert listed_task["step_count"] == len(workflow_run["steps"])

    filtered_task_list = client.get(
        "/api/tasks",
        params={
            "limit": 20,
            "offset": 0,
            "status": "completed",
            "mode": "dry_run",
            "max_risk_level": "medium",
            "requires_confirmation": "true",
        },
    )
    assert filtered_task_list.status_code == 200, filtered_task_list.text
    filtered_payload = filtered_task_list.json()
    assert filtered_payload["total"] >= 1
    filtered_tasks = filtered_payload["tasks"]
    assert filtered_tasks
    assert any(task["task_id"] == task_id for task in filtered_tasks)
    for task in filtered_tasks:
        assert task["mode"] == "dry_run"
        assert task["status"] == "completed"
        assert task["max_risk_level"] == "medium"
        assert task["requires_confirmation"] is True

    empty_page = client.get("/api/tasks?limit=1&offset=1000000")
    assert empty_page.status_code == 200, empty_page.text
    assert empty_page.json()["tasks"] == []

    # 模拟服务重启后内存缓存丢失，确认任务状态和日志可以从 SQLite 恢复。
    clear_dry_run_memory_cache()
    persisted_task_status = client.get(f"/api/tasks/{task_id}")
    assert persisted_task_status.status_code == 200, persisted_task_status.text
    assert persisted_task_status.json() == workflow_run

    persisted_task_plan = client.get(f"/api/tasks/{task_id}/plan")
    assert persisted_task_plan.status_code == 200, persisted_task_plan.text
    assert persisted_task_plan.json() == task_plan_payload

    persisted_task_steps = client.get(f"/api/tasks/{task_id}/steps")
    assert persisted_task_steps.status_code == 200, persisted_task_steps.text
    assert persisted_task_steps.json() == task_steps_payload

    persisted_task_metrics = client.get(f"/api/tasks/{task_id}/metrics")
    assert persisted_task_metrics.status_code == 200, persisted_task_metrics.text
    assert persisted_task_metrics.json() == task_metrics_payload

    persisted_task_runtime_state = client.get(f"/api/tasks/{task_id}/runtime-state")
    assert persisted_task_runtime_state.status_code == 200, persisted_task_runtime_state.text
    assert persisted_task_runtime_state.json() == task_runtime_state_payload

    persisted_task_artifacts = client.get(f"/api/tasks/{task_id}/artifacts")
    assert persisted_task_artifacts.status_code == 200, persisted_task_artifacts.text
    assert persisted_task_artifacts.json() == task_artifacts_payload

    persisted_task_tool_calls = client.get(f"/api/tasks/{task_id}/tool-calls")
    assert persisted_task_tool_calls.status_code == 200, persisted_task_tool_calls.text
    assert persisted_task_tool_calls.json() == task_tool_calls_payload

    persisted_task_logs = client.get(f"/api/tasks/{task_id}/logs")
    assert persisted_task_logs.status_code == 200, persisted_task_logs.text
    assert persisted_task_logs.json()["total"] == task_logs_payload["total"]

    persisted_task_updates = client.get(f"/api/tasks/{task_id}/updates")
    assert persisted_task_updates.status_code == 200, persisted_task_updates.text
    assert persisted_task_updates.json()["total"] == task_updates_payload["total"]

    persisted_permissions = client.get(f"/api/tasks/{task_id}/permissions")
    assert persisted_permissions.status_code == 200, persisted_permissions.text
    assert persisted_permissions.json()["total"] == confirmation_step_count

    missing_task = client.get("/api/tasks/not_exist")
    assert missing_task.status_code == 404

    missing_plan = client.get("/api/tasks/not_exist/plan")
    assert missing_plan.status_code == 404

    missing_steps = client.get("/api/tasks/not_exist/steps")
    assert missing_steps.status_code == 404

    missing_metrics = client.get("/api/tasks/not_exist/metrics")
    assert missing_metrics.status_code == 404

    missing_evaluation = client.get("/api/tasks/not_exist/evaluation")
    assert missing_evaluation.status_code == 404

    missing_runtime_state = client.get("/api/tasks/not_exist/runtime-state")
    assert missing_runtime_state.status_code == 404

    missing_artifacts = client.get("/api/tasks/not_exist/artifacts")
    assert missing_artifacts.status_code == 404

    missing_artifact_preview = client.get(f"/api/tasks/{task_id}/artifacts/not_exist/preview")
    assert missing_artifact_preview.status_code == 404

    missing_tool_calls = client.get("/api/tasks/not_exist/tool-calls")
    assert missing_tool_calls.status_code == 404

    missing_updates = client.get("/api/tasks/not_exist/updates")
    assert missing_updates.status_code == 404

    missing_permissions = client.get("/api/tasks/not_exist/permissions")
    assert missing_permissions.status_code == 404

    missing_permission_decision = client.post(
        f"/api/tasks/{task_id}/permissions/not_exist/decision",
        json={"decision": "denied"},
    )
    assert missing_permission_decision.status_code == 404

    cancel = client.post(f"/api/tasks/{task_id}/cancel")
    assert cancel.status_code == 200, cancel.text
    cancel_payload = cancel.json()
    assert cancel_payload["action"] == "cancel"
    assert cancel_payload["accepted"] is False
    assert cancel_payload["status"] == "completed"
    assert cancel_payload["workflow_run"]["task_id"] == task_id

    retry = client.post(f"/api/tasks/{task_id}/retry")
    assert retry.status_code == 200, retry.text
    retry_payload = retry.json()
    assert retry_payload["action"] == "retry"
    assert retry_payload["accepted"] is True
    assert retry_payload["new_task_id"]
    assert retry_payload["workflow_run"]["task_id"] == retry_payload["new_task_id"]
    assert retry_payload["workflow_run"]["mode"] == "dry_run"
    assert retry_payload["workflow_run"]["status"] == "completed"
    assert retry_payload["workflow_run"]["metrics"]["tool_call_total"] == len(
        retry_payload["workflow_run"]["steps"]
    )

    retry_status = client.get(f"/api/tasks/{retry_payload['new_task_id']}")
    assert retry_status.status_code == 200, retry_status.text
    assert retry_status.json() == retry_payload["workflow_run"]

    retry_metrics = client.get(f"/api/tasks/{retry_payload['new_task_id']}/metrics")
    assert retry_metrics.status_code == 200, retry_metrics.text
    assert retry_metrics.json()["metrics"]["retry_total"] == 0

    retry_runtime_state = client.get(
        f"/api/tasks/{retry_payload['new_task_id']}/runtime-state"
    )
    assert retry_runtime_state.status_code == 200, retry_runtime_state.text
    assert retry_runtime_state.json()["allowed_actions"] == ["retry"]

    retry_logs = client.get(f"/api/tasks/{retry_payload['new_task_id']}/logs")
    assert retry_logs.status_code == 200, retry_logs.text
    assert retry_logs.json()["events"][-1]["event"] == "task_completed"

    retry_artifacts = client.get(f"/api/tasks/{retry_payload['new_task_id']}/artifacts")
    assert retry_artifacts.status_code == 200, retry_artifacts.text
    assert retry_artifacts.json()["total"] == len(retry_payload["workflow_run"]["steps"])

    retry_tool_calls = client.get(f"/api/tasks/{retry_payload['new_task_id']}/tool-calls")
    assert retry_tool_calls.status_code == 200, retry_tool_calls.text
    assert retry_tool_calls.json()["total"] == len(retry_payload["workflow_run"]["steps"])

    retry_permissions = client.get(
        f"/api/tasks/{retry_payload['new_task_id']}/permissions"
    )
    assert retry_permissions.status_code == 200, retry_permissions.text
    assert retry_permissions.json()["total"] == confirmation_step_count

    auto_runtime_execute = client.post(f"/api/tasks/{task_id}/execute")
    assert auto_runtime_execute.status_code == 200, auto_runtime_execute.text
    auto_runtime_payload = auto_runtime_execute.json()
    assert auto_runtime_payload["status"] == "completed", auto_runtime_payload
    auto_runtime_task_id = auto_runtime_payload["runtime_task_id"]
    auto_runtime_permissions = client.get(
        f"/api/tasks/{auto_runtime_task_id}/permissions"
    )
    assert auto_runtime_permissions.status_code == 200, auto_runtime_permissions.text
    auto_permission_items = auto_runtime_permissions.json()["permissions"]
    assert len(auto_permission_items) == confirmation_step_count
    assert all(
        item["decision"]["decision"] == "approved"
        and item["decision"]["decided_by"] == "platform_policy:auto_approve"
        and item["request"]["details"]["policy_action"] == "allow"
        for item in auto_permission_items
    )
    auto_runtime_logs = client.get(f"/api/tasks/{auto_runtime_task_id}/logs")
    assert auto_runtime_logs.status_code == 200, auto_runtime_logs.text
    assert sum(
        event["event"] == "permission_auto_approved"
        for event in auto_runtime_logs.json()["events"]
    ) == confirmation_step_count

    reset_runtime_preferences = client.put(
        "/api/settings/runtime-preferences",
        json={"permission_policy": "smart_confirm", "personality": "warm"},
    )
    assert reset_runtime_preferences.status_code == 200, reset_runtime_preferences.text
    manual_runtime_plan = legacy_runtime_plan.model_copy(
        update={
            "preference_applied": WorkflowPlanPreferences(
                permission_policy="smart_confirm",
                personality="warm",
            )
        }
    )
    task_id = "verify_legacy_runtime_manual_permission"
    run_workflow_dry_run(
        task_id=task_id,
        plan=manual_runtime_plan,
        available_agents=legacy_runtime_agents,
    )
    assert manual_runtime_plan.preference_applied.permission_policy == "smart_confirm"

    runtime_execute = client.post(f"/api/tasks/{task_id}/execute")
    assert runtime_execute.status_code == 200, runtime_execute.text
    runtime_execute_payload = runtime_execute.json()
    assert runtime_execute_payload["source_task_id"] == task_id
    assert runtime_execute_payload["runtime_task_id"] != task_id
    assert runtime_execute_payload["accepted"] is True
    assert runtime_execute_payload["status"] == "waiting_permission"
    assert "等待用户权限确认" in runtime_execute_payload["message"]
    runtime_task_id = runtime_execute_payload["runtime_task_id"]
    runtime_run = runtime_execute_payload["workflow_run"]
    assert runtime_run["task_id"] == runtime_task_id
    assert runtime_run["mode"] == "runtime"
    assert runtime_run["status"] == "waiting_permission"
    assert runtime_run["metrics"]["tool_call_simulated"] == 0
    assert runtime_run["metrics"]["permission_request_total"] == confirmation_step_count
    runtime_steps_by_agent = {step["agent"]: step for step in runtime_run["steps"]}
    assert runtime_steps_by_agent["commander_agent"]["status"] == "completed"
    assert runtime_steps_by_agent["document_agent"]["status"] == "completed"
    assert runtime_steps_by_agent["code_agent"]["status"] == "waiting_permission"
    assert runtime_steps_by_agent["report_agent"]["status"] == "pending"
    completed_before_permission_ids = {
        step["step_id"]
        for step in runtime_run["steps"]
        if step["status"] == "completed"
    }

    runtime_state = client.get(f"/api/tasks/{runtime_task_id}/runtime-state")
    assert runtime_state.status_code == 200, runtime_state.text
    runtime_state_payload = runtime_state.json()
    assert runtime_state_payload["mode"] == "runtime"
    assert runtime_state_payload["status"] == "waiting_permission"
    assert runtime_state_payload["terminal"] is False
    assert runtime_state_payload["allowed_actions"] == ["pause", "cancel"]
    assert "running" in runtime_state_payload["allowed_next_statuses"]

    runtime_permissions = client.get(f"/api/tasks/{runtime_task_id}/permissions")
    assert runtime_permissions.status_code == 200, runtime_permissions.text
    runtime_permissions_payload = runtime_permissions.json()
    assert runtime_permissions_payload["total"] == confirmation_step_count
    assert all(
        item["decision"]["decision"] == "pending"
        for item in runtime_permissions_payload["permissions"]
    )
    assert all(
        item["request"]["details"]["runtime"] is True
        for item in runtime_permissions_payload["permissions"]
    )

    runtime_tool_calls = client.get(f"/api/tasks/{runtime_task_id}/tool-calls")
    assert runtime_tool_calls.status_code == 200, runtime_tool_calls.text
    runtime_tool_call_statuses = {
        call["status"] for call in runtime_tool_calls.json()["tool_calls"]
    }
    assert {"completed", "pending_permission", "skipped"}.issubset(
        runtime_tool_call_statuses
    )

    runtime_artifacts = client.get(f"/api/tasks/{runtime_task_id}/artifacts")
    assert runtime_artifacts.status_code == 200, runtime_artifacts.text
    assert runtime_artifacts.json()["total"] == 2

    cancel_runtime_execute = client.post(f"/api/tasks/{task_id}/execute")
    assert cancel_runtime_execute.status_code == 200, cancel_runtime_execute.text
    cancel_runtime_task_id = cancel_runtime_execute.json()["runtime_task_id"]
    cancel_runtime = client.post(f"/api/tasks/{cancel_runtime_task_id}/cancel")
    assert cancel_runtime.status_code == 200, cancel_runtime.text
    cancel_runtime_payload = cancel_runtime.json()
    assert cancel_runtime_payload["accepted"] is True
    assert cancel_runtime_payload["status"] == "cancelled"
    assert cancel_runtime_payload["workflow_run"]["status"] == "cancelled"
    cancel_runtime_state = client.get(f"/api/tasks/{cancel_runtime_task_id}/runtime-state")
    assert cancel_runtime_state.status_code == 200, cancel_runtime_state.text
    assert cancel_runtime_state.json()["terminal"] is True
    assert cancel_runtime_state.json()["allowed_actions"] == ["retry"]

    blocked_runtime_execute = client.post(f"/api/tasks/{task_id}/execute")
    assert blocked_runtime_execute.status_code == 200, blocked_runtime_execute.text
    blocked_runtime_task_id = blocked_runtime_execute.json()["runtime_task_id"]
    blocked_permissions = client.get(f"/api/tasks/{blocked_runtime_task_id}/permissions")
    assert blocked_permissions.status_code == 200, blocked_permissions.text
    blocked_permission_items = blocked_permissions.json()["permissions"]
    assert blocked_permission_items
    deny_runtime_permission = client.post(
        f"/api/tasks/{blocked_runtime_task_id}/permissions/{blocked_permission_items[0]['request']['request_id']}/decision",
        json={
            "decision": "denied",
            "decided_by": "verify_backend",
            "note": "验证拒绝权限后 Runtime 进入 blocked。",
        },
    )
    assert deny_runtime_permission.status_code == 200, deny_runtime_permission.text
    assert deny_runtime_permission.json()["decision"]["decision"] == "denied"

    blocked_runtime_resume = client.post(f"/api/tasks/{blocked_runtime_task_id}/execute")
    assert blocked_runtime_resume.status_code == 200, blocked_runtime_resume.text
    blocked_runtime_payload = blocked_runtime_resume.json()
    assert blocked_runtime_payload["accepted"] is True
    assert blocked_runtime_payload["status"] == "blocked"
    assert blocked_runtime_payload["workflow_run"]["status"] == "blocked"
    blocked_steps_by_status = {
        step["step_id"]: step["status"]
        for step in blocked_runtime_payload["workflow_run"]["steps"]
    }
    assert "blocked" in blocked_steps_by_status.values()
    blocked_runtime_state = client.get(f"/api/tasks/{blocked_runtime_task_id}/runtime-state")
    assert blocked_runtime_state.status_code == 200, blocked_runtime_state.text
    blocked_runtime_state_payload = blocked_runtime_state.json()
    assert blocked_runtime_state_payload["status"] == "blocked"
    assert blocked_runtime_state_payload["terminal"] is False
    assert blocked_runtime_state_payload["allowed_actions"] == ["retry"]
    assert "running" in blocked_runtime_state_payload["allowed_next_statuses"]
    blocked_tool_calls = client.get(f"/api/tasks/{blocked_runtime_task_id}/tool-calls")
    assert blocked_tool_calls.status_code == 200, blocked_tool_calls.text
    assert any(
        call["status"] == "blocked" and call["error"]
        for call in blocked_tool_calls.json()["tool_calls"]
    )
    blocked_logs = client.get(f"/api/tasks/{blocked_runtime_task_id}/logs")
    assert blocked_logs.status_code == 200, blocked_logs.text
    assert any(
        event["event"] == "permission_denied"
        for event in blocked_logs.json()["events"]
    )

    for item in runtime_permissions_payload["permissions"]:
        approve_runtime_permission = client.post(
            f"/api/tasks/{runtime_task_id}/permissions/{item['request']['request_id']}/decision",
            json={
                "decision": "approved",
                "decided_by": "verify_backend",
                "note": "验证 Runtime 权限批准后继续执行。",
            },
        )
        assert approve_runtime_permission.status_code == 200, approve_runtime_permission.text
        assert approve_runtime_permission.json()["decision"]["decision"] == "approved"

    runtime_resume = client.post(f"/api/tasks/{runtime_task_id}/execute")
    assert runtime_resume.status_code == 200, runtime_resume.text
    runtime_resume_payload = runtime_resume.json()
    assert runtime_resume_payload["runtime_task_id"] == runtime_task_id
    assert runtime_resume_payload["accepted"] is True
    assert runtime_resume_payload["status"] == "completed"
    completed_runtime_run = runtime_resume_payload["workflow_run"]
    assert completed_runtime_run["status"] == "completed"
    assert all(step["status"] == "completed" for step in completed_runtime_run["steps"])
    assert completed_runtime_run["metrics"]["step_completed"] == len(
        completed_runtime_run["steps"]
    )
    assert completed_runtime_run["metrics"]["tool_call_simulated"] == 0

    # interruption 恢复必须沿用同一条执行链：已完成步骤不会再次启动，日志会明确留下恢复点。
    completed_runtime_logs = client.get(f"/api/tasks/{runtime_task_id}/logs")
    assert completed_runtime_logs.status_code == 200, completed_runtime_logs.text
    completed_runtime_events = completed_runtime_logs.json()["events"]
    assert any(event["event"] == "task_resumed" for event in completed_runtime_events)
    for completed_step_id in completed_before_permission_ids:
        assert sum(
            event["event"] == "step_started" and event["step_id"] == completed_step_id
            for event in completed_runtime_events
        ) == 1

    completed_runtime_tool_calls = client.get(
        f"/api/tasks/{runtime_task_id}/tool-calls"
    )
    assert completed_runtime_tool_calls.status_code == 200, completed_runtime_tool_calls.text
    completed_tool_calls_payload = completed_runtime_tool_calls.json()["tool_calls"]
    assert all(call["status"] == "completed" for call in completed_tool_calls_payload)
    assert {
        call["tool_name"] for call in completed_tool_calls_payload
    }.issuperset(
        {
            "planner.analyze_task",
            "document.read_text",
            "code.generate_code",
            "report.compose_markdown",
        }
    )

    completed_runtime_artifacts = client.get(f"/api/tasks/{runtime_task_id}/artifacts")
    assert completed_runtime_artifacts.status_code == 200, completed_runtime_artifacts.text
    completed_artifacts_payload = completed_runtime_artifacts.json()
    assert completed_artifacts_payload["total"] == len(completed_runtime_run["steps"])
    code_artifact = next(
        artifact
        for artifact in completed_artifacts_payload["artifacts"]
        if artifact["kind"] == "code"
    )
    report_artifact = next(
        artifact
        for artifact in completed_artifacts_payload["artifacts"]
        if artifact["kind"] == "report"
    )
    # artifact 列表面向桌面端，不能再把后端绝对路径当作产品协议返回。离线验证本身已知
    # 自己创建的受控 Runtime 根，因此由固定根与 artifact 名称定位文件，并同时锁定脱敏边界。
    code_relative_path = Path(code_artifact["metadata"]["relative_path"])
    report_relative_path = Path(report_artifact["metadata"]["relative_path"])
    assert not code_relative_path.is_absolute()
    assert not report_relative_path.is_absolute()
    code_output_path = (settings.data_dir / code_relative_path).resolve()
    report_output_path = (settings.data_dir / report_relative_path).resolve()
    assert code_output_path.is_relative_to((settings.data_dir / "outputs").resolve())
    assert report_output_path.is_relative_to((settings.data_dir / "outputs").resolve())
    assert code_artifact["metadata"]["output_path"] == "<hidden>"
    assert report_artifact["metadata"]["output_path"] == "<hidden>"
    assert code_output_path.exists()
    assert report_output_path.exists()
    assert "AgentFlow 代码草稿" in code_output_path.read_text(encoding="utf-8")
    assert "AgentFlow Runtime 报告草稿" in report_output_path.read_text(encoding="utf-8")

    code_preview = client.get(
        f"/api/tasks/{runtime_task_id}/artifacts/{code_artifact['artifact_id']}/preview"
    )
    assert code_preview.status_code == 200, code_preview.text
    code_preview_payload = code_preview.json()
    assert code_preview_payload["available"] is True
    assert code_preview_payload["source"] == "runtime_output"
    assert code_preview_payload["bytes_read"] > 0
    assert code_preview_payload["truncated"] is False
    assert "AgentFlow 代码草稿" in code_preview_payload["text"]
    assert code_preview_payload["metadata"]["output_path"] == "<hidden>"
    assert str(settings.data_dir) not in json.dumps(code_preview_payload, ensure_ascii=False)

    limited_code_preview = client.get(
        f"/api/tasks/{runtime_task_id}/artifacts/{code_artifact['artifact_id']}/preview",
        params={"max_bytes": 16},
    )
    assert limited_code_preview.status_code == 200, limited_code_preview.text
    limited_code_preview_payload = limited_code_preview.json()
    assert limited_code_preview_payload["available"] is True
    assert limited_code_preview_payload["bytes_read"] == 16
    assert limited_code_preview_payload["truncated"] is True

    completed_runtime_evaluation = client.get(f"/api/tasks/{runtime_task_id}/evaluation")
    assert completed_runtime_evaluation.status_code == 200, completed_runtime_evaluation.text
    completed_runtime_evaluation_payload = completed_runtime_evaluation.json()
    assert completed_runtime_evaluation_payload["mode"] == "runtime"
    assert completed_runtime_evaluation_payload["outcome"] == "completed"
    assert completed_runtime_evaluation_payload["step_success_rate"] == 1.0
    assert completed_runtime_evaluation_payload["tool_success_rate"] == 1.0
    assert completed_runtime_evaluation_payload["failed_tool_calls"] == 0

    completed_runtime_updates = client.get(f"/api/tasks/{runtime_task_id}/updates")
    assert completed_runtime_updates.status_code == 200, completed_runtime_updates.text
    completed_runtime_updates_payload = completed_runtime_updates.json()
    runtime_update_events = {
        update["event"] for update in completed_runtime_updates_payload["updates"]
    }
    assert "artifact_created" in runtime_update_events
    assert "task_state_snapshot" in runtime_update_events
    runtime_state_update = completed_runtime_updates_payload["updates"][-1]
    assert runtime_state_update["event"] == "task_state_snapshot"
    assert runtime_state_update["status"] == "completed"
    assert runtime_state_update["payload"]["artifact_total"] == len(completed_artifacts_payload["artifacts"])
    assert runtime_state_update["payload"]["evaluation"]["outcome"] == "completed"
    assert runtime_state_update["payload"]["task_retrospective"]["outcome"] == "completed"
    assert runtime_state_update["payload"]["task_retrospective"]["facts"]["artifact_total"] == len(completed_artifacts_payload["artifacts"])

    clear_dry_run_memory_cache()
    persisted_runtime_status = client.get(f"/api/tasks/{runtime_task_id}")
    assert persisted_runtime_status.status_code == 200, persisted_runtime_status.text
    assert persisted_runtime_status.json()["status"] == "completed"
    persisted_runtime_artifacts = client.get(f"/api/tasks/{runtime_task_id}/artifacts")
    assert persisted_runtime_artifacts.status_code == 200, persisted_runtime_artifacts.text
    assert persisted_runtime_artifacts.json()["total"] == len(completed_runtime_run["steps"])
    persisted_runtime_updates = client.get(f"/api/tasks/{runtime_task_id}/updates")
    assert persisted_runtime_updates.status_code == 200, persisted_runtime_updates.text
    assert persisted_runtime_updates.json()["total"] == completed_runtime_updates_payload["total"]

    missing_execute = client.post("/api/tasks/not_exist/execute")
    assert missing_execute.status_code == 404

    missing_retry = client.post("/api/tasks/not_exist/retry")
    assert missing_retry.status_code == 404

    workspace_dir = settings.data_dir / "workspaces"
    workspace_dir.mkdir(parents=True, exist_ok=True)
    sample_document = workspace_dir / "runtime_read_success.md"
    sample_document.write_text(
        "# AgentFlow 受控文档\n\n这是 document.read_text 的成功路径验证。",
        encoding="utf-8",
    )
    read_text_plan = WorkflowPlan(
        workflow_name="runtime_read_text_plan",
        description="验证 document.read_text 可以读取受控工作区文本。",
        summary="读取受控 markdown 文件。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="document_agent",
                action="read_text",
                title="读取受控文档",
                input={"path": sample_document.name},
                reason="验证 Runtime 文件读取边界的成功路径。",
                expected_output="返回文档预览。",
            )
        ],
    )
    read_text_dry_run = run_workflow_dry_run(
        task_id="verify_runtime_read_text",
        plan=read_text_plan,
        available_agents=list_agents(),
    )
    assert read_text_dry_run.status == "completed"
    read_text_execute = client.post(f"/api/tasks/{read_text_dry_run.task_id}/execute")
    assert read_text_execute.status_code == 200, read_text_execute.text
    read_text_payload = read_text_execute.json()
    assert read_text_payload["status"] == "completed"
    read_text_run = read_text_payload["workflow_run"]
    assert read_text_run["steps"][0]["status"] == "completed"
    read_text_result = read_text_run["steps"][0]["output"]["result"]
    assert "AgentFlow 受控文档" in read_text_result["preview"]
    assert read_text_result["relative_path"] == sample_document.name
    assert read_text_result["bytes"] > 0
    read_text_tool_calls = client.get(
        f"/api/tasks/{read_text_payload['runtime_task_id']}/tool-calls"
    )
    assert read_text_tool_calls.status_code == 200, read_text_tool_calls.text
    read_text_tool_call = read_text_tool_calls.json()["tool_calls"][0]
    assert read_text_tool_call["status"] == "completed"
    assert read_text_tool_call["tool_name"] == "document.read_text"
    read_text_artifacts = client.get(
        f"/api/tasks/{read_text_payload['runtime_task_id']}/artifacts"
    )
    assert read_text_artifacts.status_code == 200, read_text_artifacts.text
    assert read_text_artifacts.json()["total"] == 1

    read_then_extract_plan = WorkflowPlan(
        workflow_name="runtime_read_then_extract_plan",
        description="验证 document.extract_requirements 可以读取前置 read_text 上下文。",
        summary="读取受控文档后归纳要点。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="document_agent",
                action="read_text",
                title="读取受控文档",
                input={"path": sample_document.name},
                reason="先读取用户明确指定的受控文档。",
                expected_output="返回文档预览。",
            ),
            WorkflowStep(
                id="step_2",
                agent="document_agent",
                action="extract_requirements",
                title="归纳文档要点",
                depends_on=["step_1"],
                reason="基于前置读取结果形成结构化要求。",
                expected_output="返回可供后续 Agent 使用的文档上下文。",
            ),
        ],
    )
    read_then_extract_dry_run = run_workflow_dry_run(
        task_id="verify_runtime_read_then_extract",
        plan=read_then_extract_plan,
        available_agents=list_agents(),
    )
    assert read_then_extract_dry_run.status == "completed"
    read_then_extract_execute = client.post(
        f"/api/tasks/{read_then_extract_dry_run.task_id}/execute"
    )
    assert read_then_extract_execute.status_code == 200, read_then_extract_execute.text
    read_then_extract_payload = read_then_extract_execute.json()
    assert read_then_extract_payload["status"] == "completed"
    read_then_extract_run = read_then_extract_payload["workflow_run"]
    assert read_then_extract_run["steps"][1]["status"] == "completed"
    read_then_extract_context = read_then_extract_run["steps"][1]["output"]["result"]["context"]
    assert read_then_extract_context["source_steps"] == ["step_1"]
    assert read_then_extract_context["read_preview_total"] == 1
    assert read_then_extract_context["search_match_total"] == 0
    assert read_then_extract_context["read_previews"][0]["relative_path"] == sample_document.name
    assert "AgentFlow 受控文档" in read_then_extract_context["read_previews"][0]["preview"]
    assert str(settings.data_dir) not in json.dumps(read_then_extract_context, ensure_ascii=False)

    search_text_plan = WorkflowPlan(
        workflow_name="runtime_search_text_plan",
        description="验证 document.search_text 可以搜索受控 workspace 文档。",
        summary="搜索受控 markdown 文件。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="document_agent",
                action="search_text",
                title="搜索受控文档",
                input={"query": "成功路径", "limit": 10},
                reason="验证 Runtime 文档精确搜索能力。",
                expected_output="返回命中文档、行号和短上下文。",
            )
        ],
    )
    search_text_dry_run = run_workflow_dry_run(
        task_id="verify_runtime_search_text",
        plan=search_text_plan,
        available_agents=list_agents(),
    )
    assert search_text_dry_run.status == "completed"
    search_text_execute = client.post(f"/api/tasks/{search_text_dry_run.task_id}/execute")
    assert search_text_execute.status_code == 200, search_text_execute.text
    search_text_payload = search_text_execute.json()
    assert search_text_payload["status"] == "completed"
    search_text_run = search_text_payload["workflow_run"]
    assert search_text_run["steps"][0]["status"] == "completed"
    search_text_result = search_text_run["steps"][0]["output"]["result"]
    assert search_text_result["total"] >= 1
    assert search_text_result["matches"][0]["document_name"] == sample_document.name
    assert "成功路径" in search_text_result["matches"][0]["preview"]
    search_text_tool_calls = client.get(
        f"/api/tasks/{search_text_payload['runtime_task_id']}/tool-calls"
    )
    assert search_text_tool_calls.status_code == 200, search_text_tool_calls.text
    search_text_tool_call = search_text_tool_calls.json()["tool_calls"][0]
    assert search_text_tool_call["status"] == "completed"
    assert search_text_tool_call["tool_name"] == "document.search_text"

    missing_context_plan = WorkflowPlan(
        workflow_name="runtime_missing_document_context_plan",
        description="验证前置文档搜索无命中时，提取要求不会假装成功。",
        summary="搜索无结果后尝试提取要求。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="document_agent",
                action="search_text",
                title="搜索不存在的内容",
                input={"query": "这段内容在验证文档中不存在", "limit": 10},
                reason="先制造一个无命中的文档搜索结果。",
                expected_output="没有命中文档。",
            ),
            WorkflowStep(
                id="step_2",
                agent="document_agent",
                action="extract_requirements",
                title="尝试归纳无命中文档",
                depends_on=["step_1"],
                reason="验证缺失文档上下文时会结构化失败。",
                expected_output="应返回 missing_document_context。",
            ),
        ],
    )
    missing_context_dry_run = run_workflow_dry_run(
        task_id="verify_runtime_missing_document_context",
        plan=missing_context_plan,
        available_agents=list_agents(),
    )
    assert missing_context_dry_run.status == "completed"
    missing_context_execute = client.post(f"/api/tasks/{missing_context_dry_run.task_id}/execute")
    assert missing_context_execute.status_code == 200, missing_context_execute.text
    missing_context_payload = missing_context_execute.json()
    assert missing_context_payload["status"] == "failed"
    missing_context_run = missing_context_payload["workflow_run"]
    assert missing_context_run["steps"][0]["status"] == "completed"
    assert missing_context_run["steps"][1]["status"] == "failed"
    assert (
        missing_context_run["steps"][1]["output"]["error"]["code"]
        == "missing_document_context"
    )
    assert missing_context_run["metrics"]["retry_total"] == 0
    missing_context_tool_calls = client.get(
        f"/api/tasks/{missing_context_payload['runtime_task_id']}/tool-calls"
    )
    assert missing_context_tool_calls.status_code == 200, missing_context_tool_calls.text
    missing_context_tool_payload = missing_context_tool_calls.json()["tool_calls"]
    assert missing_context_tool_payload[1]["tool_name"] == "document.extract_requirements"
    assert missing_context_tool_payload[1]["status"] == "failed"
    assert (
        missing_context_tool_payload[1]["result"]["error"]["code"]
        == "missing_document_context"
    )

    search_then_extract_plan = WorkflowPlan(
        workflow_name="runtime_search_then_extract_plan",
        description="验证 document.extract_requirements 可以读取前置搜索上下文。",
        summary="先搜索再归纳文档线索。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="document_agent",
                action="search_text",
                title="搜索受控文档",
                input={"query": "成功路径", "limit": 10, "auto_read_if_unique": True},
                reason="先定位相关文本。",
                expected_output="返回命中文档、行号和短上下文。",
            ),
            WorkflowStep(
                id="step_2",
                agent="document_agent",
                action="extract_requirements",
                title="归纳搜索结果",
                depends_on=["step_1"],
                reason="基于搜索命中提炼后续可用信息。",
                expected_output="结构化搜索结果摘要。",
            ),
            WorkflowStep(
                id="step_3",
                agent="code_agent",
                action="generate_code",
                title="生成代码草稿",
                depends_on=["step_2"],
                reason="验证 Code Agent 能消费 Document Agent 的短上下文。",
                expected_output="生成带来源上下文的代码草稿。",
            ),
            WorkflowStep(
                id="step_4",
                agent="report_agent",
                action="generate_report",
                title="生成报告草稿",
                depends_on=["step_2", "step_3"],
                reason="验证 Report Agent 能汇总文档上下文和工作流步骤。",
                expected_output="生成带来源上下文的 Markdown 报告。",
            ),
        ],
    )
    search_then_extract_dry_run = run_workflow_dry_run(
        task_id="verify_runtime_search_then_extract",
        plan=search_then_extract_plan,
        # 这是遗留 Code/Report 节点的底层回归，不经 Commander 客户路由；显式使用
        # 临时 runtime_ready 快照，避免当前产品 manifest 的未上架状态影响执行器测试。
        available_agents=legacy_runtime_agents,
    )
    assert search_then_extract_dry_run.status == "completed", search_then_extract_dry_run.validation_errors
    search_then_extract_execute = client.post(
        f"/api/tasks/{search_then_extract_dry_run.task_id}/execute"
    )
    assert search_then_extract_execute.status_code == 200, search_then_extract_execute.text
    search_then_extract_payload = search_then_extract_execute.json()
    assert search_then_extract_payload["status"] == "completed"
    search_then_extract_run = search_then_extract_payload["workflow_run"]
    assert search_then_extract_run["steps"][1]["status"] == "completed"
    extracted_context = search_then_extract_run["steps"][1]["output"]["result"]["context"]
    assert extracted_context["source_steps"] == ["step_1"]
    assert extracted_context["search_match_total"] >= 1
    assert extracted_context["read_preview_total"] == 1
    assert extracted_context["search_matches"][0]["document_name"] == sample_document.name
    assert "成功路径" in extracted_context["search_matches"][0]["preview"]
    assert extracted_context["read_previews"][0]["relative_path"] == sample_document.name
    assert "path" not in extracted_context["read_previews"][0]
    assert "AgentFlow 受控文档" in extracted_context["read_previews"][0]["preview"]
    assert str(settings.data_dir) not in json.dumps(extracted_context, ensure_ascii=False)
    assert search_then_extract_run["steps"][2]["status"] == "completed"
    assert search_then_extract_run["steps"][3]["status"] == "completed"
    code_result = search_then_extract_run["steps"][2]["output"]["result"]
    report_result = search_then_extract_run["steps"][3]["output"]["result"]
    code_context = code_result["document_context"]
    report_context = report_result["document_context"]
    assert code_context["read_preview_total"] == 1
    assert report_context["search_match_total"] >= 1
    assert str(settings.data_dir) not in json.dumps(code_context, ensure_ascii=False)
    assert str(settings.data_dir) not in json.dumps(report_context, ensure_ascii=False)
    assert code_result["verification"]["ok"] is True
    assert code_result["verification"]["checked_snippets"] >= 3
    assert report_result["verification"]["ok"] is True
    assert report_result["verification"]["checked_snippets"] >= 4
    code_context_path = Path(code_result["output_file"])
    report_context_path = Path(report_result["output_file"])
    code_text = code_context_path.read_text(encoding="utf-8")
    assert "成功路径" in code_text
    assert str(settings.data_dir) not in code_text
    report_text = report_context_path.read_text(encoding="utf-8")
    assert "## 文档上下文" in report_text
    assert sample_document.name in report_text
    assert "AgentFlow 受控文档" in report_text
    assert str(settings.data_dir) not in report_text
    search_then_extract_updates = client.get(
        f"/api/tasks/{search_then_extract_payload['runtime_task_id']}/updates"
    )
    assert search_then_extract_updates.status_code == 200, search_then_extract_updates.text
    search_then_extract_updates_payload = search_then_extract_updates.json()
    context_artifact_updates = [
        update
        for update in search_then_extract_updates_payload["updates"]
        if update["event"] == "artifact_created"
        and update["step_id"] in {"step_3", "step_4"}
    ]
    assert len(context_artifact_updates) == 2
    for update in context_artifact_updates:
        assert "step" in update["payload"]
        assert "tool_calls" in update["payload"]
        payload_context = update["payload"]["step"]["output"]["result"]["document_context"]
        tool_context = update["payload"]["tool_calls"][0]["result"]["document_context"]
        assert payload_context["search_match_total"] >= 1
        assert payload_context["read_preview_total"] == 1
        assert tool_context["source_steps"] == ["step_1"]

    timeout_plan = WorkflowPlan(
        workflow_name="runtime_timeout_plan",
        description="验证安全工具超时会进入 failed。",
        summary="document.read_text 立即超时。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="document_agent",
                action="read_text",
                title="读取但立即超时",
                input={"path": sample_document.name, "timeout_ms": 0},
                reason="验证 Runtime 超时结构化记录。",
                expected_output="不会产生文档摘要。",
            )
        ],
    )
    timeout_dry_run = run_workflow_dry_run(
        task_id="verify_runtime_timeout",
        plan=timeout_plan,
        available_agents=list_agents(),
    )
    assert timeout_dry_run.status == "completed"
    timeout_execute = client.post(f"/api/tasks/{timeout_dry_run.task_id}/execute")
    assert timeout_execute.status_code == 200, timeout_execute.text
    timeout_payload = timeout_execute.json()
    assert timeout_payload["status"] == "failed"
    timeout_run = timeout_payload["workflow_run"]
    assert timeout_run["steps"][0]["status"] == "failed"
    assert timeout_run["steps"][0]["output"]["error"]["code"] == "tool_timeout"
    assert timeout_run["steps"][0]["output"]["error"]["retryable"] is True
    assert timeout_run["metrics"]["retry_total"] == 2
    timeout_tool_calls = client.get(
        f"/api/tasks/{timeout_payload['runtime_task_id']}/tool-calls"
    )
    assert timeout_tool_calls.status_code == 200, timeout_tool_calls.text
    timeout_tool_call = timeout_tool_calls.json()["tool_calls"][0]
    assert timeout_tool_call["status"] == "failed"
    assert timeout_tool_call["tool_name"] == "document.read_text"
    assert timeout_tool_call["timeout_ms"] == 0
    assert timeout_tool_call["attempt"] == 3
    assert timeout_tool_call["failure_count"] == 3
    assert timeout_tool_call["result"]["error"]["code"] == "tool_timeout"
    assert timeout_tool_call["result"]["error"]["retryable"] is True
    assert timeout_tool_call["result"]["error"]["attempt"] == 3
    timeout_logs = client.get(f"/api/tasks/{timeout_payload['runtime_task_id']}/logs")
    assert timeout_logs.status_code == 200, timeout_logs.text
    assert any(event["event"] == "step_retried" for event in timeout_logs.json()["events"])
    timeout_evaluation = client.get(f"/api/tasks/{timeout_payload['runtime_task_id']}/evaluation")
    assert timeout_evaluation.status_code == 200, timeout_evaluation.text
    timeout_evaluation_payload = timeout_evaluation.json()
    assert timeout_evaluation_payload["outcome"] == "failed"
    assert timeout_evaluation_payload["retry_total"] == 2
    assert timeout_evaluation_payload["failed_tool_calls"] == 1
    assert any("自动重试" in warning for warning in timeout_evaluation_payload["warnings"])

    failing_runtime_plan = WorkflowPlan(
        workflow_name="runtime_failure_plan",
        description="验证安全工具参数错误会进入 failed。",
        summary="缺少 document.read_text 所需 input.path。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="document_agent",
                action="read_text",
                title="读取缺少路径的文档",
                reason="验证 Runtime 结构化失败记录。",
                expected_output="不会产生文档摘要。",
            )
        ],
    )
    failing_dry_run = run_workflow_dry_run(
        task_id="verify_runtime_failure",
        plan=failing_runtime_plan,
        available_agents=list_agents(),
    )
    assert failing_dry_run.status == "completed"
    failing_execute = client.post(f"/api/tasks/{failing_dry_run.task_id}/execute")
    assert failing_execute.status_code == 200, failing_execute.text
    failing_payload = failing_execute.json()
    assert failing_payload["status"] == "failed"
    failing_run = failing_payload["workflow_run"]
    assert failing_run["status"] == "failed"
    assert failing_run["steps"][0]["status"] == "failed"
    assert failing_run["steps"][0]["output"]["error"]["code"] == "invalid_parameters"
    assert failing_run["steps"][0]["output"]["error"]["retryable"] is False
    assert failing_run["steps"][0]["output"]["error"]["max_attempts"] == 3
    assert failing_run["metrics"]["step_failed"] == 1
    assert failing_run["metrics"]["tool_call_failed"] == 1
    failing_tool_calls = client.get(
        f"/api/tasks/{failing_payload['runtime_task_id']}/tool-calls"
    )
    assert failing_tool_calls.status_code == 200, failing_tool_calls.text
    failing_tool_call = failing_tool_calls.json()["tool_calls"][0]
    assert failing_tool_call["status"] == "failed"
    assert failing_tool_call["tool_name"] == "document.read_text"
    assert failing_tool_call["failure_count"] == 1
    assert failing_tool_call["attempt"] == 1
    assert failing_tool_call["max_attempts"] == 3
    assert failing_tool_call["result"]["error"]["code"] == "invalid_parameters"
    assert failing_tool_call["result"]["error"]["retryable"] is False
    failing_logs = client.get(f"/api/tasks/{failing_payload['runtime_task_id']}/logs")
    assert failing_logs.status_code == 200, failing_logs.text
    assert any(event["event"] == "step_failed" for event in failing_logs.json()["events"])
    assert any(event["event"] == "task_failed" for event in failing_logs.json()["events"])
    failing_evaluation = client.get(f"/api/tasks/{failing_payload['runtime_task_id']}/evaluation")
    assert failing_evaluation.status_code == 200, failing_evaluation.text
    failing_evaluation_payload = failing_evaluation.json()
    assert failing_evaluation_payload["outcome"] == "failed"
    assert failing_evaluation_payload["step_success_rate"] == 0.0
    assert failing_evaluation_payload["tool_success_rate"] == 0.0
    assert failing_evaluation_payload["failed_tool_calls"] == 1
    assert any("工具调用失败" in warning for warning in failing_evaluation_payload["warnings"])

    invalid_plan = WorkflowPlan(
        workflow_name="broken_plan",
        description="用于验证计划校验器能拒绝坏依赖。",
        summary="故意构造的坏计划。",
        steps=[
            WorkflowStep(
                id="step_1",
                agent="commander_agent",
                action="analyze_task",
                title="分析用户任务",
                depends_on=["missing_step"],
                reason="测试坏依赖。",
                expected_output="校验器错误。",
            )
        ],
    )
    validation_errors = validate_workflow_plan(invalid_plan, available_agents=list_agents())
    assert any("依赖不存在" in error for error in validation_errors)

    with client.websocket_connect(f"/ws/tasks/{task_id}") as websocket:
        expected_event_count = 2 + len(workflow_run["steps"]) * 2 + 1 + confirmation_step_count
        dry_run_events = [websocket.receive_json() for _ in range(expected_event_count)]
        assert dry_run_events[0]["task_id"] == task_id
        assert dry_run_events[1]["agent_id"] == "workflow_engine"
        assert dry_run_events[1]["event"] == "task_started"
        assert any(event["event"] == "confirmation_required" for event in dry_run_events)
        assert any(event["event"] == "step_completed" for event in dry_run_events)
        assert dry_run_events[-1]["event"] == "task_completed"
        assert "dry-run" in dry_run_events[-1]["message"]

    with client.websocket_connect("/ws/tasks/demo") as websocket:
        events = [websocket.receive_json() for _ in range(5)]
        assert events[0]["task_id"] == "demo"
        assert events[0]["event"] == "connected"
        assert events[-1]["event"] == "task_completed"

    print("AgentFlow backend verification passed.")


if __name__ == "__main__":
    main()
