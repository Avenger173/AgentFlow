"""知识库 K2 固定资料集的关键词/本地 Hybrid 端到端质量与延迟基线。

本脚本复用 K0 的脱敏资料和标注问题，但走 K1/K2 的真实受控导入、generation、FTS、父块
证据回读链。默认只测关键词路径；``--with-local-dense`` 仅复用已经由客户确认准备好的本机
Embedding 缓存，并把 SQLite/Chroma 索引继续隔离在系统临时目录。它绝不下载模型、读取客户
资料或调用云端 LLM。
"""

from __future__ import annotations

import argparse
import base64
import gc
from io import BytesIO
import json
import os
from pathlib import Path
import shutil
import statistics
import sys
import tempfile
import time
from collections import defaultdict
from dataclasses import dataclass

import fitz
from docx import Document as DocxDocument

BACKEND_ROOT = Path(__file__).resolve().parents[1]
FIXTURE_DIR = BACKEND_ROOT / "scripts" / "fixtures" / "knowledge_k0"
QUALITY_FIXTURE_PATH = (
    BACKEND_ROOT / "scripts" / "fixtures" / "knowledge_k6" / "retrieval_quality_cases.json"
)
VERIFY_ROOT = Path(tempfile.mkdtemp(prefix="agentflow_knowledge_retrieval_baseline_"))
os.environ["AGENTFLOW_DATA_DIR"] = str(VERIFY_ROOT / "data")
os.environ["AGENTFLOW_DATABASE_PATH"] = str(VERIFY_ROOT / "data" / "knowledge_retrieval_baseline.db")
os.environ["AGENTFLOW_CHAT_MODE"] = "mock"
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from app.database import sqlite as sqlite_service
from app.database.knowledge_repository import create_knowledge_base, import_workspace_documents_to_knowledge_base
from app.schemas.knowledge import KnowledgeRetrievalRequest
from app.services.knowledge_keyword_index import create_knowledge_index_job, run_knowledge_index_job
from app.services.knowledge_retrieval import retrieve_knowledge_evidence
from app.services.workspace_documents import import_workspace_document, import_workspace_document_base64


@dataclass(frozen=True)
class _QualityEvaluationCase:
    """K6 固定评测的一题脱敏元数据。

    题集只包含仓库内可审查的合成问题和稳定资料 ID，绝不收集客户问题、文件名或正文。
    ``required`` 用于守住既有能力；``diagnostic`` 用于暴露候选策略是否值得实验。后者的
    当前缺口不会让回归失败，否则未来策略改善时反而会被旧预期束缚。
    """

    case_id: str
    question: str
    expected_document_ids: frozenset[str]
    evaluation_group: str
    quality_gate: str
    minimum_expected_coverage: int | None
    expected_source_kind: str | None
    known_gap_candidate: bool


def _load_k6_quality_cases() -> list[_QualityEvaluationCase]:
    """读取并严格校验 K6 版本化质量夹具，避免评测自己静默漏题。"""

    raw_cases = json.loads(QUALITY_FIXTURE_PATH.read_text(encoding="utf-8"))
    if not isinstance(raw_cases, list) or not raw_cases:
        raise ValueError("K6 质量评测夹具必须是非空列表。")

    cases: list[_QualityEvaluationCase] = []
    seen_ids: set[str] = set()
    for raw_case in raw_cases:
        if not isinstance(raw_case, dict):
            raise ValueError("K6 质量评测夹具中的每一项必须是对象。")
        case_id = str(raw_case.get("id", "")).strip()
        question = str(raw_case.get("question", "")).strip()
        expected_raw = raw_case.get("expected_document_ids")
        evaluation_group = str(raw_case.get("evaluation_group", "")).strip()
        quality_gate = str(raw_case.get("quality_gate", "")).strip()
        if not case_id or case_id in seen_ids or not question or not evaluation_group:
            raise ValueError("K6 质量评测夹具包含缺失或重复的题目元数据。")
        if quality_gate not in {"required", "diagnostic"}:
            raise ValueError(f"K6 题目 {case_id} 的 quality_gate 必须是 required 或 diagnostic。")
        if not isinstance(expected_raw, list) or any(not isinstance(item, str) or not item for item in expected_raw):
            raise ValueError(f"K6 题目 {case_id} 的 expected_document_ids 非法。")
        expected_document_ids = frozenset(expected_raw)
        declared_coverage = raw_case.get("minimum_expected_coverage")
        minimum_expected_coverage = int(declared_coverage) if declared_coverage is not None else None
        if minimum_expected_coverage is not None and (
            minimum_expected_coverage < 1 or minimum_expected_coverage > len(expected_document_ids)
        ):
            raise ValueError(f"K6 题目 {case_id} 的 minimum_expected_coverage 超出预期资料范围。")
        expected_source_kind = str(raw_case.get("expected_source_kind") or "").strip() or None
        seen_ids.add(case_id)
        cases.append(
            _QualityEvaluationCase(
                case_id=case_id,
                question=question,
                expected_document_ids=expected_document_ids,
                evaluation_group=evaluation_group,
                quality_gate=quality_gate,
                minimum_expected_coverage=minimum_expected_coverage,
                expected_source_kind=expected_source_kind,
                known_gap_candidate=bool(raw_case.get("known_gap_candidate", False)),
            )
        )

    if not any(case.quality_gate == "required" for case in cases):
        raise ValueError("K6 质量评测至少要保留一个 required 回归题。")
    if not any(case.known_gap_candidate for case in cases):
        raise ValueError("K6 质量评测至少要保留一个脱敏失败模式夹具。")
    return cases


def _configure_existing_local_embedding_cache() -> None:
    """为显式 Hybrid 基准复用已确认模型，且不允许测试触发下载。

    ``AGENTFLOW_DATA_DIR`` 已在模块加载时指向临时目录，因此必须单独恢复客户先前确认的模型
    缓存位置。向量 generation 目录仍留在临时 ``data`` 下，脚本结束会整体清理；本函数只检查
    ready marker，绝不创建、下载或修改正式模型缓存。
    """

    configured_cache = os.getenv("AGENTFLOW_KNOWLEDGE_EMBEDDING_CACHE_DIR")
    cache_dir = Path(configured_cache).resolve() if configured_cache else (
        BACKEND_ROOT.parent / "data" / "knowledge_embedding_models"
    ).resolve()
    marker_path = cache_dir / "bge_small_zh_v1_5.ready"
    if not marker_path.is_file():
        raise RuntimeError(
            "--with-local-dense 只允许复用已确认的本地 Embedding 模型；"
            "当前缓存没有 ready 标记，请先在知识库页面完成客户确认的模型准备。"
        )
    os.environ["AGENTFLOW_KNOWLEDGE_EMBEDDING_CACHE_DIR"] = str(cache_dir)
    # 显式指定是为了避免调用方预设的正式向量目录意外被本基准复用。
    os.environ["AGENTFLOW_KNOWLEDGE_VECTOR_STORAGE_DIR"] = str(VERIFY_ROOT / "data" / "knowledge_vectors")


def _import_binary_format_fixtures() -> list[str]:
    """构造最小可提取 PDF/DOCX，复用客户文件的二进制受控导入协议。

    二进制夹具只写进本脚本的临时 workspace，避免把难审查的二进制文件提交到仓库；其内容
    使用稳定标识符，确保本回归验证的是解析、来源锚点和检索链，而不是某个模型的语义能力。
    """

    pdf = fitz.open()
    try:
        page = pdf.new_page()
        page.insert_text(
            (72, 72),
            "Release control RELEASE-77 requires an owner approval before production deployment.",
        )
        pdf_bytes = pdf.tobytes()
    finally:
        pdf.close()
    pdf_name = "retrieval_pdf_baseline.pdf"
    import_workspace_document_base64(
        filename=pdf_name,
        content_base64=base64.b64encode(pdf_bytes).decode("ascii"),
    )

    docx = DocxDocument()
    docx.add_paragraph("Data handoff package DS-19 must include a schema version and owner confirmation.")
    docx.add_paragraph("The acceptance owner reviews the handoff package within five business days.")
    docx_bytes = BytesIO()
    docx.save(docx_bytes)
    docx_name = "retrieval_docx_baseline.docx"
    import_workspace_document_base64(
        filename=docx_name,
        content_base64=base64.b64encode(docx_bytes.getvalue()).decode("ascii"),
    )
    return [pdf_name, docx_name]


def _import_long_markdown_fixture() -> list[str]:
    """生成跨父块长度的 Markdown，验证短编号仍能回读后段来源。

    长材料只用于评测分块与检索边界；它不模拟客户私密原文，也不要求模型理解重复背景内容。
    """

    background = "\n\n".join(
        f"第 {index} 段背景：本段说明常规交付背景、参与角色和既有流程，不包含验收编号。"
        + " 该说明用于形成足够长的章节上下文，同时保持本题的唯一命中词不出现在背景中。" * 8
        for index in range(1, 28)
    )
    content = (
        "# 长篇交付操作手册\n\n"
        f"{background}\n\n"
        "## 最终验收\n\n"
        "编号 MANUAL-11 要求在交付前完成负责人签字并保留可追溯记录。\n"
    )
    filename = "retrieval_long_baseline.md"
    import_workspace_document(filename=filename, content=content)
    return [filename]


def _percentile_95(samples: list[float]) -> float:
    """小样本采用最近秩 P95，输出只用于本机基线，不伪装成生产 SLA。"""

    ordered = sorted(samples)
    index = max(0, min(len(ordered) - 1, int(len(ordered) * 0.95 + 0.999999) - 1))
    return ordered[index]


def _evaluate_k6_quality_cases(*, knowledge_base_id: str) -> dict[str, object]:
    """运行 K6 质量题集，并输出供后续策略对比的无正文指标报告。

    这不是第二条产品检索链路：每题仍调用当前 ``retrieve_knowledge_evidence``。它的职责是将
    “某个新策略看起来不错”的主观印象转化为同一题集上的 Recall、MRR、覆盖与拒答证据。
    新策略若不能改善 diagnostic 缺口，或让 required 回归退化，就不应进入默认路径。
    """

    cases = _load_k6_quality_cases()
    group_rows: dict[str, list[dict[str, object]]] = defaultdict(list)
    required_failures: list[str] = []
    known_gap_case_ids: list[str] = []
    retrieval_modes: set[str] = set()
    latencies_ms: list[float] = []

    for case in cases:
        started = time.perf_counter()
        result = retrieve_knowledge_evidence(
            KnowledgeRetrievalRequest(
                knowledge_base_id=knowledge_base_id,
                query=case.question,
                top_k=5,
            )
        )
        latencies_ms.append((time.perf_counter() - started) * 1000)
        retrieval_modes.add(result.diagnostics.mode)
        actual_document_ids = [Path(item.document_name).stem for item in result.evidences]
        matched_document_ids = case.expected_document_ids.intersection(actual_document_ids)
        minimum_coverage = case.minimum_expected_coverage or (1 if case.expected_document_ids else 0)
        recall_at_5 = (
            len(matched_document_ids) / len(case.expected_document_ids) if case.expected_document_ids else None
        )
        matching_ranks = [
            rank
            for rank, document_id in enumerate(actual_document_ids, start=1)
            if document_id in case.expected_document_ids
        ]
        reciprocal_rank = 1.0 / matching_ranks[0] if matching_ranks else 0.0
        source_anchor_matched = not case.expected_source_kind or any(
            Path(item.document_name).stem in case.expected_document_ids
            and item.source.source_kind == case.expected_source_kind
            for item in result.evidences
        )
        if not case.expected_document_ids:
            passed = not actual_document_ids
        else:
            passed = len(matched_document_ids) >= minimum_coverage and source_anchor_matched

        group_rows[case.evaluation_group].append(
            {
                "passed": passed,
                "recall_at_5": recall_at_5,
                "reciprocal_rank": reciprocal_rank if case.expected_document_ids else None,
                "is_required": case.quality_gate == "required",
            }
        )
        if case.quality_gate == "required" and not passed:
            # 只报告合成夹具 ID 与判断条件，避免把检索结果正文或题干带进自动化日志。
            required_failures.append(case.case_id)
        if case.known_gap_candidate and not passed:
            known_gap_case_ids.append(case.case_id)

    assert not required_failures, (
        "K6 质量评测 required 回归失败：" + ", ".join(required_failures)
    )

    group_report: dict[str, dict[str, object]] = {}
    for group_name, rows in sorted(group_rows.items()):
        recall_values = [float(row["recall_at_5"]) for row in rows if row["recall_at_5"] is not None]
        mrr_values = [float(row["reciprocal_rank"]) for row in rows if row["reciprocal_rank"] is not None]
        group_report[group_name] = {
            "case_count": len(rows),
            "passed_case_count": sum(bool(row["passed"]) for row in rows),
            "required_case_count": sum(bool(row["is_required"]) for row in rows),
            "mean_recall_at_5": round(statistics.fmean(recall_values), 3) if recall_values else None,
            "mrr": round(statistics.fmean(mrr_values), 3) if mrr_values else None,
        }

    return {
        "suite": "knowledge_k6_quality_v1",
        "case_count": len(cases),
        "required_case_count": sum(case.quality_gate == "required" for case in cases),
        "diagnostic_case_count": sum(case.quality_gate == "diagnostic" for case in cases),
        "known_gap_case_ids": known_gap_case_ids,
        "retrieval_modes": sorted(retrieval_modes),
        "median_ms": round(statistics.median(latencies_ms), 2),
        "p95_ms": round(_percentile_95(latencies_ms), 2),
        "groups": group_report,
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="运行知识库 K2 关键词或本地 Hybrid 基线。")
    parser.add_argument(
        "--with-local-dense",
        action="store_true",
        help="只读复用已确认的本机 Embedding 缓存，验证真实 Hybrid；不会下载模型。",
    )
    parser.add_argument(
        "--k6-quality",
        action="store_true",
        help="追加运行 K6 脱敏质量题集，输出 required 回归与 diagnostic 缺口指标。",
    )
    parser.add_argument(
        "--require-k6-diagnostic-pass",
        action="store_true",
        help="要求 K6 已标注的失败模式全部通过；仅用于候选策略验收，不改变默认基线。",
    )
    arguments = parser.parse_args()
    if arguments.require_k6_diagnostic_pass and not arguments.k6_quality:
        parser.error("--require-k6-diagnostic-pass 必须与 --k6-quality 一起使用。")
    if arguments.with_local_dense:
        _configure_existing_local_embedding_cache()
    try:
        question_set = json.loads((FIXTURE_DIR / "question_set.json").read_text(encoding="utf-8"))
        base = create_knowledge_base(name="K2 固定检索基线")
        document_names: list[str] = []
        for fixture_path in sorted(FIXTURE_DIR.glob("*.md")):
            import_workspace_document(filename=fixture_path.name, content=fixture_path.read_text(encoding="utf-8"))
            document_names.append(fixture_path.name)
        document_names.extend(_import_binary_format_fixtures())
        document_names.extend(_import_long_markdown_fixture())
        import_workspace_documents_to_knowledge_base(
            knowledge_base_id=base.knowledge_base_id,
            workspace_document_names=document_names,
        )
        completed = run_knowledge_index_job(create_knowledge_index_job(base.knowledge_base_id).index_job_id)
        assert completed.status == "completed"

        latencies_ms: list[float] = []
        required_passed = 0
        reciprocal_ranks: list[float] = []
        recall_values: list[float] = []
        no_answer_passed = 0
        format_required_passed = 0
        retrieval_modes: set[str] = set()
        for question in question_set:
            started = time.perf_counter()
            result = retrieve_knowledge_evidence(
                KnowledgeRetrievalRequest(
                    knowledge_base_id=base.knowledge_base_id,
                    query=str(question["question"]),
                    top_k=5,
                )
            )
            latencies_ms.append((time.perf_counter() - started) * 1000)
            retrieval_modes.add(result.diagnostics.mode)
            expected = set(question["expected_document_ids"])
            actual = [Path(item.document_name).stem for item in result.evidences]
            if not expected:
                assert not actual, f"无答案题 {question['id']} 不应命中材料：{actual}"
                no_answer_passed += 1
                continue
            matching_ranks = [index for index, document_id in enumerate(actual, start=1) if document_id in expected]
            covered_expected = expected.intersection(actual)
            recall_values.append(len(expected.intersection(actual)) / len(expected))
            reciprocal_ranks.append(1.0 / matching_ranks[0] if matching_ranks else 0.0)
            declared_coverage = question.get("minimum_expected_coverage")
            if declared_coverage is not None:
                minimum_expected_coverage = int(declared_coverage)
                assert len(covered_expected) >= minimum_expected_coverage, (
                    f"覆盖题 {question['id']} 仅召回 {sorted(covered_expected)}，"
                    f"低于预期资料覆盖 {minimum_expected_coverage}。"
                )
            expected_source_kind = str(question.get("expected_source_kind") or "")
            if expected_source_kind:
                assert any(
                    Path(item.document_name).stem in expected and item.source.source_kind == expected_source_kind
                    for item in result.evidences
                ), f"格式题 {question['id']} 未保留 {expected_source_kind} 来源锚点。"
                format_required_passed += 1
            if question["baseline_required"]:
                assert matching_ranks, f"K2 基线未召回必过题 {question['id']} 的预期资料。"
                required_passed += 1

        required_total = sum(1 for item in question_set if item["baseline_required"])
        format_required_total = sum(1 for item in question_set if item.get("expected_source_kind"))
        if arguments.with_local_dense:
            # 此开关的目的就是验证真实向量 generation 与 RRF 路径；若悄悄降级，不能把结果
            # 当作 Hybrid 质量证据。无答案题也可能在阈值过滤后报告 no_result，故只要求至少有
            # 一个含证据查询真正走 Hybrid。
            assert "hybrid" in retrieval_modes, "本地 Dense 基准未进入 Hybrid 路径，不能记录为语义评测。"
        if arguments.k6_quality:
            quality_report = _evaluate_k6_quality_cases(knowledge_base_id=base.knowledge_base_id)
            # JSON 只包含合成 case ID、分组聚合与耗时；便于未来实验脚本稳定比对，且没有
            # 客户问题、文件名、正文、向量或绝对路径。
            print(
                "Knowledge K6 quality evaluation passed: "
                + json.dumps(quality_report, ensure_ascii=False, sort_keys=True)
            )
            if arguments.require_k6_diagnostic_pass:
                known_gap_case_ids = list(quality_report["known_gap_case_ids"])
                assert not known_gap_case_ids, (
                    "候选检索策略没有覆盖 K6 已标注的失败模式：" + ", ".join(known_gap_case_ids)
                )
        print(
            "Knowledge K2 retrieval baseline passed: "
            f"questions={len(question_set)} required_recall_at_5={required_passed}/{required_total} "
            f"format_anchor={format_required_passed}/{format_required_total} "
            f"mean_recall_at_5={statistics.fmean(recall_values):.3f} "
            f"mrr={statistics.fmean(reciprocal_ranks):.3f} no_answer={no_answer_passed}/1 "
            f"median_ms={statistics.median(latencies_ms):.2f} p95_ms={_percentile_95(latencies_ms):.2f} "
            f"modes={','.join(sorted(retrieval_modes))} local_dense={str(arguments.with_local_dense).lower()} "
            "network=false model_download=false"
        )
    finally:
        sqlite_service._INITIALIZED_PATHS.clear()
        # 异常断言会暂时保留局部 SQLite 引用；Windows 下先释放并有限重试，避免把夹具
        # 自己的失败掩盖成“数据库文件无法删除”。
        for _ in range(20):
            gc.collect()
            try:
                shutil.rmtree(VERIFY_ROOT, ignore_errors=False)
                break
            except PermissionError:
                time.sleep(0.05)
        else:
            shutil.rmtree(VERIFY_ROOT, ignore_errors=False)


if __name__ == "__main__":
    main()
